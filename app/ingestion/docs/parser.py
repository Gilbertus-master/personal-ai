from dataclasses import dataclass
from pathlib import Path
from datetime import datetime
from typing import Optional

from pypdf import PdfReader
from docx import Document


@dataclass
class ParsedDocument:
    title: str
    created_at: Optional[datetime]
    author: Optional[str]
    participants: list[str]
    raw_path: str
    text: str
    file_type: str


def read_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore").strip()


def read_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    parts = []

    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            parts.append(text.strip())

    return "\n\n".join(parts).strip()


def read_docx(path: Path) -> str:
    doc = Document(str(path))
    parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    parts.append(text)

    return "\n".join(parts).strip()


def read_doc(path: Path) -> str:
    raise ValueError(f"Unsupported file format .doc — convert to DOCX first: {path}")


MAX_FILE_BYTES = 100 * 1024 * 1024  # 100 MB


def parse_document_file(file_path: str | Path) -> ParsedDocument:
    path = Path(file_path)
    size = path.stat().st_size
    if size > MAX_FILE_BYTES:
        raise ValueError(f"File too large to parse: {size} bytes ({path.name})")
    suffix = path.suffix.lower()

    if suffix == ".txt":
        text = read_txt(path)
        file_type = "txt"
    elif suffix == ".pdf":
        text = read_pdf(path)
        file_type = "pdf"
    elif suffix == ".docx":
        text = read_docx(path)
        file_type = "docx"
    elif suffix == ".doc":
        text = read_doc(path)
        file_type = "doc"
    else:
        raise ValueError(f"Unsupported file type: {suffix}")

    created_at = datetime.fromtimestamp(path.stat().st_mtime)

    return ParsedDocument(
        title=path.name,
        created_at=created_at,
        author=None,
        participants=[],
        raw_path=str(path),
        text=text,
        file_type=file_type,
    )