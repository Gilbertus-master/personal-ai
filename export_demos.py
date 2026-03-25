from docx import Document
import re

for src, dst in [
    ("docs/DEMO_SCRIPT_ROCH.md", "/mnt/c/Users/jablo/Desktop/Demo_Roch_Czwartek.docx"),
    ("docs/DEMO_SCRIPT_KRYSTIAN.md", "/mnt/c/Users/jablo/Desktop/Demo_Krystian_Piatek.docx"),
]:
    with open(src, "r") as f:
        lines = f.readlines()

    doc = Document()
    for line in lines:
        line = line.rstrip()
        if not line:
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:], 0)
        elif line.startswith("## "):
            doc.add_heading(line[3:], 1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], 2)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.startswith("```"):
            continue
        elif re.match(r"^\d+\.", line):
            doc.add_paragraph(re.sub(r"^\d+\.\s*", "", line), style="List Number")
        else:
            p = doc.add_paragraph()
            parts = line.split("**")
            for i, part in enumerate(parts):
                if part:
                    run = p.add_run(part)
                    if i % 2 == 1:
                        run.bold = True

    doc.save(dst)
    print(f"Saved: {dst.split('/')[-1]}")
