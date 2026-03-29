#!/usr/bin/env python3
"""Generate Gilbertus Masterplan V5 as DOCX."""
import structlog
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT

log = structlog.get_logger(__name__)

doc = Document()

# Styles
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    return h

def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)
            for p in row_cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    return table

# ============================================================
# TITLE
# ============================================================
title = doc.add_heading('GILBERTUS ALBANS — Masterplan V5 "Jedna Piesc"', 0)
for run in title.runs:
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

doc.add_paragraph('Zaktualizowany: 2026-03-27 22:30 CET')
doc.add_paragraph('Projekt: Gilbertus Albans — prywatny mentat AI Sebastiana Jablonskiego')
doc.add_paragraph('')

# ============================================================
# STATUS 5 ZDOLNOSCI
# ============================================================
add_heading('Status 5 zdolnosci dowodcy', 1)

add_table(
    ['#', 'Zdolnosc', 'Status', 'Zrealizowane w'],
    [
        ['1', 'Petla zwrotna (dzialanie -> wynik -> nauka)', 'DONE',
         'Wave 1: Action Outcome Tracker, Decision Intelligence, Rule Reinforcement'],
        ['2', 'Autonomia wykonawcza (dzialaj bez pytania)', 'DONE',
         'Wave 1: Authority Framework (5 poziomow) + Wave 2: Adaptive Authority'],
        ['3', 'Swiadomosc finansowa (ile to kosztuje)', 'DONE',
         'Wave 3: Financial Framework + Decision Cost Estimator'],
        ['4', 'Zarzadzanie czasem (kalendarz jako zasob)', 'DONE',
         'Wave 3: Calendar Manager + Meeting ROI Tracker'],
        ['5', 'Lacznosc zwrotna (czy rozkaz wykonano?)', 'DONE',
         'Wave 2: Response Tracking, Channel Effectiveness, Delegation Chain'],
    ]
)

doc.add_paragraph('')
doc.add_paragraph('WSZYSTKIE 5 ZDOLNOSCI DOWODCY ZAADRESOWANE.', style='Intense Quote')

# ============================================================
# BASELINE
# ============================================================
add_heading('Baseline (27.03.2026)', 1)

add_table(
    ['Metryka', 'Poczatek sesji', 'Koniec sesji', 'Delta'],
    [
        ['MCP tools', '21', '36', '+15'],
        ['DB tables', '40', '64', '+24'],
        ['App modules', '76', '125', '+49'],
        ['Scripts', '60', '78', '+18'],
        ['Cron jobs', '28', '46', '+18'],
        ['API endpoints', '~21', '~95', '+~74'],
        ['Feedback loops', '0', '6', '+6'],
        ['Linie kodu (sesja)', '-', '+14,140', '4 commity'],
    ]
)

# ============================================================
# ZREALIZOWANE
# ============================================================
add_heading('Zrealizowane — pelna historia', 1)

add_heading('Sesja 23-24.03 (~20h) — Od zera do mentata', 2)
doc.add_paragraph(
    '95k chunkow z 10 zrodel (email, Teams, WhatsApp, Plaud, calendar, ChatGPT, dokumenty). '
    'RAG pipeline: query interpreter -> Qdrant search -> Claude answering. '
    'Morning brief (5 sekcji). MCP 18 tools. WhatsApp integration. Proaktywne alerty.'
)

add_heading('Sesja 26.03 (~14h) — Stabilizacja + Intelligence + Omnius', 2)
doc.add_paragraph(
    'Faza 0 (19/19): Connection pooling, Teams grouping, entity dedup, event taxonomy 15 typow, '
    'chunk dedup, prompt caching, pre-commit hooks, API cost tracking, graceful shutdown.\n'
    'Faza 1 (6/6): Calendar sync, morning brief 5 sekcji, person-aware retrieval, extraction 100%, '
    'decision journal, cross-domain correlation.\n'
    'Fazy 2-5: Evaluation pipeline, Omnius codebase (REH+REF), action pipeline, '
    'scorecard, inefficiency detector, QC agents.'
)

add_heading('Sesja 27.03 — Intelligence Layer (13 modulow)', 2)
add_table(
    ['#', 'Modul', 'Co robi'],
    [
        ['1', 'Commitment Tracker', 'Ekstrakcja obietnic z chunkow + overdue detection + fulfillment scan'],
        ['2', 'Meeting Prep Brief', 'Auto-brief 30 min przed spotkaniem na WhatsApp'],
        ['3', 'Meeting Minutes Generator', 'Structured minutes z nagran Plaud -> commitments'],
        ['4', 'Smart Response Drafter', 'Auto-odpowiedzi email/Teams/WhatsApp per standing order'],
        ['5', 'Weekly Executive Synthesis', 'Niedzielny raport strategiczny'],
        ['6', 'Sentiment Trend Monitor', 'Tygodniowy sentiment per osoba (1-5)'],
        ['7', 'Wellbeing Monitor', 'Wellbeing score Sebastiana (stress, family, health)'],
        ['8', 'Contract Intelligence', 'Ekstrakcja kontraktow + auto-alert 30/14/7 dni'],
        ['9', 'Delegation Effectiveness', 'Completion rate, on-time rate per osoba'],
        ['10', 'Knowledge Blind Spots', 'Nieznane osoby, brak docs, luki zrodel'],
        ['11', 'Communication Network', 'Graf komunikacji, silosy, bottlenecki'],
        ['12', 'Predictive Alerts', 'Predykcja eskalacji i deadline risks'],
        ['13', 'Cron Registry', 'DB-driven per-user cron control (26 jobow)'],
    ]
)

add_heading('Wave 1: Zamkniete petle zwrotne (4 moduly)', 2)
add_table(
    ['Modul', 'Co robi', 'Znaczenie'],
    [
        ['Action Outcome Tracker', 'Po 24h/72h/7d sprawdza czy akcja zadzialala',
         'Gilbertus uczy sie z wynikow akcji'],
        ['Decision Intelligence', 'Auto-capture decyzji, confidence calibration, bias detection',
         'Gilbertus uczy sie z decyzji'],
        ['Self-Rule Reinforcement', 'Effectiveness scoring, conflict detection, auto-deprecation',
         'Reguly sie wzmacniaja lub slabna'],
        ['Authority Framework', '5 poziomow (inform -> never alone), 22 kategorie',
         'Gilbertus dziala autonomicznie w ramach autoryzacji'],
    ]
)

add_heading('Wave 2: Lancuch dowodzenia (5 modulow)', 2)
add_table(
    ['Modul', 'Co robi', 'Znaczenie'],
    [
        ['Delegation Chain', 'Deleguj -> trackuj -> remind -> eskaluj -> complete',
         'Gilbertus wydaje rozkazy i trackuje wykonanie'],
        ['Response Tracking', 'Czy email przeczytany, czas, sentiment, auto-follow-up',
         'Zamyka petle komunikacyjna'],
        ['Standing Order Effectiveness', 'Response rate per order, ineffective topics',
         'Optymalizuje komunikacje'],
        ['Channel Effectiveness', 'Per-person optimal channel (email vs Teams vs WhatsApp)',
         'Wie ktory kanal uzyc z kim'],
        ['Adaptive Authority', 'Approval patterns -> sugestie zmian poziomow',
         'System sam rozszerza autonomie'],
    ]
)

add_heading('Wave 3: Strategia + Finanse + Czas (6 modulow)', 2)
add_table(
    ['Modul', 'Co robi', 'Adresuje brak'],
    [
        ['Financial Framework', 'KPI per spolka, budzety, API costs, alerts',
         '#3 Swiadomosc finansowa'],
        ['Decision Cost Estimator', 'LLM szacowanie kosztow propozycji (ROI, payback)',
         '#3 Swiadomosc finansowa'],
        ['Calendar Manager', 'Graph API ReadWrite: deep work, meeting suggestions, conflicts',
         '#4 Zarzadzanie czasem'],
        ['Meeting ROI Tracker', 'Scoring spotkan, pattern analysis, waste elimination',
         '#4 Zarzadzanie czasem'],
        ['Strategic Goal Tracker', 'Cel -> KPI -> auto-tracking -> risks -> sub-goals',
         'Strategia'],
        ['Org Health Score', '1-100 tygodniowo, 8 wymiarow',
         'Strategia'],
    ]
)

# ============================================================
# ARCHITEKTURA
# ============================================================
add_heading('Architektura techniczna', 1)

doc.add_paragraph('Stack: Python, FastAPI, PostgreSQL 16, Qdrant, Claude (Anthropic), OpenAI embeddings, Docker')
doc.add_paragraph('')

add_heading('MCP Tools (36)', 2)
add_table(
    ['Kategoria', 'Narzedzia'],
    [
        ['Core (11)', 'ask, timeline, summary, brief, alerts, status, db_stats, decide, people, lessons, costs'],
        ['Extended (7)', 'evaluate, propose_action, pending_actions, self_rules, opportunities, inefficiency, correlate'],
        ['Intelligence (6)', 'commitments, meeting_prep, sentiment, wellbeing, delegation, network'],
        ['Command (4)', 'crons, authority, decision_patterns, delegation_chain'],
        ['Feedback (2)', 'response_stats, finance'],
        ['Strategy (2)', 'goals, org_health'],
        ['Calendar (1)', 'calendar'],
        ['Omnius (3)', 'omnius_ask, omnius_command, omnius_status'],
    ]
)

add_heading('Cron Registry (26 jobow w 6 kategoriach)', 2)
add_table(
    ['Kategoria', 'Joby', 'Czestotliwosc'],
    [
        ['backup (4)', 'backup_db, backup_daytime, prune, pg_restore', '3:00 + co 4h + @reboot'],
        ['ingestion (5)', 'index_chunks, live_ingest, plaud, corporate_sync, claude_archive', 'co 5-30 min'],
        ['extraction (2)', 'turbo_extract, extract_commitments', 'co 30 min'],
        ['intelligence (6)', 'morning_brief, commitments, self_improvement, intelligence_scan, weekly_analysis, weekly_synthesis', 'daily/weekly'],
        ['communication (5)', 'meeting_prep, response_drafter, minutes, insights, task_monitor', 'co 2-30 min (8-20)'],
        ['qc (4)', 'code_quality, data_quality, session_context, quarterly_eval', 'daily + quarterly'],
    ]
)

add_heading('Baza danych (64 tabele)', 2)
doc.add_paragraph(
    'Core: documents, chunks, sources, entities, events + join tables\n'
    'People: people, relationships, open_loops, roles_history, timeline\n'
    'Intelligence: opportunities, alerts, insights, decisions, summaries, self_rules\n'
    'Commitments: commitments, chunks_commitment_checked\n'
    'Analytics: sentiment_scores, wellbeing_scores, predictive_alerts, contracts, communication_edges, meeting_minutes\n'
    'Communication: standing_orders, sent_communications, action_items, response_drafts\n'
    'Command: authority_levels, authority_log, delegation_tasks, rule_applications, action_outcomes\n'
    'Finance: financial_metrics, budget_items, financial_alerts\n'
    'Strategy: strategic_goals, goal_progress, goal_dependencies, org_health_scores\n'
    'Operations: api_costs, prompt_versions, lessons_learned, cron_registry, cron_user_assignments'
)

# ============================================================
# CO DALEJ
# ============================================================
add_heading('Co dalej', 1)

add_heading('Wave 4: Proaktywne dowodzenie (~7 dni)', 2)
add_table(
    ['Modul', 'Co robi', 'Effort'],
    [
        ['Scenario Analyzer', '"Co jesli Roch odejdzie?" — impact, replacement, financial cost', '2 dni'],
        ['Market Intelligence', 'Ceny energii, regulacje URE, przetargi OZE, KRS', '3 dni'],
        ['Competitor Intelligence', 'KRS, hiring patterns, media monitoring konkurencji', '2 dni'],
    ]
)

add_heading('Faza A: Deploy + Discovery', 2)
add_table(
    ['Task', 'Status', 'Blocker'],
    [
        ['A1: Migracja Hetzner + aktywacja cronow z registry', 'TODO', 'Sebastian (secrets)'],
        ['A2: Omnius REH deploy', 'TODO', 'Roch + IT Admin'],
        ['A3: Omnius REF deploy', 'TODO', 'Krystian + IT Admin'],
        ['A4: Onboarding Roch + Krystian (per-user crony)', 'TODO', 'A2, A3'],
        ['A5: NDA + compliance', 'TODO', 'Milosz Awedyk'],
        ['A6-A8: Discovery + cost-to-value ranking', 'TODO', 'A2, A3'],
    ]
)

add_heading('Faza C: Voice', 2)
doc.add_paragraph('C1: Voice interface (STT+TTS+WebSocket) | C3: Speaker ID | C4: Meeting boundaries')

add_heading('Faza D: Organizacja = Omnius', 2)
doc.add_paragraph('D1-D9: Infrastructure audit -> co-development -> director augmentation -> adoption monitoring')

add_heading('Faza E: Scale + Protect', 2)
doc.add_paragraph('E1: Wlasny LLM | E2: IP Box 5% | E4: Komercyjny Omnius')

# ============================================================
# KOMENDY WHATSAPP
# ============================================================
add_heading('Komendy WhatsApp', 1)
add_table(
    ['Kategoria', 'Komenda', 'Co robi'],
    [
        ['Zadania', 'Gilbertusie task: [opis]', 'Nowe zadanie'],
        ['Zadania', 'gtd: [mysl]', 'Zapisz do przemyslenia'],
        ['Decyzje', 'decyzja: [tekst]', 'Zapisz decyzje'],
        ['Decyzje', 'outcome #ID: [wynik] rating: [1-5]', 'Ocen wynik decyzji'],
        ['Decyzje', 'skip #ID', 'Pomin ocene'],
        ['Akcje', 'tak #ID / nie #ID / edit #ID: [zmiany]', 'Zatwierdzenie akcji'],
        ['Standing orders', 'authorize: [kanal] [odbiorca] [zakres]', 'Nowy standing order'],
        ['Standing orders', 'revoke #ID | list orders | digest', 'Zarzadzanie orderami'],
        ['Delegacje', 'remind #ID / cancel #ID / extend #ID [dni]', 'Zarzadzanie delegacjami'],
        ['Autonomia', 'authority [kategoria] [0-4]', 'Zmien poziom autonomii'],
        ['Feedback', '+1 / -1', 'Ocena ostatniej odpowiedzi'],
    ]
)

# ============================================================
# WIZUALIZACJA
# ============================================================
add_heading('Droga od obserwatora do dowodcy', 1)

stages = [
    ('Sesja 23-24.03: BAZA', 'Ingestion -> Extraction -> RAG -> Brief -> WhatsApp\n"Gilbertus widzi"'),
    ('Sesja 26.03: INTELIGENCJA', '+ Evaluation + Opportunity detector + Action pipeline + Omnius\n"Gilbertus analizuje i proponuje"'),
    ('27.03 INTELLIGENCE LAYER', '+ Commitments + Meeting prep + Auto-response + Sentiment + Wellbeing\n"Gilbertus dziala proaktywnie"'),
    ('27.03 WAVE 1: NAUKA', '+ Action outcomes + Decision learning + Rule reinforcement + Authority\n"Gilbertus uczy sie z wynikow i dziala autonomicznie"'),
    ('27.03 WAVE 2: DOWODZENIE', '+ Delegation chain + Response tracking + Channel optimization\n"Gilbertus wydaje rozkazy i trackuje wykonanie"'),
    ('27.03 WAVE 3: STRATEGIA', '+ Financial framework + Calendar manager + Strategic goals + Org health\n"Gilbertus widzi finanse, zarzadza czasem, sledzi cele"'),
    ('WAVE 4: SWIADOMOSC', '+ Scenario analyzer + Market intelligence + Competitor tracking\n"Gilbertus widzi pole bitwy"'),
    ('DEPLOY (Faza A)', '+ Hetzner + Omnius REH/REF + Onboarding\n"Gilbertus dziala 24/7 na produkcji"'),
    ('VOICE (Faza C)', '+ Real-time voice + Speaker ID\n"Gilbertus rozmawia z Sebastianem"'),
    ('OMNIUS (Faza D)', '+ Kazda organizacja = Omnius + Director augmentation\n"Gilbertus dowodzi calym imperium"'),
]

add_table(['Etap', 'Opis'], stages)

# ============================================================
# COMMITY
# ============================================================
add_heading('Commity sesji 27.03', 1)
add_table(
    ['Commit', 'Opis', 'Pliki', 'Linie'],
    [
        ['842c7d1', 'Intelligence Layer: 13 modulow', '28', '+6,210'],
        ['0cd5547', 'Wave 1: feedback loops + authority', '11', '+2,687'],
        ['b485b4e', 'Wave 2: delegation chain + communication', '12', '+1,973'],
        ['2a4fa68', 'Wave 3: finanse + kalendarz + strategia', '12', '+3,270'],
        ['64a679f', 'Plaud pipeline fixes', '5', '+286'],
        ['RAZEM', '28 nowych modulow', '68', '+14,426'],
    ]
)

# Save
output = '/mnt/c/Users/jablo/Desktop/Gilbertus_Masterplan_V5.docx'
doc.save(output)
log.info(f'Saved to {output}')
