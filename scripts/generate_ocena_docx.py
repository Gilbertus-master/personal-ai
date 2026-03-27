#!/usr/bin/env python3
"""Generate Cele 2025 employee evaluation DOCX -- based on 2025 data only."""
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h


def add_section_label(doc, label):
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x2E, 0x5C, 0x8A)
    p.space_after = Pt(4)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        run.font.size = Pt(10)
    return p


def add_summary_box(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.italic = True
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)
    return p


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def score_color(score):
    if score >= 4.5:
        return "1B7A3D"
    elif score >= 3.5:
        return "4CAF50"
    elif score >= 2.5:
        return "FFC107"
    elif score >= 1.5:
        return "FF9800"
    return "F44336"


def add_score_table(doc, scores, overall, bonus_rec):
    add_section_label(doc, "OCENA SKWANTYFIKOWANA (skala 1\u20135):")
    p = doc.add_paragraph()
    run = p.add_run("1 = poni\u017cej oczekiwa\u0144  |  2 = cz\u0119\u015bciowo  |  3 = spe\u0142nia  |  4 = ponad oczekiwania  |  5 = wybitny")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    p.paragraph_format.space_after = Pt(4)

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, label in enumerate(["Wymiar", "Ocena", "Komentarz"]):
        header_cells = table.rows[0].cells
        header_cells[i].text = label
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(header_cells[i], "2E5C8A")

    for dim, score, comment in scores:
        row = table.add_row()
        row.cells[0].text = dim
        row.cells[1].text = f"{score:.1f}"
        row.cells[2].text = comment
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
        set_cell_shading(row.cells[1], score_color(score))
        for paragraph in row.cells[1].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.bold = True

    row = table.add_row()
    row.cells[0].text = "\u015aREDNIA OCENA"
    row.cells[1].text = f"{overall:.1f}"
    row.cells[2].text = f"Rekomendacja premii: {bonus_rec}"
    for cell in row.cells:
        set_cell_shading(cell, "1A1A2E")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row in table.rows:
        row.cells[0].width = Cm(5.5)
        row.cells[1].width = Cm(1.5)
        row.cells[2].width = Cm(10)
    doc.add_paragraph()


def main():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    title = doc.add_heading("CELE 2025 \u2014 Ocena Sebastian Jab\u0142o\u0144ski", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Bonus 2025 | Kalibracja | Data: 26.03.2026\nOkres oceny: stycze\u0144\u2013grudzie\u0144 2025")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_paragraph()

    # =========================================================================
    # 1. KULPA MARCIN
    # =========================================================================
    add_heading_styled(doc, "1. Kulpa Marcin \u2014 Head of OTC", level=1)

    add_section_label(doc, "WHAT (Co osi\u0105gn\u0105\u0142 w 2025):")
    add_bullet(doc, "Uruchomi\u0142 biuro OTC Trading \u2014 najwi\u0119ksza zmiana priorytetowa w 2025. Plan ekspansji przedstawiony w marcu, pierwszy deal z mar\u017c\u0105 wykonany w lipcu (20 MW DE CAL 26, EDF, mar\u017ca 0.5 EUR).")
    add_bullet(doc, "120 transakcji handlowych w 2025 \u2014 rynki PL, DE, BG, RO, SK. Produkty: baseload Y26-Y28 (431\u2013462 PLN), peak, CAL. Kontrahenci: Veolia, Axpo, EDF, E.ON, Tauron, Ecosa, Cognor, RE Alloys.")
    add_bullet(doc, "Ekspansja kontrahent\u00f3w EFET \u2014 nowe podpisane umowy ramowe, aktywacja na platformach Enmacc (6 ofert, zainteresowanie BP i Engie GM), EEX.")
    add_bullet(doc, "Nowi klienci: Beskidzka Energia (8k PLN/mc + 50 gr/MWh, zamkni\u0119ty listopad 2025), Ecosa (referral gaz/energy swap), Tauron (\u201epozytywnie zaskoczony szybko\u015bci\u0105\u201d).")
    add_bullet(doc, "Podwoi\u0142 bud\u017cet tradingowy z 40M do 80M PLN (lipiec 2025).")
    add_bullet(doc, "Zaprojektowa\u0142 system bonusowy trader\u00f3w OTC: skala 3\u201322% PnL, max 25% po kosztach (3x roczna pensja + koszty transakcyjne). Zatwierdzony przez Sebastiana.")
    add_bullet(doc, "Dostarczy\u0142 PowerBI PnL report (kwiecie\u0144 2025).")
    add_bullet(doc, "Rekrutacja nowej traderki (listopad 2025, start kwiecie\u0144 2026, 20k PLN) \u2014 focus na relacjach z elektrowniami i konsumentami.")
    add_bullet(doc, "Wsparcie Dispatch Center \u2014 pokrycie godzin biznesowych + cz\u0119\u015b\u0107 drugiej zmiany.")
    add_bullet(doc, "Samoocena KPI: OTC 15% wagi / 90% realizacji. Sebastian zatwierdzi\u0142.")

    add_section_label(doc, "HOW (Jak pracuje):")
    add_bullet(doc, "Samodzielny i inicjatywny \u2014 sam zaprojektowa\u0142 bonus scheme, plan ekspansji, buduje relacje z kontrahentami bez nadzoru.")
    add_bullet(doc, "Aktywny na wielu rynkach jednocze\u015bnie \u2014 PL, DE, BG, RO, SK, cross-border JAO aukcje.")
    add_bullet(doc, "Szybki w egzekucji \u2014 Tauron \u201epozytywnie zaskoczony szybko\u015bci\u0105\u201d, Enmacc 6 ofert po jednym dniu.")
    add_bullet(doc, "Komunikatywny \u2014 szeroka sie\u0107 wsp\u00f3\u0142pracy: Kalinowska (PM mandaty), Kuzminski (tech), Juchacz (REF/gaz), Skotnicka (risk), Baranowski (strategia).")

    add_section_label(doc, "S\u0141ABE STRONY / PORA\u017bKI W 2025:")
    add_bullet(doc, "RE OBR\u00d3T \u2014 CEL NIEZREALIZOWANY (waga 25%): Sp\u00f3\u0142ka za\u0142o\u017cona i dodana do cashpool, ale wniosek o koncesj\u0119 niezłożony. Pow\u00f3d: URE mog\u0142oby \u017c\u0105da\u0107 zabezpieczenia ~70M PLN. Sam obni\u017cy\u0142 wag\u0119 z 25% do 10% w samoocenie.")
    add_bullet(doc, "B\u0141\u0104D LUCY \u2014 STRATA 150K EUR (maj 2025): Import danych przez system Lucy spowodowa\u0142 strat\u0119 150 000 EUR. Kulpa zidentyfikowa\u0142 przyczyn\u0119 i doda\u0142 mechanizm por\u00f3wnywania stan\u00f3w, ale strata ju\u017c wyst\u0105pi\u0142a.")
    add_bullet(doc, "OKTE S\u0142owacja \u2014 OBNI\u017bONY RATING (lipiec 2025): Sp\u00f3\u017aniona p\u0142atno\u015b\u0107 200 EUR spowodowa\u0142a obni\u017cenie ratingu tradingowego. OKTE odm\u00f3wi\u0142o wyj\u0105tku.")
    add_bullet(doc, "AZOTY \u2014 ODRZUCENIE OFERTY (grudzie\u0144 2025): 10 MW @ 465 PLN odrzucone \u2014 \u201erozbili\u015bmy si\u0119 o zabezpieczenie\u201d i warunki p\u0142atno\u015bci (klient chcia\u0142 60 dni, RE oferowa\u0142o 20).")
    add_bullet(doc, "PROBLEM DKR+3 (sierpie\u0144 2025): Brak mo\u017cliwo\u015bci osi\u0105gni\u0119cia ceny DKR+3 na bliskich kontraktach \u2014 rynek oferowa\u0142 ta\u0144sze alternatywy.")
    add_bullet(doc, "WIRE PROJECT \u2014 LUKA KOMUNIKACYJNA (maj 2025): Sebastian zapyta\u0142 kto z IT prowadzi projekt WIRE \u2014 Kulpa nie wiedzia\u0142.")
    add_bullet(doc, "PROBLEMY OSOBISTE: wypadek motocyklowy (lipiec), kryzys zdrowotny \u017cony (wrzesie\u0144), operacja (grudzie\u0144), sprawy rozwodowe \u2014 wp\u0142yw na dyspozycyjno\u015b\u0107.")

    add_section_label(doc, "PODSUMOWANIE:")
    add_summary_box(doc,
        "Rok 2025 by\u0142 rokiem budowy OTC \u2014 Marcin dostarczy\u0142 120 transakcji, nowych kontrahent\u00f3w (Tauron, "
        "Beskidzka Energia), podwojony bud\u017cet i system bonusowy. Kluczowa pora\u017cka to niezrealizowany cel "
        "RE Obr\u00f3t (25% wagi) oraz strata 150K EUR z b\u0142\u0119du Lucy. Rating OKTE obni\u017cony przez sp\u00f3\u017anion\u0105 "
        "p\u0142atno\u015b\u0107 200 EUR \u2014 drobny b\u0142\u0105d z powa\u017cnymi konsekwencjami. Og\u00f3lnie silny rok operacyjny, "
        "ale governance i dyscyplina procesowa wymagaj\u0105 poprawy."
    )

    add_score_table(doc, [
        ("Realizacja cel\u00f3w biznesowych", 4.0, "OTC uruchomiony, 120 transakcji, nowi klienci. Ale RE Obr\u00f3t (25% wagi) niezrealizowany."),
        ("Inicjatywa i przedsi\u0119biorczo\u015b\u0107", 4.5, "Samodzielnie zaprojektowa\u0142 ekspansj\u0119, bonus scheme, rekrutacj\u0119, Enmacc."),
        ("Jako\u015b\u0107 i dok\u0142adno\u015b\u0107 pracy", 3.0, "Strata 150K EUR (Lucy), rating OKTE za 200 EUR, luka WIRE. Problemy procesowe."),
        ("Wsp\u00f3\u0142praca i komunikacja", 4.0, "Szeroka sie\u0107, szybki response. Luka: nie wiedzia\u0142 kto prowadzi WIRE."),
        ("Zarz\u0105dzanie zespo\u0142em", 4.0, "Rekrutacja traderki, bonus model, wsparcie Dispatch Center."),
        ("Zarz\u0105dzanie ryzykiem", 3.0, "150K EUR strata, OKTE rating, zabezpieczenia Azoty. Poprawki reaktywne, nie prewencyjne."),
    ], overall=3.8, bonus_rec="70\u201385% premii docelowej")

    # =========================================================================
    # 2. KALINOWSKA MAJA
    # =========================================================================
    add_heading_styled(doc, "2. Kalinowska Maja \u2014 Head of Portfolio Management", level=1)

    add_section_label(doc, "WHAT (Co osi\u0105gn\u0119\u0142a w 2025):")
    add_bullet(doc, "GoldenPeaks Capital \u2014 wsparcie analityczne i PM przy negocjacjach: kalkulacje zysk\u00f3w, wyceny portfela 690 MWp, propozycja warunk\u00f3w ryzyka (23 PLN/MWh, podzia\u0142 10/90). Potencjalny upside 20+ mln PLN. UWAGA: funkcja wsparciowo-analityczna, nie prowadzi\u0142a negocjacji samodzielnie.")
    add_bullet(doc, "NOFAR \u2014 przegl\u0105d warunk\u00f3w kontraktowych, koordynacja odpowiedzi z Mart\u0105 Batory. Klient zagrozi\u0142 post\u0119powaniem s\u0105dowym z powodu op\u00f3\u017anie\u0144.")
    add_bullet(doc, "MIROVA \u2014 pricing wy\u0142\u0105cze\u0144 mocy (SPOT -14/-16.5 PLN), klauzula 6h cen ujemnych. Rola analityczna.")
    add_bullet(doc, "Bu\u0142garia wiatr \u2014 wyceny WINDKRAFT, BALCHIK (4 EUR mar\u017ca, 3.14 EUR bilansowanie), oferty 2026-2028.")
    add_bullet(doc, "Rumunia \u2014 analiza komponent\u00f3w cenowych, analiza terminacji 14 propozycji.")
    add_bullet(doc, "Zasady renegocjacji \u2014 tylko >900 PLN, max do 2027, max 4 lata. Cap na mar\u017c\u0119 merchant. Realizacja zlecenia Sebastiana.")
    add_bullet(doc, "CR portfela wzr\u00f3s\u0142 z ~80% (2024) do ~91% (2025) po wy\u0142\u0105czeniu redysponowania.")
    add_bullet(doc, "Pozytywna opinia PSE (luty 2025): \u201enajlepiej zbilansowany portfel na rynku i najlepsze predykcje\u201d.")
    add_bullet(doc, "Zarz\u0105dzanie zespo\u0142em: podwy\u017cki 4000 PLN/mc (Rupniewska +2500, Batory +500, Bartosik +1000). Utworzenie osobnego zespo\u0142u PM do zarz\u0105dzania pozycjami.")
    add_bullet(doc, "Sprzeda\u017c ~50 MW na TGE Q4 2025, 40 MW CAL 2028 @ 443.5 PLN, hedging sub-sezonalny.")

    add_section_label(doc, "HOW (Jak pracuje):")
    add_bullet(doc, "Analityczna \u2014 precyzyjna w kalkulacjach, kwestionuje warto\u015bci (Fixing I, klauzula BRP), identyfikuje zawy\u017cone pozycje biura handlu.")
    add_bullet(doc, "REAKTYWNA \u2014 realizuje zlecone zadania rzetelnie, ale nie wykazuje inicjatywy w proponowaniu usprawnie\u0144 proces\u00f3w czy narz\u0119dzi. Nie generuje w\u0142asnych pomys\u0142\u00f3w na optymalizacj\u0119.")
    add_bullet(doc, "FUNKCJA WSPARCIOWA \u2014 nie prowadzi\u0142a samodzielnie \u017cadnych negocjacji z kontrahentami. Zawsze pe\u0142ni\u0142a rol\u0119 analityczn\u0105/PM, nie biznesow\u0105.")
    add_bullet(doc, "Pokrywa 15+ rynk\u00f3w europejskich (Nordic, DE, ES, IT, HU, RO, BG, Ba\u0142kany).")
    add_bullet(doc, "Comiesi\u0119czne spotkania z Sebastianem (od maja) + cotygodniowe ze sprzeda\u017c\u0105 i wytw\u00f3rcami.")

    add_section_label(doc, "S\u0141ABE STRONY / PORA\u017bKI W 2025:")
    add_bullet(doc, "B\u0141\u0118DNIE DOMYKANA POZYCJA \u2014 STRATA 239 627 PLN (stycze\u0144\u2013listopad 2025): Biuro tradingu sk\u0142ada\u0142o zawy\u017cone pozycje na RDN, \u0142\u0105cz\u0105c pozycje tradingowe i portfelowe wbrew procedurze z 23.12.2024. Maja nie mia\u0142a pe\u0142nego dost\u0119pu do danych tradingu. Sebastian: \u201eto nie s\u0105 du\u017ce kwoty\u201d, ale \u201eto nie t\u0142umaczy rozbie\u017cno\u015bci kilku milion\u00f3w miesi\u0119cznie\u201d.")
    add_bullet(doc, "B\u0141\u0104D RAPORTOWY \u2014 400 GWh R\u00d3\u017bNICY (pa\u017adziernik 2025): Odfiltrowany SB przy ostatnim wgraniu, r\u00f3\u017cnica 400 GWh w raporcie. Szybko poprawione, ale decyzje mog\u0142y by\u0107 podejmowane na b\u0142\u0119dnych danych.")
    add_bullet(doc, "OP\u00d3\u0179NIENIE IMPLEMENTACJI SUMY (luty 2025): Sebastian prosi\u0142 od 28 stycznia, suma nadal nie wdro\u017cona. Maja przeprosi\u0142a za op\u00f3\u017anienie.")
    add_bullet(doc, "BRAK MIGRACJI WOLUMEN\u00d3W SB\u2192EBIKOM (marzec 2025): Wolumeny za przesz\u0142e miesi\u0105ce niezmigrowane. Umowy wisia\u0142y na statusach, cz\u0119\u015b\u0107 na li\u015bcie b\u0142\u0119d\u00f3w.")
    add_bullet(doc, "LUCY \u2014 B\u0141\u0118DNE POZYCJE SHORT PM_BG (wrzesie\u0144 2025): Wolumen kupowany od wytw\u00f3rc\u00f3w nie wprowadzony do Lucy. B\u0142\u0105d systemowy (kontrakty automatycznie przesuwaj\u0105 dat\u0119 ko\u0144ca o 1 dzie\u0144).")
    add_bullet(doc, "NOFAR \u2014 ESKALACJA KLIENTA: Klient zagrozi\u0142 post\u0119powaniem s\u0105dowym z powodu op\u00f3\u017anie\u0144 w odpowiedziach.")
    add_bullet(doc, "~7 ABSENCJI L4: g\u0142\u00f3wnie choroby dziecka (wrzesie\u0144, listopad, grudzie\u0144). Brak widocznego planu zast\u0119pstw \u2014 single point of failure w PM.")
    add_bullet(doc, "BRAK PROAKTYWNO\u015aCI: Nie proponuje usprawnie\u0144 proces\u00f3w ani narz\u0119dzi z w\u0142asnej inicjatywy. JIRA/ticketing w 2024 by\u0142 reakcj\u0105 na kryzys (Intereuropol), nie proaktywn\u0105 inicjatyw\u0105. Czeka na zlecenia zamiast identyfikowa\u0107 problemy i proponowa\u0107 rozwi\u0105zania.")
    add_bullet(doc, "BRAK SAMODZIELNO\u015aCI NEGOCJACYJNEJ: Mimo tytu\u0142u Head of PM, nie prowadzi\u0142a w 2025 \u017cadnych negocjacji z kontrahentami samodzielnie. Zawsze funkcja wsparciowo-analityczna. Na poziomie Head oczekuje si\u0119 wi\u0119kszej autonomii biznesowej.")

    add_section_label(doc, "PODSUMOWANIE:")
    add_summary_box(doc,
        "Maja dostarczy\u0142a solidne wsparcie analityczne przy kluczowych dealach (GoldenPeaks, MIROVA, Bu\u0142garia) "
        "i podnios\u0142a CR portfela z 80% do 91%. PSE potwierdzi\u0142o najlepszy portfel na rynku. Jednocze\u015bnie "
        "wyst\u0105pi\u0142y powa\u017cne problemy: strata 239K PLN, 400 GWh b\u0142\u0105d raportowy, eskalacja NOFAR. "
        "Kluczowy deficyt: reaktywno\u015b\u0107 zamiast proaktywno\u015bci \u2014 realizuje zlecenia rzetelnie, ale nie "
        "generuje inicjatyw usprawniaj\u0105cych. Nie prowadzi negocjacji samodzielnie mimo tytu\u0142u Head. "
        "Cz\u0119ste L4 bez planu zast\u0119pstw to dodatkowe ryzyko."
    )

    add_score_table(doc, [
        ("Realizacja cel\u00f3w biznesowych", 4.0, "CR 91%, PSE \u201enajlepszy portfel\u201d, wsparcie GoldenPeaks/MIROVA/BG. Ale rola wsparciowa, nie wiod\u0105ca."),
        ("Inicjatywa i proaktywno\u015b\u0107", 2.5, "Reaktywna \u2014 realizuje zlecenia, nie generuje usprawnie\u0144. JIRA by\u0142o reakcj\u0105 na kryzys."),
        ("Jako\u015b\u0107 i dok\u0142adno\u015b\u0107 pracy", 3.0, "Strata 239K PLN, 400 GWh b\u0142\u0105d, op\u00f3\u017anienia wdro\u017ce\u0144, b\u0142\u0119dy Lucy PM_BG."),
        ("Wsp\u00f3\u0142praca i komunikacja", 3.5, "Eskalacja NOFAR, ale dobra koordynacja wewn\u0119trzna i comiesi\u0119czne synce."),
        ("Zarz\u0105dzanie zespo\u0142em", 3.0, "Podwy\u017cki OK, ale 7x L4 bez planu zast\u0119pstw. Zesp\u00f3\u0142 nie realizuje priorytet\u00f3w bez nadzoru."),
        ("Samodzielno\u015b\u0107 i leadership", 2.5, "Nie prowadzi negocjacji samodzielnie. Funkcja wsparciowa mimo tytu\u0142u Head of PM."),
    ], overall=3.1, bonus_rec="55\u201370% premii docelowej")

    # =========================================================================
    # 3. MRUK MARIA
    # =========================================================================
    add_heading_styled(doc, "3. Mruk Maria \u2014 Head of Environmental Commodities / Portfolio Management", level=1)

    add_section_label(doc, "WHAT (Co osi\u0105gn\u0119\u0142a w 2025):")
    add_bullet(doc, "UTWORZENIE BIURA EC PM (pa\u017adziernik 2025): Przygotowa\u0142a 6-stronicowy dokument \u201ePropozycja Powo\u0142ania Biura Strategii Portfela\u201d + prezentacj\u0119. Zarz\u0105d podj\u0105\u0142 uchwa\u0142\u0119, Maria mianowana Head of EC PM.")
    add_bullet(doc, "IDENTYFIKACJA RYZYKA IREC \u2014 EUR 71K: Kary za certyfikaty >21 miesi\u0119cy: 43 863 EUR natychmiast + 27 141 EUR do ko\u0144ca pa\u017adziernika. Podnios\u0142a temat zanim sta\u0142 si\u0119 problemem.")
    add_bullet(doc, "METODOLOGIA CENNIKA GP: Opracowa\u0142a pe\u0142n\u0105 metodologi\u0119 \u2014 podzia\u0142 Y/Y-1, amortyzacja named/unnamed, technologie, fixed/full offtake. Zatwierdzona przez Sebastiana i Rocha.")
    add_bullet(doc, "MANDAT ZAKUPU GP: 1 TWh GP z cen\u0105 max 40 gr/MWh (wrzesie\u0144), potem podwy\u017cszona do 50 gr dla 958 485 MWh.")
    add_bullet(doc, "PROCEDURY OFERT WI\u0104\u017b\u0104CYCH (lipiec 2025): BG/RO deadline 12:00, inne rynki 14:00, aktualizacja SB w 30 min od podpisu.")
    add_bullet(doc, "WYCENY MI\u0118DZYNARODOWE: Bu\u0142garia wiatr 10-letni kontrakt (zatwierdzony maj), Rumunia PPA 114 MW (oferta wi\u0105\u017c\u0105ca sierpie\u0144), Chorwacja wycena pilna (czerwiec), Macedonia analiza 7-10 lat.")
    add_bullet(doc, "CFP \u2014 hedging na EEX: Wycena i oferta Park I (67 EUR) i Park II (69 EUR) na Q3-Q4 2025 i 2026. Zatwierdzona i wys\u0142ana.")
    add_bullet(doc, "SKYBORN \u2014 gotowosc do deadline\u2019u z ofert\u0105 wi\u0105\u017c\u0105c\u0105.")
    add_bullet(doc, "POLITYKA ODMOWY: Odrzuci\u0142a PPA Czarnog\u00f3ra (350 GWh, 10 lat) zgodnie z polityk\u0105 max 3 lata dla kraj\u00f3w spoza PL/DE/RO. Sebastian potwierdzi\u0142.")
    add_bullet(doc, "REKRUTACJA: 2 FTE zatwierdzone przez Zarz\u0105d (Senior Pricing Specialist 14-22K + Analyst 11-14K PLN).")
    add_bullet(doc, "PODWY\u017bKA 4000 PLN/mc od listopada \u2014 powi\u0105zana ze zmian\u0105 stanowiska.")
    add_bullet(doc, "Codzienne cenniki PL (Cennik PL) dla Origination, EC-Traders, PM. Cenniki GP z pe\u0142n\u0105 metodologi\u0105.")

    add_section_label(doc, "HOW (Jak pracuje):")
    add_bullet(doc, "Systematyczna \u2014 codzienne cenniki, procedury z precyzyjnymi deadlinami, formalne wnioski (g\u0142osowania obiegowe).")
    add_bullet(doc, "Proaktywna w identyfikacji ryzyk \u2014 IREC EUR 71K podniesione zanim sta\u0142o si\u0119 kryzysem.")
    add_bullet(doc, "Multi-market \u2014 PL, DE, BG, RO, Chorwacja, Macedonia, Czarnog\u00f3ra. R\u00f3\u017cne metodologie per region.")
    add_bullet(doc, "Wsp\u00f3\u0142praca z Kalinowsk\u0105 (PM), Kulp\u0105 (OTC mandaty), Jankowskim (PAN-EU), Makarukiem (BD).")

    add_section_label(doc, "S\u0141ABE STRONY / PORA\u017bKI W 2025:")
    add_bullet(doc, "OP\u00d3\u0179NIENIE RAPORTU GP (koniec sierpnia 2025): Braki w danych sprzed 2024, dane AIB niekompletne. Raport prze\u0142o\u017cony na poniedzia\u0142ek/wtorek.")
    add_bullet(doc, "KOREKTA CENNIKOWA (lipiec 2025): Sebastian zarz\u0105dzi\u0142 piln\u0105 korekt\u0119 Buy SPOT PL indicative. Na poziomie Head cenniki powinny by\u0107 poprawne bez interwencji.")
    add_bullet(doc, "WAHANIE PRZY MACEDONII (maj 2025): Pyta\u0142a o pozwolenie na z\u0142o\u017cenie oferty wi\u0105\u017c\u0105cej 7-10 lat. Obawy s\u0142uszne, ale na poziomie Head oczekuje si\u0119 rekomendacji z uzasadnieniem, nie pytania o zgod\u0119.")
    add_bullet(doc, "UMORZENIA ZAGRANICZNYCH GP \u2014 RYZYKO COMPLIANCE (sierpie\u0144 2025): 800+ GWh w 2024, ~150 GWh w 2025. Brak odpowiedzi prawnika. Temat podniesiony, ale pozostawa\u0142 nierozwi\u0105zany.")
    add_bullet(doc, "ZESP\u00d3\u0141 W BUDOWIE: 2 FTE zatwierdzone dopiero w grudniu \u2014 przez ca\u0142y rok operowa\u0142a de facto jednoosobowo w nowej roli.")

    add_section_label(doc, "PODSUMOWANIE:")
    add_summary_box(doc,
        "Maria przesz\u0142a w 2025 z Portfolio Managera na Head of EC PM \u2014 awans potwierdzony uchwa\u0142\u0105 Zarz\u0105du "
        "i podwy\u017ck\u0105 4K PLN. Kluczowe osi\u0105gni\u0119cia: identyfikacja ryzyka IREC (71K EUR), pe\u0142na metodologia GP, "
        "procedury ofert wi\u0105\u017c\u0105cych, wyceny na 6 rynkach, polityka odmowy (Czarnog\u00f3ra). S\u0142abo\u015bci: op\u00f3\u017anienia "
        "raportowe, korekty cennikowe wymagaj\u0105ce interwencji, wahanie decyzyjne (Macedonia). Solidny rok transformacji, "
        "ale pe\u0142na ocena b\u0119dzie mo\u017cliwa po ramp-upie zespo\u0142u w 2026."
    )

    add_score_table(doc, [
        ("Realizacja cel\u00f3w biznesowych", 4.5, "Biuro EC PM utworzone, IREC 71K, GP metodologia, wyceny 6 rynk\u00f3w, CFP/SKYBORN."),
        ("Inicjatywa i przedsi\u0119biorczo\u015b\u0107", 4.5, "IREC risk, propozycja biura, polityka odmowy, procedury ofertowe."),
        ("Jako\u015b\u0107 i dok\u0142adno\u015b\u0107 pracy", 3.5, "Codzienne cenniki OK, ale korekty cennikowe i op\u00f3\u017anienia raport\u00f3w."),
        ("Wsp\u00f3\u0142praca i komunikacja", 4.0, "Dobra koordynacja z PM, OTC, PAN-EU. Formalne wnioski do Zarz\u0105du."),
        ("Zarz\u0105dzanie zespo\u0142em", 3.0, "2 FTE dopiero w grudniu. Ca\u0142y rok solo \u2014 ryzyko przeci\u0105\u017cenia."),
        ("Zarz\u0105dzanie ryzykiem", 4.5, "IREC wczesna identyfikacja, compliance GP, r\u00f3\u017cne metodologie per rynek."),
    ], overall=4.0, bonus_rec="80\u2013100% premii docelowej")

    # =========================================================================
    # 4. AGATA MORSKA
    # =========================================================================
    add_heading_styled(doc, "4. Agata Morska \u2014 Executive Assistant", level=1)

    add_section_label(doc, "WHAT (Co osi\u0105gn\u0119\u0142a w 2025):")
    add_bullet(doc, "234 zarejestrowane zdarzenia w 2025 \u2014 \u015brednio ~1 zdarzenie na dzie\u0144 roboczy.")
    add_bullet(doc, "SPRZEDA\u017b HONDAJET HA 420: Samodzielna korespondencja z AMS Aircraft (UK) w j. angielskim. Negocjacja prowizji z 2.25% do 2%. Walka o non-exclusive (przegrana \u2014 AMS wymusi\u0142 exclusive). Koordynacja z Jet Story.")
    add_bullet(doc, "PILATUS PC-24 (nowy samolot): Dokumenty mLeasing, projekt umowy, koordynacja szkole\u0144 pilot\u00f3w, protok\u00f3\u0142 odbioru (wrzesie\u0144), subskrypcje CPDLC/datalink ($6995/rok + $1495/rok).")
    add_bullet(doc, "REMONT D\u0104BROWSKIEGO 45 (ca\u0142y rok): Recyrkulacja/klimatyzacja (maj), okna (czerwiec, mimo braku PnB), TV/audio (lipiec), oran\u017ceria Lassart (wrzesie\u0144), ceramika (pa\u017adziernik), TOM-BUD prace wyko\u0144czeniowe + renowacja piaskowca (grudzie\u0144). Wynegocjowa\u0142a obni\u017ck\u0119 15 000 PLN (z 94 200 PLN).")
    add_bullet(doc, "FAKTURY: Miliony PLN przetworzonych \u2014 od 40 PLN (Orange) po 798 615 PLN (pa\u017adziernik SJ Sp.k.). Jet Story, Honda Aircraft (USD), mLeasing, KADOM, Ground Frost, REF, RESA, BYS, Solid Security, UDT.")
    add_bullet(doc, "KOLEKCJA SZTUKI: Instalacja rze\u017aby Mitoraja \u201eCenturion I\u201d (~400kg) w North Gate. Faktura za dzie\u0142o Ewy Juszkiewicz. Wycena DESA (100 PLN/obiekt vs 500 PLN PolswissArt).")
    add_bullet(doc, "DOKUMENTACJA KORPORACYJNA: SF 2023 (zaleg\u0142y, pilnie z\u0142o\u017cony), pe\u0142nomocnictwa e-Dor\u0119czenia, KSeF, weksle mLeasing, Fundacja Rodzinna SJ rejestracja, bank Santander.")
    add_bullet(doc, "POSIEDZENIA ZARZ\u0104DU: Organizacja i dystrybucja protoko\u0142\u00f3w (maj, czerwiec).")
    add_bullet(doc, "UBEZPIECZENIA: Ferrari + Bentley (OC 800 PLN/szt, AC Ferrari 26 685, Bentley 20 390 PLN). Dom Jaros\u0142awiec 900K PLN. Polisa Bitwa pod Rokitn\u0105.")
    add_bullet(doc, "PODRÓŻE I LOGISTYKA: Helikopter Sky Poland, Columbia Beach Resort Cypr (3 Junior Suites 487 EUR/noc), dokumenty wizowe Indie.")
    add_bullet(doc, "SPRAWY OSOBISTE SJ: Thames British School (Winter Camp, posi\u0142ki), kwiaty i czekoladki na zako\u0144czenie roku, ortodonta Wojtka, paszport znaleziony.")

    add_section_label(doc, "HOW (Jak pracuje):")
    add_bullet(doc, "Wyj\u0105tkowo wszechstronna \u2014 lotnictwo, budownictwo, prawo korporacyjne, finanse, sztuka, logistyka, sprawy osobiste.")
    add_bullet(doc, "Zero powa\u017cnych b\u0142\u0119d\u00f3w operacyjnych w 234 zdarzeniach.")
    add_bullet(doc, "Skuteczna negocjatorka \u2014 15K TOM-BUD, prowizja AMS 2.25%\u21922%, DESA 100 vs 500 PLN/obiekt.")
    add_bullet(doc, "Proaktywna \u2014 flaguje koszty (TV/audio), reklamuje brakuj\u0105ce faktury (KADOM), sygnalizuje ryzyka terminowe (okna bez PnB).")
    add_bullet(doc, "Stabilna \u2014 tylko 2 tygodnie urlopu w ca\u0142ym roku (28.07\u201308.08).")

    add_section_label(doc, "S\u0141ABE STRONY / PORA\u017bKI W 2025:")
    add_bullet(doc, "AMS AIRCRAFT \u2014 PRZEGRANA NEGOCJACJA EXCLUSIVE: Nie uda\u0142o si\u0119 uzyska\u0107 non-exclusive na remarketing HondaJet. AMS wymusi\u0142 exclusive mandate. Kompromis (exclusive z wyj\u0105tkiem Jet Story), ale nie pierwotny cel.")
    add_bullet(doc, "SINGLE POINT OF FAILURE: Jedyna osoba obs\u0142uguj\u0105ca lotnictwo, remont, faktury SJ Sp.k., dokumentacj\u0119. Backup (Piotr D\u0105browski) tylko na odbieranie dzieci. Jej 2-tygodniowa absencja = parali\u017c operacji.")
    add_bullet(doc, "WYPADEK SAMOCHODOWY (sierpie\u0144 2025): Kolizja na ul. Stawki \u2014 tylko samoch\u00f3d uszkodzony, bez obra\u017ce\u0144. Ale ryzyko zwi\u0105zane z SPOF.")
    add_bullet(doc, "MIESZANIE R\u00d3L FIRMOWYCH I PRYWATNYCH: Thames, ortodonta, hotele wakacyjne, kwiaty na zako\u0144czenie roku szkolnego \u2014 obok KSeF, protoko\u0142\u00f3w zarz\u0105du i mLeasing. Brak jasnej granicy.")
    add_bullet(doc, "BRAK DEVELOPMENT PLANU: Rola rozros\u0142a si\u0119 organicznie, ale bez formalnego planu rozwoju kompetencji. Szeroko\u015b\u0107 zada\u0144 mo\u017ce hamowa\u0107 g\u0142\u0119boko\u015b\u0107 ekspertyzy.")

    add_section_label(doc, "PODSUMOWANIE:")
    add_summary_box(doc,
        "Agata to operacyjny kr\u0119gos\u0142up \u2014 234 zdarzenia, zero powa\u017cnych b\u0142\u0119d\u00f3w. W 2025 obs\u0142u\u017cy\u0142a "
        "sprzeda\u017c HondaJet, odbi\u00f3r Pilatusa, pe\u0142ny remont D\u0105browskiego 45, miliony PLN w fakturach, "
        "kolekcj\u0119 sztuki i dokumentacj\u0119 korporacyjn\u0105 4+ podmiot\u00f3w. Jedyna powa\u017cna s\u0142abo\u015b\u0107 to ryzyko SPOF \u2014 "
        "brak zast\u0119pstw i plan\u00f3w sukcesji. Niezawodna, wszechstronna, niezast\u0105piona."
    )

    add_score_table(doc, [
        ("Realizacja cel\u00f3w operacyjnych", 5.0, "234 zdarzenia, zero b\u0142\u0119d\u00f3w. Lotnictwo, remont, faktury, dokumentacja."),
        ("Inicjatywa i proaktywno\u015b\u0107", 4.5, "Flaguje koszty, reklamuje braki, negocjuje bez pytania."),
        ("Jako\u015b\u0107 i dok\u0142adno\u015b\u0107 pracy", 5.0, "Miliony PLN miesi\u0119cznie bez b\u0142\u0119d\u00f3w. Pe\u0142na dokumentacja."),
        ("Wsp\u00f3\u0142praca i komunikacja", 4.5, "Sprawna koordynacja z dostawcami, kancelariami, HR, ksi\u0119gowo\u015bci\u0105."),
        ("Wszechstronno\u015b\u0107 i adaptacja", 5.0, "Lotnictwo + remont + sztuka + dokumentacja \u2014 bez straty jako\u015bci."),
        ("Planowanie i ci\u0105g\u0142o\u015b\u0107", 3.0, "SPOF \u2014 brak backup\u00f3w, brak planu sukcesji, brak development planu."),
    ], overall=4.5, bonus_rec="100\u2013120% premii docelowej")

    # =========================================================================
    # RANKING
    # =========================================================================
    add_heading_styled(doc, "Podsumowanie \u2014 Ranking ko\u0144cowy", level=1)

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, label in enumerate(["Pracownik", "\u015arednia", "Rekomendacja premii", "Priorytet retencji"]):
        header_cells = table.rows[0].cells
        header_cells[i].text = label
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(header_cells[i], "1A1A2E")

    ranking = [
        ("Agata Morska", "4.5", "100\u2013120%", "KRYTYCZNY"),
        ("Mruk Maria", "4.0", "80\u2013100%", "WYSOKI"),
        ("Kulpa Marcin", "3.8", "70\u201385%", "WYSOKI"),
        ("Kalinowska Maja", "3.1", "55\u201370%", "\u015aREDNI"),
    ]
    colors = ["1B7A3D", "4CAF50", "4CAF50", "FFC107"]

    for (name, score, bonus, retention), color in zip(ranking, colors):
        row = table.add_row()
        row.cells[0].text = name
        row.cells[1].text = score
        row.cells[2].text = bonus
        row.cells[3].text = retention
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
        row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(row.cells[1], color)
        for paragraph in row.cells[1].paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.bold = True

    doc.add_paragraph()

    # Disclaimer
    add_section_label(doc, "Uwagi metodologiczne:")
    p2 = doc.add_paragraph()
    run2 = p2.add_run(
        "Poziom ufno\u015bci ocen: OBNI\u017bONY. Dane \u017ar\u00f3d\u0142owe obejmuj\u0105 wy\u0142\u0105cznie korespondencj\u0119 "
        "z perspektywy Sebastiana Jab\u0142o\u0144skiego (Teams, email, WhatsApp). Brak dost\u0119pu do danych "
        "z perspektywy ocenianych pracownik\u00f3w. Okres oceny: stycze\u0144\u2013grudzie\u0144 2025."
    )
    run2.font.size = Pt(9)
    run2.italic = True
    run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    output_path = sys.argv[1] if len(sys.argv) > 1 else "Cele_2025_Ocena_SJ.docx"
    doc.save(output_path)
    print(f"DOCX saved: {output_path}")


if __name__ == "__main__":
    main()
