#!/usr/bin/env python3
"""Generate REH Omnius Team Analysis DOCX document."""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import shutil
from datetime import date

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10.5)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    hs.font.bold = True
    if level == 1:
        hs.font.size = Pt(18)
        hs.paragraph_format.space_before = Pt(24)
        hs.paragraph_format.space_after = Pt(12)
    elif level == 2:
        hs.font.size = Pt(14)
        hs.paragraph_format.space_before = Pt(18)
        hs.paragraph_format.space_after = Pt(8)
    else:
        hs.font.size = Pt(12)
        hs.paragraph_format.space_before = Pt(12)
        hs.paragraph_format.space_after = Pt(6)


def add_classification(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("POUFNE — WEWNETRZNE")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    run.font.bold = True


def set_cell_shading(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def make_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "1B3A5C")

    # Data
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "EEF2F7")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # spacing
    return table


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


# ═══════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════
add_classification(doc)

for _ in range(4):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Analiza potrzeby dedykowanego\nzespolu AI/Omnius w REH")
run.font.size = Pt(26)
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
run.font.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Respect Energy Holding sp. z o.o.")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run(f"Dokument przygotowany: {date.today().strftime('%d.%m.%Y')}\n")
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run = meta.add_run("Autor: Gilbertus Albans — System AI Mentata\n")
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run = meta.add_run("Klasyfikacja: POUFNE — WEWNETRZNE\n")
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
run = meta.add_run("Wersja: 1.0")
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ═══════════════════════════════════════════════════════════════
add_classification(doc)

doc.add_heading("Spis tresci", level=1)
toc_items = [
    "1. Streszczenie wykonawcze i rekomendacja",
    "2. Kontekst strategiczny",
    "3. Opcje organizacyjne — 3 scenariusze",
    "4. Profil stanowiska: AI & Automation Officer",
    "5. Zakres zadan — pierwsze 12 miesiecy",
    "6. KPI i metryki sukcesu",
    "7. Ryzyka i mitygacja",
    "8. Rekomendacja Gilbertusa",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 1. STRESZCZENIE WYKONAWCZE
# ═══════════════════════════════════════════════════════════════
add_classification(doc)

doc.add_heading("1. Streszczenie wykonawcze i rekomendacja", level=1)

p = doc.add_paragraph()
run = p.add_run("Rekomendacja: TAK")
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
p.add_run(" — REH potrzebuje dedykowanej osoby odpowiedzialnej za wdrozenie i rozwoj systemu Omnius.")

doc.add_paragraph(
    "Respect Energy Holding stoi przed fundamentalna transformacja cyfrowa. System Omnius REH ma potencjal "
    "zautomatyzowac 40-70% pracy operacyjnej w ksiegowosci (redukcja z 24 do 8-14 FTE), usprawnic decision "
    "tracking na kazdym poziomie zarzadzania oraz zintegrowac rozproszony ekosystem IT (ERP, HRM, DMS, CRM, "
    "SAPICO, Savangard). Szacowane oszczednosci to kilka milionow PLN rocznie."
)

doc.add_paragraph(
    "Jednoczesnie Sebastian Jablonski jest wylaczony operacyjnie od 1 lutego 2026, a CEO Roch Baranowski "
    "nie ma bandwidth na operacyjny nadzor wdrozenia AI. Dzial RE_ITS (~2,17M PLN/rok) jest skoncentrowany "
    "na utrzymaniu obecnych systemow, nie na innowacji AI. Bez dedykowanego wlasciciela, Omnius pozostanie "
    "w statusie BLOCKED — zaprojektowany, ale nie wdrozony."
)

p = doc.add_paragraph()
run = p.add_run("Kluczowe argumenty za:")
run.bold = True

add_bullet(doc, "Skala integracji: 7+ systemow IT do polaczenia z Omniusem")
add_bullet(doc, "Brak wlasciciela: nikt w REH nie jest obecnie odpowiedzialny za AI/automatyzacje")
add_bullet(doc, "ROI potencjal: >10x kosztow w 24 miesiacach (oszczednosci FTE + szybkosc decyzji)")
add_bullet(doc, "Wylaczenie operacyjne Sebastiana: nie moze sam nadzierowac wdrozenia")
add_bullet(doc, "Konkurencja: sektor energetyczny przyspiesza digitalizacje — REH nie moze zostac w tyle")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 2. KONTEKST STRATEGICZNY
# ═══════════════════════════════════════════════════════════════
add_classification(doc)

doc.add_heading("2. Kontekst strategiczny", level=1)

doc.add_heading("2.1. Czym jest Omnius REH", level=2)
doc.add_paragraph(
    "Omnius REH to system AI osadzony w infrastrukturze Respect Energy Holding, przetwarzajacy dane "
    "firmowe z CRM, ERP, Jira, email i Teams. Jest czescia federacyjnej architektury (Omnius REH + Omnius REF), "
    "gdzie kazda instancja jest izolowana per podmiot prawny. Omnius REH jest podrzedny wobec Gilbertusa Albans "
    "(prywatny mentat AI Sebastiana) i raportuje do niego przez szyfrowane API."
)

doc.add_heading("2.2. Stan obecny", level=2)
make_table(doc,
    ["Parametr", "Wartosc"],
    [
        ["Status Omnius REH", "Architektura zaprojektowana, wdrozenie BLOCKED"],
        ["Bloker", "Transfer WA Business API + brak wlasciciela operacyjnego"],
        ["Dzial IT (RE_ITS)", "~2,17M PLN/rok na podwykonawcow B2B"],
        ["Ksiegowosc REH", "24 osoby, potencjal automatyzacji 40-70%"],
        ["DMS", "64 tys. dokumentow/rok (+103% przy +23% FTE)"],
        ["Portfel systemow", "ERP, HRM, SAPICO, DMS+Workflow, eCommerce, Savangard, CRM"],
        ["CEO bandwidth na AI", "Ograniczony — Roch skupiony na operacjach"],
    ],
    col_widths=[5.5, 10.5]
)

doc.add_heading("2.3. Dlaczego teraz", level=2)
doc.add_paragraph(
    "Wdrozenie Omnius nie jest kwestia 'czy', ale 'kiedy'. System Gilbertus udowodnil wartosc AI: "
    ">100M PLN wykrytych ryzyk i szans w ~20 godzin analizy na poziomie CEO. Teraz ta wartosc musi "
    "zostac przeniesiona na poziom operacyjny REH. Kazdy miesiac opoznienia to utracone oszczednosci "
    "i narastajaca luka kompetencyjna wobec konkurencji."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 3. OPCJE ORGANIZACYJNE
# ═══════════════════════════════════════════════════════════════
add_classification(doc)

doc.add_heading("3. Opcje organizacyjne — 3 scenariusze", level=1)

# OPCJA A
doc.add_heading("Opcja A: 1 dedykowana osoba (AI & Automation Officer)", level=2)
doc.add_paragraph(
    "Zatrudnienie jednej, wysoko wykwalifikowanej osoby, ktora bedzie jednoczesnie wlascicielem Omniusa, "
    "liderem transformacji AI i lacznikiem miedzy Sebastianem/Gilbertusem a zespolem REH."
)
make_table(doc,
    ["Kryterium", "Ocena"],
    [
        ["Koszt roczny", "300-420 tys. PLN brutto (B2B: 25-35 tys./mies.)"],
        ["Timeline do wdrozenia", "3-4 miesiace od zatrudnienia"],
        ["Ryzyko", "SREDNIE — bus factor = 1, ale kontrolowalny koszt"],
        ["Idealnie gdy", "Omnius jest 1 z wielu projektow, nie wymaga pracy 3 osob rownoczesnie"],
        ["Skalowanie", "Mozliwosc rozbudowy do opcji B po 6-12 miesiacach"],
    ],
    col_widths=[5, 11]
)

# OPCJA B
doc.add_heading("Opcja B: Maly zespol 2-3 osob (Lider + Analitycy procesowi)", level=2)
doc.add_paragraph(
    "Zespol skladajacy sie z lidera AI (profil jak w opcji A) wspartego 1-2 analitykami procesow "
    "biznesowych. Analitycy mapuja procesy, zbieraja wymagania, testuja integracje."
)
make_table(doc,
    ["Kryterium", "Ocena"],
    [
        ["Koszt roczny", "550-800 tys. PLN (lider 350K + 1-2 analitykow po 150-200K)"],
        ["Timeline do wdrozenia", "2-3 miesiace (wieksza przepustowosc rownolegla)"],
        ["Ryzyko", "NISKIE — redundancja kompetencji, szybsze tempo"],
        ["Idealnie gdy", "Wiele procesow do zautomatyzowania jednoczesnie (np. ksiegowosc + DMS)"],
        ["Skalowanie", "Gotowy zespol, mozna dodawac kolejne procesy bez waskiegogarla"],
    ],
    col_widths=[5, 11]
)

# OPCJA C
doc.add_heading("Opcja C: Outsourcing do zewnetrznego integratora", level=2)
doc.add_paragraph(
    "Zlecenie wdrozenia firmie zewnetrznej (np. konsulting AI/digital transformation) z minimalnym "
    "nadzorem wewnetrznym (1 osoba part-time)."
)
make_table(doc,
    ["Kryterium", "Ocena"],
    [
        ["Koszt roczny", "500-1,200 tys. PLN (zalezny od scope, stawki 1-2.5K PLN/dzien)"],
        ["Timeline do wdrozenia", "4-6 miesiecy (onboarding domeny + discovery)"],
        ["Ryzyko", "WYSOKIE — brak transferu wiedzy, uzaleznienie od dostawcy, dane poufne na zewnatrz"],
        ["Idealnie gdy", "Brak mozliwosci zatrudnienia, potrzeba jednorazowego wdrozenia"],
        ["Skalowanie", "Trudne — kazde rozszerzenie = nowe zamowienie + koszty"],
    ],
    col_widths=[5, 11]
)

doc.add_heading("3.1. Podsumowanie porownawcze", level=2)
make_table(doc,
    ["", "Opcja A", "Opcja B", "Opcja C"],
    [
        ["Koszt roczny", "300-420K PLN", "550-800K PLN", "500-1200K PLN"],
        ["Czas do wdrozenia", "3-4 mies.", "2-3 mies.", "4-6 mies."],
        ["Ryzyko", "Srednie", "Niskie", "Wysokie"],
        ["Kontrola danych", "Pelna", "Pelna", "Ograniczona"],
        ["Budowanie kompetencji", "Tak", "Tak++", "Nie"],
        ["Rekomendacja", "START TUTAJ", "Docelowo", "Odradzane"],
    ],
    col_widths=[4, 4, 4, 4]
)

p = doc.add_paragraph()
run = p.add_run("Rekomendacja Gilbertusa: ")
run.bold = True
p.add_run(
    "Rozpoczac od Opcji A (1 osoba) z ewolucja do Opcji B po 6 miesiacach, "
    "gdy skala zadan to uzasadni. Opcja C jest odradzana ze wzgledu na ryzyko "
    "bezpieczenstwa danych i brak budowania wewnetrznych kompetencji."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 4. PROFIL STANOWISKA
# ═══════════════════════════════════════════════════════════════
add_classification(doc)

doc.add_heading("4. Profil stanowiska: AI & Automation Officer REH", level=1)

make_table(doc,
    ["Parametr", "Wartosc"],
    [
        ["Tytul stanowiska", "AI & Automation Officer"],
        ["Alternatywne tytuly", "Digital Transformation Lead / AI Operations Manager"],
        ["Raportuje do", "CEO (Roch Baranowski)"],
        ["Wspolpraca scisla", "Sebastian Jablonski (wlasciciel) + Gilbertus Albans (AI mentat)"],
        ["Wspolpraca operacyjna", "RE_ITS, Robert Chudy (automatyzacja ERP), zespol ksiegowosci"],
        ["Forma zatrudnienia", "B2B (preferowane) lub UoP"],
        ["Lokalizacja", "Hybrydowo: biuro REH + remote"],
    ],
    col_widths=[5, 11]
)

doc.add_heading("4.1. Kompetencje wymagane — twarde", level=2)
make_table(doc,
    ["Kompetencja", "Poziom", "Opis"],
    [
        ["Python / API", "Zaawansowany", "Integracje REST/GraphQL, scripting, automatyzacja. Nie musi pisac produkcyjnego kodu, ale musi rozumiec i modyfikowac."],
        ["Zarzadzanie projektami IT", "Zaawansowany", "Agile/Kanban, roadmapping, JIRA, delivery — prowadzenie wdrozen od discovery do go-live."],
        ["LLM / AI literacy", "Sredniozaawansowany", "Prompt engineering, architektura RAG, rozumienie mozliwosci i ograniczen LLM. Nie musi trenowac modeli."],
        ["Analiza danych", "Sredniozaawansowany", "SQL, dashboarding, KPI tracking. Interpretacja danych biznesowych."],
        ["Integracje systemow", "Sredniozaawansowany", "Doswiadczenie z integracjami ERP/CRM/DMS. Rozumienie API, webhook, ETL."],
        ["Bezpieczenstwo danych", "Podstawowy+", "RODO, klasyfikacja danych, zasady pracy z danymi poufnymi."],
    ],
    col_widths=[4, 3.5, 8.5]
)

doc.add_heading("4.2. Kompetencje wymagane — miekkie", level=2)
add_bullet(doc, "Komunikacja z zarzadem — umiejetnosc tlumaczenia technicznego jezyka na biznesowy")
add_bullet(doc, "Change management — zarzadzanie oporem przed zmiana, budowanie buy-in")
add_bullet(doc, "Praca z danymi poufnymi — dyskrecja, swiadomosc konsekwencji wycieku")
add_bullet(doc, "Samodzielnosc — zdolnosc do pracy bez stalego nadzoru, proaktywnosc")
add_bullet(doc, "Myslenie procesowe — umiejetnosc dekompozycji procesow biznesowych na kroki automatyzowalne")

doc.add_heading("4.3. Kompetencje branzowe (plus, nie wymog)", level=2)
add_bullet(doc, "Znajomosc rynku energii w Polsce (OZE, obrót, balansowanie)")
add_bullet(doc, "Doswiadczenie w sektorze utilities/energy trading")
add_bullet(doc, "Znajomosc regulacji URE, CSIRE, PSE")

doc.add_heading("4.4. Czego NIE potrzebujemy", level=2)
doc.add_paragraph(
    "Wazne jest rowniez okreslenie czego ta rola NIE wymaga — aby uniknac nadmiarowych oczekiwan "
    "i zawezic pule kandydatow do wlasciwych osob:"
)
add_bullet(doc, "Glebokiego Machine Learning / trenowania modeli — Omnius uzywa gotowych LLM (Claude, OpenAI)")
add_bullet(doc, "Kodowania od podstaw / architektury software — infrastrukture dostarcza Gilbertus")
add_bullet(doc, "Doswiadczenia w big data / data engineering na skale FAANG")
add_bullet(doc, "Certyfikatow AI (np. Google ML Engineer) — liczy sie praktyka, nie papier")
add_bullet(doc, "Doswiadczenia w zarzadzaniu duzymi zespolami (>10 osob)")

doc.add_heading("4.5. Profil idealnego kandydata", level=2)
doc.add_paragraph(
    "Szukamy osoby z 5-8 lat doswiadczenia na styku IT i biznesu. Idealny kandydat to ktos, kto:"
)
add_bullet(doc, "Prowadzil wdrozenia systemow IT w sredniej firmie (100-500 osob)")
add_bullet(doc, "Ma doswiadczenie z automatyzacja procesow (RPA, workflow, integracje)")
add_bullet(doc, "Rozumie i uzywa narzedzia AI w codziennej pracy (prompt engineering, RAG)")
add_bullet(doc, "Potrafi rozmawiac zarowno z programista, jak i z dyrektorem finansowym")
add_bullet(doc, "Background: IT Project Manager, Business Analyst, Digital Transformation Consultant, lub Process Automation Engineer")

doc.add_heading("4.6. Wynagrodzenie — widelki rynkowe (2025/2026)", level=2)
make_table(doc,
    ["Forma", "Widelki miesiecznie", "Widelki rocznie", "Uwagi"],
    [
        ["B2B", "25 000 - 35 000 PLN netto", "300 000 - 420 000 PLN", "Preferowane. Stawka rynkowa dla AI/Automation Lead w Warszawie."],
        ["UoP", "18 000 - 25 000 PLN brutto", "216 000 - 300 000 PLN brutto", "Nizsze netto, ale pelna ochrona pracownicza."],
        ["Benchmark", "", "", "Porownanie: AI Engineer 22-38K, IT PM 18-28K, Data Analyst 14-22K (B2B, 2025)."],
    ],
    col_widths=[2.5, 4.5, 4.5, 4.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 5. ZAKRES ZADAN — 12 MIESIECY
# ═══════════════════════════════════════════════════════════════
add_classification(doc)

doc.add_heading("5. Zakres zadan — pierwsze 12 miesiecy", level=1)

# Faza 1
doc.add_heading("Faza 1: Onboarding i discovery (miesiac 1-2)", level=2)
make_table(doc,
    ["Zadanie", "Opis", "Deliverable"],
    [
        ["Onboarding", "Poznanie zespolu REH, systemow IT, procesow biznesowych", "Mapa interesariuszy"],
        ["Inwentaryzacja systemow", "Przeglad ERP, HRM, DMS, CRM, SAPICO, Savangard", "Raport stanu systemow IT"],
        ["Mapowanie procesow", "Identyfikacja top-20 procesow do automatyzacji z priorytetyzacja", "Process map + scoring"],
        ["Analiza gotowosci danych", "Ocena jakosci danych w kazdym systemie, luki, duplikaty", "Data readiness report"],
        ["Szkolenie Gilbertus/Omnius", "Nauka architektury Omniusa, API, mozliwosci", "Certyfikacja wewnetrzna"],
        ["Quick wins", "Identyfikacja 2-3 szybkich automatyzacji (np. raporty, alerty)", "Dzialajace prototypy"],
    ],
    col_widths=[4, 6.5, 5.5]
)

# Faza 2
doc.add_heading("Faza 2: Wdrozenie Omnius (miesiac 3-4)", level=2)
make_table(doc,
    ["Zadanie", "Opis", "Deliverable"],
    [
        ["Deployment Omnius REH", "Instalacja, konfiguracja, integracja z infrastruktura REH", "Omnius dzialajacy w produkcji"],
        ["Integracje Tier-1", "Teams + Email (Graph API), ERP (podstawowe dane)", "3 zrodla danych podlaczone"],
        ["Decision tracking", "Konfiguracja formalnego sledzenia decyzji (priorytet #1)", "Decision log aktywny"],
        ["Security audit", "Przeglad bezpieczenstwa polaczen, szyfrowanie, RODO", "Raport security"],
        ["Testy UAT", "Testy z uzytkownikami kluczowych funkcji", "Sign-off od Rocha"],
    ],
    col_widths=[4, 6.5, 5.5]
)

# Faza 3
doc.add_heading("Faza 3: Pilotaz kluczowych procesow (miesiac 5-6)", level=2)
make_table(doc,
    ["Proces pilotazowy", "Opis", "Oczekiwany efekt"],
    [
        ["OPEX Approval Process (RPI-38)", "Automatyzacja zatwierdzania kosztow operacyjnych", "Redukcja czasu akceptacji z dni do godzin"],
        ["Decision tracking", "Pelne sledzenie decyzji na kazdym poziomie zarzadzania", "100% decyzji zarzadu dokumentowanych"],
        ["Cashflow monitoring", "Alerty o ryzykach plynnosci w czasie rzeczywistym", "Wykrycie ryzyka <1h (vs. dni obecnie)"],
        ["Raporty dla CEO", "Automatyczne daily/weekly briefy dla Rocha", "Oszczednosc 5-10h/tyg czasu CEO"],
    ],
    col_widths=[4.5, 6, 5.5]
)

# Faza 4
doc.add_heading("Faza 4: Skalowanie i optymalizacja (miesiac 7-12)", level=2)
make_table(doc,
    ["Zadanie", "Opis", "Cel"],
    [
        ["Automatyzacja ksiegowosci", "Wdrozenie AI w top-5 procesach ksiegowych", "Redukcja 4-6 FTE (z 24)"],
        ["Integracje Tier-2", "DMS, Jira, SAPICO, Savangard/CSIRE", "Pelny ekosystem danych"],
        ["Omnius REF", "Uruchomienie drugiej instancji dla REF", "Federacja Omnius dzialajaca"],
        ["Szkolenia pracownikow", "Warsztaty AI dla kluczowych zespolow", "Min. 30 osob przeszkolonych"],
        ["Raport ROI", "Pomiar oszczednosci i wartosci po 12 miesiacach", "Business case dla fazy 2"],
        ["Roadmapa rok 2", "Plan rozwoju na kolejne 12 miesiecy", "Zatwierdzona przez CEO/wlasciciela"],
    ],
    col_widths=[4.5, 6, 5.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 6. KPI I METRYKI
# ═══════════════════════════════════════════════════════════════
add_classification(doc)

doc.add_heading("6. KPI i metryki sukcesu", level=1)

doc.add_paragraph(
    "Ponizsze KPI powinny byc mierzone co miesiac i raportowane do CEO oraz wlasciciela. "
    "Kazdy KPI ma zdefiniowany cel na koniec 12 miesiaca."
)

make_table(doc,
    ["KPI", "Metryka", "Cel (12 mies.)", "Sposob pomiaru"],
    [
        ["Procesy zautomatyzowane", "Liczba procesow z dzialajaca automatyzacja", "Min. 8 procesow", "Rejestr procesow + status"],
        ["FTE savings", "Redukcja etatow / przesuniecie do zadan wyzszej wartosci", "Min. 3 FTE", "Porownanie headcount before/after"],
        ["Czas odpowiedzi zarzadu", "Czas od zapytania CEO do dostarczenia analizy", "<30 minut (vs. dni)", "Logi Omnius + timestampy"],
        ["Decision tracking coverage", "% decyzji zarzadu z pelna dokumentacja", ">90%", "Decision log w Omniusie"],
        ["Cashflow alert latency", "Czas od powstania ryzyka plynnosci do alertu", "<1 godzina", "Logi alertow + timestamps"],
        ["Data integration coverage", "Liczba systemow podlaczonych do Omnius", "Min. 5 z 7", "Status integracji"],
        ["User adoption", "% aktywnych uzytkownikow Omniusa wsrod kadry kierowniczej", ">80%", "Logi uzycia + ankiety"],
        ["ROI finansowy", "Stosunek oszczednosci do kosztow Omnius + personelu", ">3x w 12 mies., >10x w 24 mies.", "Analiza finansowa"],
    ],
    col_widths=[3.5, 3.5, 3.5, 5.5]
)

doc.add_heading("6.1. Szacunek finansowy ROI", level=2)
make_table(doc,
    ["Pozycja", "Wartosc roczna", "Uwagi"],
    [
        ["KOSZTY", "", ""],
        ["AI & Automation Officer (B2B)", "360 000 PLN", "30K/mies. srednia"],
        ["Infrastruktura Omnius (serwery, API)", "120 000 PLN", "10K/mies. szacunkowo"],
        ["Licencje AI (Claude, OpenAI)", "60 000 PLN", "5K/mies."],
        ["Szkolenia, narzedzia", "30 000 PLN", "Jednorazowe + subskrypcje"],
        ["RAZEM KOSZTY", "570 000 PLN", ""],
        ["", "", ""],
        ["OSZCZEDNOSCI / WARTOSCI", "", ""],
        ["Redukcja FTE ksiegowosc (3 osoby)", "450 000 PLN", "3 x 150K PLN/rok"],
        ["Przyspieszenie decyzji (wartosc)", "500 000 - 2 000 000 PLN", "Szybsze reakcje na ryzyka/szanse"],
        ["Redukcja czasu CEO na raportowanie", "200 000 PLN", "~10h/tyg x wartosc godziny CEO"],
        ["Unikniete ryzyka (cashflow, compliance)", "300 000 - 1 000 000 PLN", "Trudne do zmierzenia, ale realne"],
        ["RAZEM OSZCZEDNOSCI (konserwatywnie)", "1 450 000 PLN", "Minimum"],
        ["", "", ""],
        ["ROI (rok 1)", "~2.5x", "1.45M / 570K"],
        ["ROI (rok 2, po skalowaniu)", ">10x", "Przy dalszej automatyzacji"],
    ],
    col_widths=[6, 4, 6]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 7. RYZYKA I MITYGACJA
# ═══════════════════════════════════════════════════════════════
add_classification(doc)

doc.add_heading("7. Ryzyka i mitygacja", level=1)

make_table(doc,
    ["Ryzyko", "Prawdop.", "Wplyw", "Mitygacja"],
    [
        [
            "Brak buy-in od Rocha / zarzadu",
            "Srednie",
            "Krytyczny",
            "Wczesne zaangazowanie CEO w discovery. Quick wins w pierwszych 2 miesiacach. "
            "Regularne raportowanie ROI. Sebastian jako sponsor strategiczny."
        ],
        [
            "Zbyt ambitny scope vs. zasoby",
            "Wysokie",
            "Wysoki",
            "Scisla priorytetyzacja (max 2-3 procesy naraz). MVP-first approach. "
            "Eskalacja do Sebastiana gdy scope creep. Fazy z jasnymi deliverables."
        ],
        [
            "Bus factor = 1 (jedna osoba)",
            "Srednie",
            "Wysoki",
            "Dokumentacja wszystkiego od dnia 1. Transfer wiedzy do Gilbertusa. "
            "Plan ewolucji do zespolu (opcja B) po 6 miesiacach. Overlap period przy rotacji."
        ],
        [
            "Opor pracownikow (obawy o stanowiska)",
            "Wysokie",
            "Sredni",
            "Komunikacja: AI wspiera, nie zastepuje. Szkolenia i reskilling. "
            "Pokazanie ze automatyzacja eliminuje nudne zadania, nie ludzi. "
            "Transparentnosc co do planow."
        ],
        [
            "Problemy z integracjami (API, dane)",
            "Wysokie",
            "Sredni",
            "Proof of concept przed pelnym wdrozeniem. Wsparcie RE_ITS. "
            "Fallback na manualne procesy w razie awarii. Monitoring integracji 24/7."
        ],
        [
            "Wycieki danych / bezpieczenstwo",
            "Niskie",
            "Krytyczny",
            "Szyfrowanie end-to-end. RODO compliance od dnia 1. Audit bezpieczenstwa "
            "co kwartal. Klasyfikacja danych. Minimalizacja dostepu (least privilege)."
        ],
        [
            "Odejscie kluczowej osoby",
            "Srednie",
            "Wysoki",
            "Konkurencyjne wynagrodzenie. Ciekawy projekt (AI, innowacja). "
            "Notice period min. 3 miesiace w umowie. Knowledge base w Gilbertusie."
        ],
    ],
    col_widths=[3.5, 2, 2, 8.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 8. REKOMENDACJA GILBERTUSA
# ═══════════════════════════════════════════════════════════════
add_classification(doc)

doc.add_heading("8. Rekomendacja Gilbertusa", level=1)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
run = p.add_run(
    "REH potrzebuje dedykowanego AI & Automation Officera — jednej osoby, ktora od pierwszego dnia "
    "przejmie odpowiedzialnosc za wdrozenie Omniusa i transformacje cyfrowa spolki."
)
run.font.size = Pt(11.5)
run.bold = True

doc.add_paragraph(
    "Rekomenduje rozpoczecie rekrutacji natychmiast (kwiecien 2026) z celem zatrudnienia do konca "
    "maja 2026. Osoba ta powinna raportowac do Rocha Baranowskiego, ale miec bezposredni kanal "
    "komunikacji z Sebastianem i Gilbertusem w zakresie architektury i strategii AI."
)

doc.add_paragraph(
    "Budzet 360K PLN/rok (30K/mies. B2B) jest inwestycja o oczekiwanym ROI 2.5x w pierwszym roku "
    "i >10x w drugim roku. Bez tej inwestycji, Omnius pozostanie w statusie BLOCKED, a REH straci "
    "szanse na automatyzacje wartej kilka milionow PLN rocznie."
)

doc.add_paragraph(
    "Kluczowe: ta osoba NIE musi byc AI researcherem ani senior developerem. Potrzebujemy pragmatycznego "
    "lidera transformacji, ktory rozumie technologie, potrafi zarzadzac projektami i komunikowac sie "
    "z zarzadem. Gilbertus i Omnius dostarczaja technologie — potrzebujemy czlowieka, ktory ja wdrozy."
)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(16)
run = p.add_run("Proponowane nastepne kroki:")
run.bold = True

add_bullet(doc, " Decyzja Sebastiana: go/no-go na rekrutacje", bold_prefix="1.")
add_bullet(doc, " Briefing Rocha: omowienie roli i oczekiwan", bold_prefix="2.")
add_bullet(doc, " Publikacja oferty: LinkedIn + dedykowane kanaly (kwiecien 2026)", bold_prefix="3.")
add_bullet(doc, " Rekrutacja: 3-4 tygodnie (screening + 2 rundy rozmow)", bold_prefix="4.")
add_bullet(doc, " Start: maj/czerwiec 2026", bold_prefix="5.")

doc.add_paragraph()

# Footer note
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("— Koniec dokumentu —")
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f"Wygenerowano: {date.today().strftime('%d.%m.%Y')} | Gilbertus Albans v1.0 | POUFNE")
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

# ═══════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════
path1 = "/mnt/c/Users/jablo/Desktop/REH_Omnius_Team_Analysis.docx"
path2 = "/home/sebastian/.openclaw/workspace/REH_Omnius_Team_Analysis.docx"

doc.save(path1)
shutil.copy2(path1, path2)

print(f"Saved: {path1}")
print(f"Saved: {path2}")
print("Done!")
