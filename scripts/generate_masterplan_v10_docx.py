#!/usr/bin/env python3
"""
Generate Gilbertus Masterplan V10 as DOCX with architecture diagram.
Reflects state as of 2026-03-28: all agentic patterns implemented.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from PIL import Image, ImageDraw, ImageFont

# ============================================================
# Architecture diagram generation (Pillow -> PNG)
# ============================================================

def draw_rounded_rect(draw, xy, radius, fill, outline=None, width=1):
    x0, y0, x1, y1 = xy
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)

def generate_architecture_diagram(output_path):
    """Generate Gilbertus+Omnius architecture diagram as PNG."""
    W, H = 2400, 1800
    img = Image.new('RGB', (W, H), '#FFFFFF')
    draw = ImageDraw.Draw(img)

    # Try to use a system font
    try:
        font_title = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 28)
        font_heading = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 20)
        font_normal = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 16)
        font_small = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 13)
    except Exception:
        font_title = ImageFont.load_default()
        font_heading = font_title
        font_normal = font_title
        font_small = font_title

    # Colors
    C_GILB = '#1a1a2e'       # Gilbertus dark blue
    C_GILB_BG = '#e8eaf6'    # Gilbertus light bg
    C_OMNIUS = '#1b5e20'     # Omnius dark green
    C_OMNIUS_BG = '#e8f5e9'  # Omnius light bg
    C_DATA = '#0d47a1'       # Data sources blue
    C_DATA_BG = '#e3f2fd'
    C_DELIVERY = '#e65100'   # Delivery orange
    C_DELIVERY_BG = '#fff3e0'
    C_ARROW = '#455a64'

    # ── Title ──
    draw.text((W//2 - 300, 15), "GILBERTUS ALBANS + OMNIUS — Architektura Systemu", fill=C_GILB, font=font_title)
    draw.text((W//2 - 100, 50), "Stan: 2026-03-28", fill='#666666', font=font_small)

    # ── GILBERTUS main box ──
    gx, gy, gw, gh = 50, 80, 1500, 900
    draw_rounded_rect(draw, (gx, gy, gx+gw, gy+gh), 15, fill=C_GILB_BG, outline=C_GILB, width=3)
    draw.text((gx+20, gy+10), "GILBERTUS (Hetzner / WSL2)", fill=C_GILB, font=font_heading)

    # ── Layer 1: Input ──
    ly = gy + 50
    draw_rounded_rect(draw, (gx+20, ly, gx+480, ly+120), 10, fill='#fff9c4', outline='#f9a825', width=2)
    draw.text((gx+30, ly+5), "WARSTWA 1: INPUT", fill='#f57f17', font=font_heading)
    draw.text((gx+30, ly+30), "FastAPI (110+ endpointow)", fill='#333', font=font_normal)
    draw.text((gx+30, ly+50), "/ask  /timeline  /summary  /status", fill='#555', font=font_small)
    draw.text((gx+30, ly+70), "/voice/ask  /teams  /conversation", fill='#555', font=font_small)
    draw.text((gx+30, ly+90), "AskRequest + session_id + channel", fill='#555', font=font_small)

    # ── Layer 2: Understanding ──
    draw_rounded_rect(draw, (gx+500, ly, gx+960, ly+120), 10, fill='#e1f5fe', outline='#0288d1', width=2)
    draw.text((gx+510, ly+5), "WARSTWA 2: UNDERSTANDING", fill='#01579b', font=font_heading)
    draw.text((gx+510, ly+30), "Query Interpreter (Haiku)", fill='#333', font=font_normal)
    draw.text((gx+510, ly+50), "-> InterpretedQuery:", fill='#555', font=font_small)
    draw.text((gx+510, ly+70), "  question_type, analysis_depth,", fill='#555', font=font_small)
    draw.text((gx+510, ly+90), "  alternate_queries, sub_questions", fill='#555', font=font_small)

    # Arrow 1->2
    draw.line([(gx+480, ly+60), (gx+500, ly+60)], fill=C_ARROW, width=3)

    # ── Layer 3: Retrieval ──
    ly2 = ly + 140
    draw_rounded_rect(draw, (gx+20, ly2, gx+960, ly2+200), 10, fill='#f3e5f5', outline='#7b1fa2', width=2)
    draw.text((gx+30, ly2+5), "WARSTWA 3: RETRIEVAL (Agentic Patterns)", fill='#4a148c', font=font_heading)

    # Sub-boxes for patterns
    px = gx + 40
    py = ly2 + 35
    patterns = [
        ("ROUTING", "#e8eaf6", "#3f51b5", "chronology -> timeline SQL\nanalysis -> parallel\nsummary -> cache\nretrieval -> standard"),
        ("PARALLELIZATION", "#e3f2fd", "#1565c0", "ThreadPoolExecutor\n2-3 alternate queries\nfan-out + dedup merge"),
        ("ORCHESTRATOR", "#fce4ec", "#c62828", "Sub-question decomp\nParallel retrieve+answer\nSonnet synthesis"),
    ]
    for i, (name, bg, fg, desc) in enumerate(patterns):
        bx = px + i * 300
        draw_rounded_rect(draw, (bx, py, bx+280, py+155), 8, fill=bg, outline=fg, width=2)
        draw.text((bx+10, py+5), name, fill=fg, font=font_normal)
        for j, line in enumerate(desc.split('\n')):
            draw.text((bx+10, py+28 + j*18), line, fill='#333', font=font_small)

    # Qdrant + PostgreSQL
    draw_rounded_rect(draw, (px+920, py, px+1080, py+70), 8, fill='#e8f5e9', outline='#2e7d32', width=2)
    draw.text((px+930, py+5), "Qdrant", fill='#1b5e20', font=font_normal)
    draw.text((px+930, py+28), "99k vectors", fill='#555', font=font_small)
    draw.text((px+930, py+45), "embeddings", fill='#555', font=font_small)

    draw_rounded_rect(draw, (px+920, py+80, px+1080, py+155), 8, fill='#e3f2fd', outline='#1565c0', width=2)
    draw.text((px+930, py+85), "PostgreSQL", fill='#0d47a1', font=font_normal)
    draw.text((px+930, py+108), "102 tables", fill='#555', font=font_small)
    draw.text((px+930, py+125), "94k events", fill='#555', font=font_small)

    # ── Layer 4: Generation ──
    ly3 = ly2 + 215
    draw_rounded_rect(draw, (gx+20, ly3, gx+700, ly3+100), 10, fill='#fff3e0', outline='#e65100', width=2)
    draw.text((gx+30, ly3+5), "WARSTWA 4: GENERATION", fill='#bf360c', font=font_heading)
    draw.text((gx+30, ly3+30), "Claude Sonnet 4.6 (answer_question)", fill='#333', font=font_normal)
    draw.text((gx+30, ly3+50), "+ Conversation Context (sliding window 20 msg)", fill='#555', font=font_small)
    draw.text((gx+30, ly3+70), "+ Prompt Caching (ephemeral system prompt)", fill='#555', font=font_small)

    # ── Evaluator ──
    draw_rounded_rect(draw, (gx+720, ly3, gx+960, ly3+100), 10, fill='#fce4ec', outline='#c62828', width=2)
    draw.text((gx+730, ly3+5), "EVALUATOR", fill='#b71c1c', font=font_heading)
    draw.text((gx+730, ly3+30), "Haiku: relevance,", fill='#333', font=font_small)
    draw.text((gx+730, ly3+48), "grounding, depth", fill='#333', font=font_small)
    draw.text((gx+730, ly3+66), "score < 0.6 -> retry", fill='#c62828', font=font_small)
    draw.text((gx+730, ly3+84), "with feedback", fill='#c62828', font=font_small)

    # Arrow gen -> eval
    draw.line([(gx+700, ly3+50), (gx+720, ly3+50)], fill='#c62828', width=3)

    # ── Layer 5: Action ──
    ly4 = ly3 + 115
    draw_rounded_rect(draw, (gx+20, ly4, gx+700, ly4+100), 10, fill='#efebe9', outline='#4e342e', width=2)
    draw.text((gx+30, ly4+5), "WARSTWA 5: ACTION (LangGraph HITL)", fill='#3e2723', font=font_heading)
    draw.text((gx+30, ly4+30), "propose -> [WhatsApp approval] -> execute/reject", fill='#333', font=font_normal)
    draw.text((gx+30, ly4+50), "5-tier authority (INFORM -> NEVER_ALONE)", fill='#555', font=font_small)
    draw.text((gx+30, ly4+70), "Actions: email, ticket, meeting, WhatsApp, Omnius cmd", fill='#555', font=font_small)

    # ── Automation ──
    draw_rounded_rect(draw, (gx+720, ly4, gx+960, ly4+100), 10, fill='#e0f2f1', outline='#00695c', width=2)
    draw.text((gx+730, ly4+5), "AUTOMATION", fill='#004d40', font=font_heading)
    draw.text((gx+730, ly4+30), "39 cron jobs", fill='#333', font=font_small)
    draw.text((gx+730, ly4+48), "24 extraction workers", fill='#333', font=font_small)
    draw.text((gx+730, ly4+66), "Auto-actions triggers", fill='#333', font=font_small)
    draw.text((gx+730, ly4+84), "Cost budget guards", fill='#333', font=font_small)

    # ── MCP Tools ──
    draw_rounded_rect(draw, (gx+980, ly4-50, gx+1480, ly4+100), 10, fill='#f5f5f5', outline='#616161', width=2)
    draw.text((gx+990, ly4-45), "MCP SERVER (43 tools)", fill='#212121', font=font_heading)
    mcp_lines = [
        "Core (11): ask, brief, decide, timeline...",
        "People (9): sentiment, commitments, delegation...",
        "Business (8): process_intel, workforce, scenarios...",
        "Finance (3): finance, market, competitors",
        "Omnius (4): ask, command, status, bridge",
        "System (11): status, costs, crons, authority...",
    ]
    for i, line in enumerate(mcp_lines):
        draw.text((gx+990, ly4-20+i*20), line, fill='#555', font=font_small)

    # ── DATA SOURCES (left side) ──
    dsy = gy + gh + 30
    draw_rounded_rect(draw, (50, dsy, 780, dsy+200), 12, fill=C_DATA_BG, outline=C_DATA, width=2)
    draw.text((60, dsy+5), "ZRODLA DANYCH (10+)", fill=C_DATA, font=font_heading)
    sources = [
        "Email (Graph API)      Teams (Graph API)      WhatsApp (OpenClaw)",
        "Plaud Pin S (audio)    Calendar (Graph API)   Claude Code sessions",
        "Documents (upload)     Spreadsheets           ChatGPT export",
        "WhatsApp Live (JSONL)  Audio transcripts (Whisper)",
    ]
    for i, s in enumerate(sources):
        draw.text((70, dsy+35+i*22), s, fill='#333', font=font_small)
    draw.text((70, dsy+130), "-> 33k documents -> 100k chunks -> 36k entities -> 94k events", fill=C_DATA, font=font_normal)
    draw.text((70, dsy+155), "Ingestion co 5 min | Extraction co 30 min | 100% coverage", fill='#555', font=font_small)

    # ── DELIVERY (right side) ──
    draw_rounded_rect(draw, (800, dsy, 1550, dsy+200), 12, fill=C_DELIVERY_BG, outline=C_DELIVERY, width=2)
    draw.text((810, dsy+5), "DELIVERY", fill=C_DELIVERY, font=font_heading)
    deliveries = [
        "WhatsApp (primary): briefs, alerts, HITL approval",
        "Teams Bot: RBAC-filtered answers, task management",
        "Voice REST: Whisper STT -> /ask -> TTS",
        "HTTP API: 110+ endpoints, MCP server (stdio)",
        "Morning Brief (7:00) | Weekly Synthesis (Sun 20:00)",
        "Meeting Prep (15 min before) | Response Drafter",
    ]
    for i, d in enumerate(deliveries):
        draw.text((820, dsy+35+i*22), d, fill='#333', font=font_small)
    draw.text((820, dsy+175), "Conversation Memory: sliding window 20 msg / 8000 chars", fill=C_DELIVERY, font=font_small)

    # ── OMNIUS box (bottom right) ──
    ox, oy = 1570, 80
    ow, oh = 780, 900
    draw_rounded_rect(draw, (ox, oy, ox+ow, oy+oh), 15, fill=C_OMNIUS_BG, outline=C_OMNIUS, width=3)
    draw.text((ox+20, oy+10), "OMNIUS (per-company infra)", fill=C_OMNIUS, font=font_heading)

    # Omnius REF
    draw_rounded_rect(draw, (ox+20, oy+45, ox+370, oy+250), 10, fill='#c8e6c9', outline='#2e7d32', width=2)
    draw.text((ox+30, oy+50), "OMNIUS REF", fill='#1b5e20', font=font_heading)
    draw.text((ox+30, oy+75), "Respect Energy Fuels", fill='#333', font=font_normal)
    draw.text((ox+30, oy+100), "CEO: Krystian Juchacz", fill='#555', font=font_small)
    draw.text((ox+30, oy+118), "Board: Edgar, Witold", fill='#555', font=font_small)
    draw.text((ox+30, oy+136), "Operator: Michal Schulte", fill='#555', font=font_small)
    draw.text((ox+30, oy+160), "RBAC: 7 roles, classifications", fill='#555', font=font_small)
    draw.text((ox+30, oy+178), "Data: SharePoint, Teams, Email", fill='#555', font=font_small)
    draw.text((ox+30, oy+196), "Azure AD SSO + Teams Bot", fill='#555', font=font_small)
    draw.text((ox+30, oy+220), "Governance: non-regression", fill='#555', font=font_small)

    # Omnius REH
    draw_rounded_rect(draw, (ox+390, oy+45, ox+740, oy+250), 10, fill='#c8e6c9', outline='#2e7d32', width=2)
    draw.text((ox+400, oy+50), "OMNIUS REH", fill='#1b5e20', font=font_heading)
    draw.text((ox+400, oy+75), "Respect Energy Holding", fill='#333', font=font_normal)
    draw.text((ox+400, oy+100), "CEO: Roch Baranowski", fill='#555', font=font_small)
    draw.text((ox+400, oy+118), "(planned — clone REF)", fill='#888', font=font_small)
    draw.text((ox+400, oy+160), "Same architecture:", fill='#555', font=font_small)
    draw.text((ox+400, oy+178), "FastAPI + PostgreSQL + Qdrant", fill='#555', font=font_small)
    draw.text((ox+400, oy+196), "RBAC + Governance + Audit", fill='#555', font=font_small)

    # Omnius arch details
    omy = oy + 270
    omnius_details = [
        ("Architektura Omnius", C_OMNIUS, font_heading),
        ("FastAPI per tenant (port 8001/8002)", '#333', font_normal),
        ("PostgreSQL: omnius_* tables (roles, users,", '#555', font_small),
        ("  documents, chunks, permissions, audit_log,", '#555', font_small),
        ("  operator_tasks, config, api_keys)", '#555', font_small),
        ("Qdrant: per-tenant vectors", '#555', font_small),
        ("", '#555', font_small),
        ("Kluczowe zasady:", C_OMNIUS, font_normal),
        ("1. One-way: Gilbertus -> Omnius (nigdy odwrotnie)", '#333', font_small),
        ("2. Data isolation: dane firmy na infra firmy", '#333', font_small),
        ("3. RBAC: operator ZERO access do danych biznes.", '#333', font_small),
        ("4. Governance: CEO nie moze usunac features", '#333', font_small),
        ("5. Audit: immutable log kazdej operacji", '#333', font_small),
        ("", '#555', font_small),
        ("Gilbertus Bridge:", C_OMNIUS, font_normal),
        ("cross_tenant_search() — query REF+REH, merge", '#333', font_small),
        ("aggregated_dashboard() — unified stats", '#333', font_small),
        ("cross_tenant_audit() — unified audit trail", '#333', font_small),
        ("sync_all_tenants() — trigger all data sync", '#333', font_small),
    ]
    for i, (text, color, fnt) in enumerate(omnius_details):
        draw.text((ox+30, omy + i*22), text, fill=color, font=fnt)

    # ── Arrow Gilbertus -> Omnius ──
    draw.line([(gx+gw, gy+200), (ox, oy+200)], fill=C_ARROW, width=4)
    draw.text((gx+gw+5, gy+180), "API Key auth", fill=C_ARROW, font=font_small)
    draw.text((gx+gw+5, gy+200), "-> commands", fill=C_ARROW, font=font_small)
    draw.text((gx+gw+5, gy+220), "-> config push", fill=C_ARROW, font=font_small)

    # ── Arrow Sources -> Gilbertus ──
    draw.line([(400, dsy), (400, gy+gh)], fill=C_DATA, width=3)

    # ── Arrow Gilbertus -> Delivery ──
    draw.line([(1100, gy+gh), (1100, dsy)], fill=C_DELIVERY, width=3)

    # Save
    img.save(output_path, 'PNG', quality=95)
    return output_path


# ============================================================
# DOCX Generation
# ============================================================

doc = Document()

# Styles
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)


def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    return h


def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)
            for p in row_cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    return table


# ============================================================
# TITLE
# ============================================================
title = doc.add_heading('GILBERTUS ALBANS — Masterplan V10 "Agentic Architecture"', 0)
for run in title.runs:
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

doc.add_paragraph('Stan na: 2026-03-28 15:30 CET')
doc.add_paragraph('Projekt: Gilbertus Albans — prywatny mentat AI Sebastiana Jablonskiego')
doc.add_paragraph('Omnius — multi-tenant corporate AI platform (REH + REF)')
doc.add_paragraph('')

# ============================================================
# BASELINE
# ============================================================
add_heading('Aktualny baseline (28.03.2026)', 1)

add_table(
    ['Metryka', 'Wartosc', 'Zmiana vs V5'],
    [
        ['MCP tools', '43', '+7 (z 36)'],
        ['DB tables', '102', '+38 (z 64)'],
        ['App modules', '131+', '+55+'],
        ['Scripts', '107', '+29'],
        ['Cron jobs', '39 (registered)', '-7 (konsolidacja)'],
        ['API endpoints', '110+', '+15'],
        ['Documents', '33,617', '+daly ingestion'],
        ['Chunks', '99,905', '+daly ingestion'],
        ['Entities', '36,271', '+daly extraction'],
        ['Events', '94,370', '+daly extraction'],
        ['Extraction coverage', '~100%', 'utrzymane'],
        ['Agentic patterns', '5/5 Anthropic', 'NEW (Phase 1-3)'],
    ]
)

# ============================================================
# AGENTIC ARCHITECTURE
# ============================================================
add_heading('Architektura Agentic (Anthropic "Building Effective Agents")', 1)

doc.add_paragraph(
    'Gilbertus implementuje WSZYSTKIE 5 wzorcow z whitepapera Anthropic '
    '"Building Effective Agents" + Autonomous Agent feedback loop.'
)

add_table(
    ['Wzorzec Anthropic', 'Implementacja w Gilbertus', 'Feature flag'],
    [
        ['1. Prompt Chaining', 'interpret_query -> search_chunks -> answer_question (3 kroki)', 'zawsze aktywny'],
        ['2. Routing', 'Deep Routing: chronology->timeline SQL, analysis->parallel, summary->cache', 'ENABLE_DEEP_ROUTING'],
        ['3. Parallelization', 'ThreadPoolExecutor: 2-3 alternate queries rownolegly retrieval + dedup', 'ENABLE_PARALLEL_RETRIEVAL'],
        ['4. Orchestrator-Workers', 'Sub-question decomposition: multi-entity -> parallel sub-Q -> synthesis', 'ENABLE_ORCHESTRATOR'],
        ['5. Evaluator-Optimizer', 'Haiku eval (relevance/grounding/depth), retry z feedback jesli score<0.6', 'ENABLE_ANSWER_EVAL'],
        ['Agent Feedback Loop', 'Conversation-aware: interpreter celuje w luki z poprzedniej odpowiedzi', 'ENABLE_CONV_AWARE'],
    ]
)

# ============================================================
# 5 WARSTW PIPELINE
# ============================================================
add_heading('5 warstw pipeline /ask', 1)

add_table(
    ['#', 'Warstwa', 'Technologia', 'Co robi'],
    [
        ['1', 'Input', 'FastAPI + AskRequest', 'Walidacja, channel defaults, session_id, conversation history (sliding window)'],
        ['2', 'Understanding', 'Claude Haiku (query_interpreter)', 'Normalizacja query, date extraction, question_type, analysis_depth, alternate_queries, sub_questions'],
        ['3', 'Retrieval', 'Qdrant + PostgreSQL + Router', 'Deep routing per question_type, parallel retrieval, timeline SQL merge, dedup'],
        ['4', 'Generation', 'Claude Sonnet + Evaluator', 'Answer z conversation context + answer eval gate (Haiku) + retry z feedback'],
        ['5', 'Action', 'LangGraph StateGraph + HITL', 'Propozycja -> WhatsApp approval -> execute/reject, 5-tier authority'],
    ]
)

# ============================================================
# GILBERTUS MODULY
# ============================================================
add_heading('Moduly Gilbertusa', 1)

add_heading('Intelligence Layer (13 modulow)', 2)
add_table(
    ['Modul', 'Co robi'],
    [
        ['Commitment Tracker', 'Ekstrakcja obietnic + overdue detection + fulfillment scan'],
        ['Meeting Prep Brief', 'Auto-brief 30 min przed spotkaniem (WhatsApp)'],
        ['Meeting Minutes Generator', 'Structured minutes z nagran Plaud'],
        ['Smart Response Drafter', 'Auto-odpowiedzi email/Teams/WhatsApp per standing order'],
        ['Weekly Executive Synthesis', 'Niedzielny raport strategiczny'],
        ['Sentiment Trend Monitor', 'Sentiment per osoba (1-5) tygodniowo'],
        ['Wellbeing Monitor', 'Wellbeing score Sebastiana (stress, family, health)'],
        ['Delegation Effectiveness', 'Completion rate, on-time rate per osoba'],
        ['Communication Network', 'Graf komunikacji, silosy, bottlenecki'],
        ['Predictive Alerts', 'Predykcja eskalacji i deadline risks'],
        ['Scenario Analyzer', 'Impact analysis, replacement planning, financial cost'],
        ['Market Intelligence', 'Ceny energii, regulacje URE, przetargi OZE'],
        ['Competitor Intelligence', 'KRS, hiring patterns, media monitoring'],
    ]
)

add_heading('Faza F: Zastypowalnosc (zrealizowana 28.03)', 2)
add_table(
    ['Modul', 'Co robi'],
    [
        ['Process Intel (18 akcji)', 'Dashboard, process mining, app analysis, tech radar, tech roadmap'],
        ['Workforce Analysis', 'Employee automation potential, replacement risk, CEO-only access'],
        ['Tech Radar', 'Technology scanning, solution ranking, alignment scoring'],
        ['App Deep Analysis', 'Per-app cost/user/replacement analysis'],
    ]
)

add_heading('Command & Control', 2)
add_table(
    ['Modul', 'Co robi'],
    [
        ['Authority Framework', '5 poziomow (INFORM -> NEVER_ALONE), 22 kategorie'],
        ['Action Pipeline', 'LangGraph StateGraph, Postgres checkpointing, HITL'],
        ['Delegation Chain', 'Deleguj -> trackuj -> remind -> eskaluj -> complete'],
        ['Decision Intelligence', 'Auto-capture decyzji, confidence calibration, bias detection'],
        ['Rule Reinforcement', 'Self-improving rules: effectiveness scoring, conflict detection'],
        ['Auto-Actions', 'Market/competitor/goal/commitment triggers -> proposals'],
    ]
)

# ============================================================
# OMNIUS
# ============================================================
add_heading('Omnius — Multi-tenant Corporate AI Platform', 1)

doc.add_paragraph(
    'Omnius to firmowy AI agent per spolka, odizolowany od danych prywatnych. '
    'Gilbertus jest kontrolerem (admin level 99), Omnius podlega governance.'
)

add_table(
    ['Aspekt', 'Omnius REF', 'Omnius REH'],
    [
        ['Firma', 'Respect Energy Fuels', 'Respect Energy Holding'],
        ['CEO', 'Krystian Juchacz', 'Roch Baranowski'],
        ['Status', 'Aktywny (Azure)', 'Planowany (klon REF)'],
        ['Infra', 'FastAPI + PostgreSQL + Qdrant', 'Identyczna'],
        ['Auth', 'Azure AD SSO + API keys', 'Azure AD SSO + API keys'],
        ['RBAC', '7 rol (gilbertus_admin -> specialist)', 'Identyczne'],
        ['Governance', 'Non-regression, frozen features', 'Identyczne'],
        ['Dane', 'SharePoint, Teams, Email, Plaud', 'SharePoint, Teams, Email, Plaud'],
        ['Teams Bot', 'Aktywny (task management)', 'Planowany'],
    ]
)

add_heading('Kluczowe zasady Omnius', 2)
doc.add_paragraph(
    '1. ONE-WAY FLOW: Gilbertus -> Omnius (nigdy odwrotnie)\n'
    '2. DATA ISOLATION: dane firmy na infra firmy\n'
    '3. RBAC: operator (Michal Schulte) ZERO access do danych biznesowych\n'
    '4. GOVERNANCE: CEO nie moze usunac ani zredukowac features\n'
    '5. AUDIT: immutable log kazdej operacji\n'
    '6. BRIDGE: Gilbertus widzi obie firmy (cross_tenant_search, aggregated_dashboard)'
)

# ============================================================
# MCP TOOLS
# ============================================================
add_heading('MCP Tools (43)', 1)

add_table(
    ['Grupa', 'Narzedzia', 'Ilosc'],
    [
        ['Core', 'ask, brief, decide, propose_action, pending_actions, timeline, alerts, router', '8'],
        ['People', 'people, sentiment, commitments, delegation, delegation_chain, response_stats, network, meeting_prep, evaluate', '9'],
        ['Business', 'opportunities, inefficiency, correlate, process_intel, workforce_analysis, org_health, scenarios, goals, decision_patterns', '9 (8?)'],
        ['Finance', 'finance, market, competitors', '3'],
        ['Omnius', 'omnius_ask, omnius_command, omnius_status, omnius_bridge', '4'],
        ['System', 'status, db_stats, lessons, costs, crons, authority, self_rules, wellbeing, calendar, summary', '10'],
    ]
)

# ============================================================
# CRON JOBS
# ============================================================
add_heading('Cron Jobs (39)', 1)

add_table(
    ['Kategoria', 'Joby', 'Harmonogram'],
    [
        ['Ingestion', 'live_ingest, plaud_monitor, sync_corporate, archive_claude, index_chunks', 'co 5-30 min'],
        ['Extraction', 'turbo_extract, extract_events, extract_entities, extract_commitments', 'co 30-60 min'],
        ['Intelligence', 'morning_brief, meeting_prep, response_drafter, intelligence_scan, weekly_synthesis, weekly_analysis', 'daily/weekly/15min 8-20'],
        ['Automation', 'auto_actions, task_monitor, calendar_check', 'co 5-30 min 8-20'],
        ['Maintenance', 'backup_db (3:00 + co 4h), prune_backups, QC, conv_cleanup, code_fix', 'daily + co 6h'],
        ['Observability', 'session_context, obs_alerts', 'co 30 min'],
    ]
)

# ============================================================
# CONVERSATION MEMORY
# ============================================================
add_heading('Conversation Memory (Sliding Window)', 1)

doc.add_paragraph(
    'Zaimplementowane 28.03.2026. Kazdy kanal (WhatsApp, Teams, voice, API) ma '
    'wlasne okno rozmowy (max 20 wiadomosci, max 8000 znakow). '
    'Historia wstrzykiwana do promptu przed materialem zrodlowym. '
    'Conversation-aware follow-ups: interpreter celuje w luki z poprzedniej odpowiedzi.'
)

add_table(
    ['Kanal', 'session_id', 'Przyklad'],
    [
        ['WhatsApp', '+48505441635', 'whatsapp:+48505441635'],
        ['Teams', 'conversation_id', 'teams:conv-abc-123'],
        ['Voice REST', 'session_id (Form)', 'voice:anonymous'],
        ['API', 'brak (stateless)', 'api:anonymous'],
    ]
)

# ============================================================
# CO DALEJ
# ============================================================
add_heading('Co dalej — pozostale zadania', 1)

add_heading('Faza A: Deploy + Discovery (BLOKOWANA przez Sebastiana)', 2)
add_table(
    ['Task', 'Status', 'Blocker'],
    [
        ['A1: Migracja Hetzner VPS + aktywacja cronow', 'TODO', 'Sebastian (Hetzner account + secrets)'],
        ['A2: Omnius REH deploy', 'TODO', 'Roch + IT Admin (Azure AD)'],
        ['A3: Omnius REF deploy na Azure', 'TODO', 'Krystian + IT Admin (Azure AD)'],
        ['A4: Onboarding Roch + Krystian', 'TODO', 'A2 + A3'],
        ['A5: NDA + compliance (Milosz Awedyk)', 'TODO', 'Sebastian'],
    ]
)

add_heading('Faza B: Hardening + Quality (czesciowo zrobiona)', 2)
add_table(
    ['Task', 'Status', 'Opis'],
    [
        ['B1: Dedup pipeline', 'DONE', 'Entity dedup, chunk dedup, event dedup'],
        ['B2: Ingestion Health Monitor', 'DONE', 'WhatsApp, email, calendar, Teams monitoring'],
        ['B3: Cost budget guards', 'DONE', 'Per-module budgets + WhatsApp alerts'],
        ['B4: QC agents', 'DONE', 'Auto code review + fixer loops'],
        ['B5: Observability dashboard', 'DONE', '/observability/ endpoints + alerts'],
    ]
)

add_heading('Faza C: Voice Interface', 2)
add_table(
    ['Task', 'Status', 'Opis'],
    [
        ['C1: Voice REST (STT->ask->TTS)', 'DONE', '/voice/ask endpoint, Whisper + edge-tts'],
        ['C2: Voice WebSocket (real-time)', 'DONE', 'voice_ws.py, sliding window 20 msg'],
        ['C3: Speaker ID', 'TODO', 'Rozpoznawanie glosu per osoba w nagraniach Plaud'],
        ['C4: Meeting boundaries', 'TODO', 'Auto-detekcja granic spotkan w nagraniach'],
    ]
)

add_heading('Faza D: Organizacja = Omnius', 2)
add_table(
    ['Task', 'Status', 'Opis'],
    [
        ['D1: Infrastructure audit per firma', 'TODO', 'SharePoint structure, Teams topology'],
        ['D2: Data classification rules', 'DONE', 'public/internal/confidential/ceo_only'],
        ['D3: Director augmentation', 'TODO', 'Per-director Omnius profile + auto-brief'],
        ['D4: Adoption monitoring', 'TODO', 'Usage metrics, engagement per role'],
    ]
)

add_heading('Faza E: Scale + Protect', 2)
add_table(
    ['Task', 'Status', 'Opis'],
    [
        ['E1: Wlasny LLM (self-hosted)', 'TODO', 'Llama/Mistral fine-tuned na danych Sebastiana'],
        ['E2: IP Box 5% CIT', 'IN PROGRESS', 'Dokumentacja R&D, ewidencja czasu, wniosek'],
        ['E3: Komercyjny Omnius SaaS', 'TODO', 'Multi-tenant, billing, onboarding flow'],
    ]
)

add_heading('Zmiana 6: MCP Tool Poka-Yoke (ciagla)', 2)
doc.add_paragraph(
    'Upgrade 43 narzedzi MCP: examples w inputSchema, negatywne przyklady '
    '("NIE uzywaj do..."), walidacja inputow z helpfulnymi error messages. '
    'Zalecanane przez Anthropic whitepaper — tool design = HCI design.'
)

# ============================================================
# ARCHITECTURE DIAGRAM
# ============================================================
add_heading('Diagram architektury', 1)

arch_png = '/tmp/gilbertus_architecture_v10.png'
generate_architecture_diagram(arch_png)
doc.add_picture(arch_png, width=Inches(7.5))
last_paragraph = doc.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('')
doc.add_paragraph('Diagram: Gilbertus (5 warstw pipeline + agentic patterns) + Omnius (per-company AI, RBAC, governance)')

# ============================================================
# COMMITY SESJI 28.03
# ============================================================
add_heading('Commity sesji 28.03.2026', 1)

add_table(
    ['Commit', 'Opis', 'Pliki', 'Linie'],
    [
        ['1d1c456', 'LangGraph StateGraph for Action Pipeline + HITL', '-', '-'],
        ['f407199', 'Observability + cost budgets + code review fixes', '-', '-'],
        ['9910d0e', 'Sliding Window Memory: multi-turn conversation', '7', '+284'],
        ['fd1e4b1', 'Agentic Phase 1: Answer Evaluator + Conv-Aware', '4', '+180'],
        ['00d1fa4', 'Agentic Phase 2: Deep Routing + Parallel Retrieval', '4', '+180'],
        ['d02d853', 'Agentic Phase 3: Orchestrator-Workers', '3', '+208'],
    ]
)

# ============================================================
# KOSZTY
# ============================================================
add_heading('Koszty API (daily budget: $20)', 1)
doc.add_paragraph(
    'Anthropic: Claude Sonnet (answering), Claude Haiku (interpreter, evaluator, extraction)\n'
    'OpenAI: text-embedding-3-large (embeddings)\n'
    'Budget guards: per-module limits + hard daily cap\n'
    'Evaluator: +~$0.003/request average (Haiku eval + 15% retry Sonnet)\n'
    'Parallel retrieval: +$0.0001/request (extra embedding + Qdrant)\n'
    'Orchestrator: 2-4x za decomposed queries (~10-15% traffic)'
)

# ============================================================
# SAVE
# ============================================================
output = '/mnt/c/Users/jablo/Desktop/Gilbertus_Masterplan_V10.docx'
doc.save(output)
print(f'Saved to {output}')
print(f'Architecture diagram: {arch_png}')
