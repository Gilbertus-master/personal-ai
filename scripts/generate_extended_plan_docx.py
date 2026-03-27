#!/usr/bin/env python3
"""Generate extended Gilbertus+Omnius plan DOCX with Security, QC, Compliance, Legal, Finance, Resources."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


def shading(cell, color):
    cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>'))

def h(doc, text, level=1):
    x = doc.add_heading(text, level=level)
    for r in x.runs: r.font.color.rgb = RGBColor(0x1A,0x1A,0x2E)

def lbl(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x2E,0x5C,0x8A)
    p.space_after = Pt(4)

def b(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    for r in p.runs: r.font.size = Pt(10)

def bd(doc, text):
    p = doc.add_paragraph(text)
    for r in p.runs: r.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(4)

def tbl(doc, headers, rows, hdr_color="2E5C8A"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"; table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, hd in enumerate(headers):
        c = table.rows[0].cells[i]; c.text = hd; shading(c, hdr_color)
        for p in c.paragraphs:
            for r in p.runs: r.bold=True; r.font.size=Pt(9); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
    for rd in rows:
        row = table.add_row()
        for i, v in enumerate(rd):
            row.cells[i].text = v
            for p in row.cells[i].paragraphs:
                for r in p.runs: r.font.size = Pt(9)
    doc.add_paragraph()

def title_page(doc, t1, t2):
    doc.add_paragraph(); doc.add_paragraph()
    x = doc.add_heading(t1, level=0); x.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in x.runs: r.font.color.rgb = RGBColor(0x1A,0x1A,0x2E)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(t2); r.font.size = Pt(13); r.font.color.rgb = RGBColor(0x55,0x55,0x55)
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Data: 26.03.2026  |  Przygotowa\u0142: Sebastian Jab\u0142o\u0144ski")
    r2.font.size = Pt(10); r2.font.color.rgb = RGBColor(0x88,0x88,0x88)
    doc.add_page_break()

def main():
    doc = Document()
    for s in doc.sections:
        s.top_margin=Cm(2); s.bottom_margin=Cm(2); s.left_margin=Cm(2.5); s.right_margin=Cm(2.5)

    title_page(doc, "GILBERTUS ALBANS + OMNIUS", "Plan Rozwoju Q2\u2013Q3 2026\nSecurity | Quality | Compliance | Legal | Finance | Resources")

    # === SPIS TRESCI ===
    h(doc, "Spis tre\u015bci", 1)
    bd(doc, "1. Kontekst i cel\n2. Diagnoza (8 problem\u00f3w)\n3. Plan \u2014 6 faz w 6 miesi\u0119cy\n4. Security\n5. Quality Control (agenci CI/CD)\n6. Compliance & RODO\n7. Legal\n8. Finance (koszty, bud\u017cet, ROI)\n9. Zasoby ludzkie \u2014 REH\n10. Zasoby ludzkie \u2014 REF\n11. Zasoby \u2014 Sebastian\n12. Ryzyka\n13. Weryfikacja")
    doc.add_page_break()

    # === 1. KONTEKST ===
    h(doc, "1. Kontekst i cel", 1)
    bd(doc, "Gilbertus Albans to prywatny mentat AI Sebastiana Jab\u0142o\u0144skiego. V1 operacyjne: 123k chunk\u00f3w, 35k event\u00f3w, 16k encji, 15 API endpoint\u00f3w, 17+ cron job\u00f3w. Delivery: WhatsApp, Teams, HTTP.")
    bd(doc, "Cel docelowy: Gilbertus pod nadzorem Sebastiana zarz\u0105dza sp\u00f3\u0142kami (REH, REF), optymalizuje procesy, ocenia ludzi, wy\u0142apuje nieefektywno\u015bci i wspiera decyzje tradingowe.")
    bd(doc, "Omnius to agent firmowy \u2014 \u017cyje w infrastrukturze sp\u00f3\u0142ki. Ka\u017cda sp\u00f3\u0142ka ma swojego Omniusa. Raportuj\u0105 do Gilbertusa przez encrypted API (summaries, nie raw data).")

    # === 2. DIAGNOZA ===
    h(doc, "2. Diagnoza \u2014 8 problem\u00f3w do naprawy", 1)
    tbl(doc, ["#", "Problem", "Impact", "Fix"], [
        ("1", "Brak connection poolingu", "Skalowalno\u015b\u0107", "psycopg_pool.ConnectionPool"),
        ("2", "Teams: 17k osobnych dokument\u00f3w", "Fragmentacja kontekstu", "Grupowanie per chat per 2h"),
        ("3", "Entities odci\u0119te od People", "Brief bez relacji", "Bridge entity_id"),
        ("4", "Entity dedup surface-level", "Zaszumiony graph", "pg_trgm + fuzzy"),
        ("5", "Event taxonomy za w\u0105ska", "Brak deadline/escalation", "+7 typ\u00f3w"),
        ("6", "Alerty = 0 (martwy system)", "Brak proaktywno\u015bci", "Aktywacja po >80%"),
        ("7", "Brak pipeline\u2019u ocen", "R\u0119czne oceny", "app/evaluation/"),
        ("8", "Produkcja na laptopie", "Brak GPU/HTTPS", "Hetzner AX102"),
    ])

    # === 3. PLAN 6 FAZ ===
    h(doc, "3. Plan \u2014 6 faz w 6 miesi\u0119cy", 1)
    tbl(doc, ["Faza", "Okres", "Cel", "Effort"], [
        ("0: Stabilizacja", "Tyg 1\u20132 (kwiecie\u0144)", "Naprawienie fundament\u00f3w", "6 dni"),
        ("1: Proaktywna inteligencja", "Tyg 3\u20136 (kw\u2013maj)", "Gilbertus m\u00f3wi co wa\u017cne", "12 dni"),
        ("2: Serwer + Oceny", "Tyg 7\u201310 (maj\u2013cze)", "Hetzner + pipeline ocen", "10 dni"),
        ("3: Omnius Lite REH", "Tyg 11\u201316 (cze\u2013lip)", "Roch ma AI asystenta", "12 dni"),
        ("4: Pe\u0142ne oceny + REF", "Tyg 17\u201320 (sie)", "Multi-perspective + Krystian", "11 dni"),
        ("5: Autonomiczne akcje", "Tyg 21\u201326 (wrz)", "Proactive management", "12 dni"),
    ])

    for phase, title_text, bullets in [
        ("0", "Faza 0: STABILIZACJA (kwiecie\u0144, 2 tygodnie)", [
            "Connection pooling \u2014 psycopg_pool singleton",
            "Teams message grouping \u2014 per chat per 2h",
            "Entity dedup \u2014 jednorazowy + pg_trgm",
            "Bridge people\u2194entities \u2014 entity_id w people",
            "Event taxonomy +7 typ\u00f3w: deadline, commitment, escalation, blocker, task_assignment, approval, rejection",
            "Aktywacja alert\u00f3w jako cron",
        ]),
        ("1", "Faza 1: PROAKTYWNA INTELIGENCJA (kwiecie\u0144\u2013maj)", [
            "Sync kalendarza z Microsoft Graph API",
            "Morning brief z kontekstem kalendarza + relacji + open loops",
            "Person-aware retrieval (auto-expand alias\u00f3w)",
            "Decision journal via WhatsApp (\u201egtd:\u201d keyword)",
            "Cross-domain correlation MVP",
        ]),
        ("2", "Faza 2: SERWER + PIPELINE OCEN (maj\u2013czerwiec)", [
            "Migracja Hetzner AX102 (128GB RAM, GPU-ready)",
            "HTTPS + monitoring",
            "Pipeline ocen: POST /evaluate \u2192 data_collector \u2192 evaluator (Claude) \u2192 DOCX",
            "Ka\u017cda ocena z explicit confidence score",
            "Kwartalny cron auto-ocen",
        ]),
        ("3", "Faza 3: OMNIUS LITE \u2014 REH (czerwiec\u2013lipiec)", [
            "Omnius REH: osobny Docker Compose, Postgres, Qdrant",
            "Graph API: SharePoint + shared Teams (NIE prywatne czaty)",
            "Interfejs Rocha: Teams Bot lub web chat",
            "Gilbertus\u2194Omnius read-only sync (nightly summaries)",
            "Compliance: Omnius V1 przetwarza TYLKO dane firmowe",
            "R\u00f3wnolegle: compliance review z prawnikiem",
        ]),
        ("4", "Faza 4: PE\u0141NE OCENY + OMNIUS REF (sierpie\u0144)", [
            "Enhanced evaluations z danymi Omnius",
            "Omnius REF \u2014 klon REH dla Krystiana",
            "Anomaly detection per osoba",
            "Employee scorecard dashboard",
            "Decision journal outcomes tracking (7/30/90 dni)",
        ]),
        ("5", "Faza 5: AUTONOMICZNE AKCJE (wrzesie\u0144)", [
            "Action items: Gilbertus proponuje \u2192 Sebastian zatwierdza \u2192 wykonanie",
            "Weekly reports: per-company + cross-company + personal",
            "Kwartalny evaluation cycle (auto)",
            "Process inefficiency detector",
            "Feedback loop: thumbs up/down",
        ]),
    ]:
        h(doc, title_text, 2)
        for bullet_text in bullets:
            b(doc, bullet_text)

    doc.add_page_break()

    # === 4. SECURITY ===
    h(doc, "4. Security", 1)

    h(doc, "Faza 0\u20131: Fundamenty", 2)
    b(doc, "Secrets management \u2014 SOPS/Vault zamiast plaintext .env")
    b(doc, "Audit log tabela \u2014 who, what, when, from_ip na ka\u017cdy endpoint")
    b(doc, "API key auth na wszystkich endpointach")

    h(doc, "Faza 2: Hardening serwera", 2)
    b(doc, "Firewall UFW: 443 (HTTPS), 22 (SSH key-only)")
    b(doc, "HTTPS everywhere: Let\u2019s Encrypt + nginx")
    b(doc, "SSH hardening: disable password, ed25519 keys, fail2ban")
    b(doc, "LUKS disk encryption na data volumes")
    b(doc, "GPG-encrypted backups off-site")
    b(doc, "Docker: non-root, read-only FS, resource limits")

    h(doc, "Faza 3: Multi-tenant security", 2)
    b(doc, "mTLS mi\u0119dzy Gilbertus\u2194Omnius")
    b(doc, "JWT auth z role claims (god/ceo/manager/employee)")
    b(doc, "RBAC per tenant \u2014 izolacja danych")
    b(doc, "Rate limiting per user")

    h(doc, "Faza 5: Action security", 2)
    b(doc, "Action approval: TOTP/biometric (docelowo, nie tylko WhatsApp \u201etak\u201d)")
    b(doc, "Principle of least privilege na Graph API scopes")
    b(doc, "Penetration test przed autonomicznymi akcjami")

    doc.add_page_break()

    # === 5. QUALITY CONTROL ===
    h(doc, "5. Quality Control \u2014 Agenci w p\u0119tlach", 1)

    h(doc, "Agent Kontroler Jako\u015bci Kodu (daily loop)", 2)
    bd(doc, "Trigger: co 24h (cron 6:00) + pre-commit hook.")
    b(doc, "Static analysis: ruff (linter+formatter), mypy (types), bandit (security)")
    b(doc, "Convention enforcement: connection pool, parameterized SQL, structured logging")
    b(doc, "Dead code detection: vulture")
    b(doc, "Test coverage: pytest-cov, alert <80%")
    b(doc, "Dependency audit: pip-audit na CVE")
    b(doc, "Output: JSON raport \u2192 insights (code_quality). Krytyczne \u2192 morning brief.")

    h(doc, "Agent Analizuj\u0105cy Architektur\u0119 (weekly loop)", 2)
    bd(doc, "Trigger: co niedziel\u0119 22:00 (przed poniedzia\u0142kowym briefem).")
    b(doc, "Module coupling: analiza import\u00f3w, granice modu\u0142\u00f3w")
    b(doc, "DB schema drift: por\u00f3wnanie aktualnego schematu z dokumentacj\u0105")
    b(doc, "Performance regression: avg latency /ask vs poprzedni tydzie\u0144")
    b(doc, "Tech debt tracker: TODO/FIXME/HACK trendline")
    b(doc, "Plan adherence: aktualny stan vs harmonogram faz")
    b(doc, "Output: tygodniowy raport \u2192 insights (architecture_review). Wchodzi do weekly reportu.")

    doc.add_page_break()

    # === 6. COMPLIANCE ===
    h(doc, "6. Compliance & RODO", 1)

    h(doc, "Podstawa prawna per warstwa", 2)
    tbl(doc, ["Warstwa", "Artyku\u0142 RODO", "Uwagi"], [
        ("Gilbertus (dane Sebastiana)", "Art. 6.1(a) zgoda", "Sebastian przetwarza w\u0142asne dane"),
        ("Gilbertus (firmowe email/Teams SJ)", "Art. 6.1(f) uzasad. interes", "SJ jest cz\u0142onkiem zarz\u0105du"),
        ("Omnius V1 (shared docs/channels)", "Art. 6.1(f) uzasad. interes", "Organizacja informacji firmowej"),
        ("Omnius V2 (people analytics)", "Art. 6.1(f) + DPIA", "BLOCKER \u2014 wymaga DPIA przed wdro\u017ceniem"),
    ])

    h(doc, "DPIA (Data Protection Impact Assessment)", 2)
    bd(doc, "Wymagana przed Faz\u0105 4 (people analytics). Przygotowuje: Mi\u0142osz Awedyk + DPO + Sebastian. Czas: ~20h. Zawiera: opis przetwarzania, ocena konieczno\u015bci, ocena ryzyk, \u015brodki bezpiecze\u0144stwa.")

    h(doc, "Polityka informacyjna pracownik\u00f3w", 2)
    bd(doc, "Art. 13/14 RODO \u2014 pracownicy musz\u0105 wiedzie\u0107, \u017ce shared channels/docs s\u0105 indeksowane. Dla people analytics: rozszerzona informacja + prawo sprzeciwu (Art. 21).")

    h(doc, "Retencja danych", 2)
    tbl(doc, ["Typ danych", "Retencja", "Podstawa"], [
        ("Surowe dane (chunki, dokumenty)", "5 lat", "Wymogi handlowe (faktury, kontrakty)"),
        ("Oceny pracownik\u00f3w", "3 lata po zako\u0144czeniu wsp\u00f3\u0142pracy", "Kodeks pracy + uzasad. interes"),
        ("Logi audytowe", "2 lata", "Best practice security"),
        ("Dane osobiste (WhatsApp, audio)", "Nieokre\u015blona", "Dane Sebastiana o sobie (zgoda)"),
    ])

    doc.add_page_break()

    # === 7. LEGAL ===
    h(doc, "7. Legal", 1)

    tbl(doc, ["Dokument", "Kiedy", "Kto", "Effort", "Koszt est."], [
        ("DPA Gilbertus", "Faza 0", "Mi\u0142osz Awedyk", "8h", "3 200\u20134 800 PLN"),
        ("IP ownership (kod)", "Faza 0", "Mi\u0142osz", "4h", "1 600\u20132 400 PLN"),
        ("DPA Omnius REH", "Faza 3", "Mi\u0142osz", "8h", "3 200\u20134 800 PLN"),
        ("Regulamin Omnius", "Faza 3", "Mi\u0142osz", "6h", "2 400\u20133 600 PLN"),
        ("Incident response plan", "Faza 3", "Mi\u0142osz", "4h", "1 600\u20132 400 PLN"),
        ("Umowa Hetzner review", "Faza 2", "Mi\u0142osz", "2h", "800\u20131 200 PLN"),
        ("DPIA people analytics", "Przed Faz\u0105 4", "Mi\u0142osz+DPO+SJ", "20h", "8 000\u201312 000 PLN"),
        ("Polityka informacyjna", "Przed Faz\u0105 4", "Mi\u0142osz+HR", "12h", "4 800\u20137 200 PLN"),
        ("DPA Omnius REF", "Faza 4", "Mi\u0142osz", "4h", "1 600\u20132 400 PLN"),
        ("Audit licencji open source", "Faza 0", "Mi\u0142osz", "4h", "1 600\u20132 400 PLN"),
    ])
    bd(doc, "\u0141\u0105czny effort prawnika: ~70h (~9 dni roboczych). Szacunkowy koszt: 28 000 \u2013 42 000 PLN.")

    doc.add_page_break()

    # === 8. FINANCE ===
    h(doc, "8. Finance", 1)

    h(doc, "Koszty bie\u017c\u0105ce (miesi\u0119czne)", 2)
    tbl(doc, ["Pozycja", "Koszt/mc", "Uwagi"], [
        ("Anthropic API (Claude)", "2 000\u20134 000 PLN", "Sonnet odpowiedzi + Haiku ekstrakcja"),
        ("OpenAI API (embeddings)", "200\u2013400 PLN", "text-embedding-3-large, przyrostowo"),
        ("Hetzner AX102 (od Fazy 2)", "~800 PLN", "128GB RAM, Ryzen 9, 2x1TB NVMe"),
        ("Backup off-site", "~100 PLN", "S3-compatible (Backblaze B2)"),
        ("Domena + SSL", "~5 PLN", "Let\u2019s Encrypt free"),
        ("RAZEM", "3 150\u20135 350 PLN/mc", ""),
    ])

    h(doc, "Koszty jednorazowe", 2)
    tbl(doc, ["Pozycja", "Koszt", "Kiedy"], [
        ("Prawnik (pe\u0142en scope)", "28 000\u201342 000 PLN", "Roz\u0142o\u017cone 6 mc"),
        ("GPU (opcjonalnie, real-time audio)", "8 000\u201315 000 PLN", "Faza 4+"),
        ("Pentest zewn\u0119trzny (opcjonalnie)", "5 000\u201315 000 PLN", "Przed Faz\u0105 5"),
        ("RAZEM jednorazowe", "33 000\u201372 000 PLN", ""),
    ])

    h(doc, "Bud\u017cet miesi\u0119czny per faza", 2)
    tbl(doc, ["Faza", "API", "Infra", "Prawnik", "Razem/mc"], [
        ("0\u20131 (kw\u2013maj)", "2 500 PLN", "0 (laptop)", "5 000 PLN", "7 500 PLN"),
        ("2 (maj\u2013cze)", "3 000 PLN", "800 PLN", "5 000 PLN", "8 800 PLN"),
        ("3 (cze\u2013lip)", "3 500 PLN", "800 PLN", "7 000 PLN", "11 300 PLN"),
        ("4 (sie)", "4 000 PLN", "800 PLN", "8 000 PLN", "12 800 PLN"),
        ("5 (wrz)", "4 000 PLN", "800 PLN", "3 000 PLN", "7 800 PLN"),
    ])

    h(doc, "ROI estimate (12 miesi\u0119cy)", 2)
    tbl(doc, ["Kategoria", "Oszcz\u0119dno\u015b\u0107/rok"], [
        ("Czas SJ na szukanie informacji (260h \u00d7 500 PLN/h)", "130 000 PLN"),
        ("Automatyczne oceny (60h/rok oszcz\u0119dno\u015bci)", "30 000 PLN"),
        ("Prewencja ryzyk (IREC 71K EUR, pozycja 239K PLN)", "100 000+ PLN"),
        ("Omnius REH: czas Rocha (156h \u00d7 300 PLN/h)", "47 000 PLN"),
        ("Omnius REF: czas Krystiana", "47 000 PLN"),
        ("Szybsze onboardingi (2 sp\u00f3\u0142ki)", "40 000 PLN"),
        ("\u0141\u0104CZNE OSZCZ\u0118DNO\u015aCI", "~394 000 PLN/rok"),
        ("\u0141\u0104CZNE KOSZTY (12 mc)", "~110 000 PLN"),
        ("ROI", "~3.6x w pierwszym roku"),
    ])

    doc.add_page_break()

    # === 9. ZASOBY REH ===
    h(doc, "9. Zasoby ludzkie \u2014 REH", 1)

    tbl(doc, ["Osoba", "Rola", "Czas", "Kiedy", "Zadania"], [
        ("Roch Baranowski\n(CEO)", "Sponsor +\npierwszy user", "16h\n(2h/tyg \u00d7 8)", "Faza 3\u20134", "Admin consent Graph API,\nfeedback, testowanie, scope"),
        ("Diana Skotnicka\n(CFO)", "Bud\u017cet +\nrisk oversight", "4h", "Faza 2\u20133", "Zatwierdzenie bud\u017cetu,\nprzegl\u0105d koszt\u00f3w"),
        ("Monika Grudzie\u0144\n(HR)", "Compliance +\nevaluations", "12h", "Faza 3\u20134", "Polityka informacyjna,\npotwierdzenia, wsparcie ocen"),
        ("Pawe\u0142 Kowalik\n(Chief of Staff)", "Use case ID", "4h", "Faza 3", "SharePoint libraries,\nTeams channels, procesy"),
        ("IT Admin REH\n(do ID)", "Graph API +\ninfra", "8h", "Faza 3", "Azure AD consent,\nscope\u2019y, firewall, VPN"),
        ("A. Jaguszewska\n(Biuro Zarz\u0105du)", "Uchwa\u0142y", "4h", "Faza 3\u20134", "Uchwa\u0142a zarz\u0105du\ndt. Omnius"),
        ("Maja Kalinowska\n(Head of PM)", "Beta tester", "4h", "Faza 3", "Test pyta\u0144 o portfel,\nkontrakty, pricing"),
        ("Marcin Kulpa\n(Head of OTC)", "Beta tester", "4h", "Faza 3", "Test pyta\u0144 o trading,\nEFET, kontrahent\u00f3w"),
    ])
    bd(doc, "\u0141\u0105czny czas REH (bez prawnika): ~52h = ~7 dni roboczych roz\u0142o\u017conych na 6 miesi\u0119cy.")
    bd(doc, "Z prawnikiem (Mi\u0142osz Awedyk, 70h): ~122h = ~15 dni roboczych.")

    doc.add_page_break()

    # === 10. ZASOBY REF ===
    h(doc, "10. Zasoby ludzkie \u2014 REF", 1)

    tbl(doc, ["Osoba", "Rola", "Czas", "Kiedy", "Zadania"], [
        ("Krystian Juchacz\n(CEO)", "Sponsor +\npierwszy user", "12h\n(2h/tyg \u00d7 6)", "Faza 4\u20135", "Admin consent,\nfeedback, testowanie"),
        ("IT Admin REF\n(do ID)", "Graph API +\ninfra", "6h", "Faza 4", "Azure AD consent,\nscope\u2019y"),
        ("HR REF\n(do ID)", "Compliance", "6h", "Faza 4\u20135", "Polityka informacyjna,\npotwierdzenia"),
        ("Key manager REF\n(do ID)", "Beta tester", "4h", "Faza 4", "Test pyta\u0144 o gaz,\nlogistyk\u0119, LNG"),
    ])
    bd(doc, "\u0141\u0105czny czas REF: ~28h = ~3.5 dnia roboczego roz\u0142o\u017conych na 3 miesi\u0105ce.")

    # === 11. SEBASTIAN ===
    h(doc, "11. Zasoby \u2014 Sebastian Jab\u0142o\u0144ski", 1)

    tbl(doc, ["Aktywno\u015b\u0107", "Czas", "Kiedy"], [
        ("Przegl\u0105d i zatwierdzenie planu", "2h", "Faza 0"),
        ("Review prawny z Mi\u0142oszem", "4h \u0142\u0105cznie", "Faza 0, 3, 4"),
        ("Testowanie (brief/oceny/alerty)", "26h (1h/tyg)", "Ca\u0142y okres"),
        ("Onboarding Rocha", "4h", "Faza 3"),
        ("Onboarding Krystiana", "3h", "Faza 4"),
        ("Decyzje bud\u017cetowe", "2h", "Faza 0, 2"),
        ("Feedback na weekly reports", "6h (0.5h/tyg)", "Faza 4\u20135"),
        ("RAZEM", "~47h", "6 miesi\u0119cy"),
    ])

    h(doc, "Development (Claude Code + Sebastian)", 2)
    tbl(doc, ["Faza", "Effort", "Zakres"], [
        ("0: Stabilizacja+Pami\u0119\u0107+MCP", "12 dni", "Pool, grouping, dedup, CLAUDE.md, MCP, rollback"),
        ("1: Proaktywna inteligencja", "12 dni", "Kalendarz, brief, correlation"),
        ("2: Serwer + Oceny", "10 dni", "Hetzner, evaluation pipeline"),
        ("3: Omnius REH", "12 dni", "Docker, Graph API, Teams Bot"),
        ("4: Oceny + REF", "11 dni", "Enhanced evals, anomaly, REF"),
        ("5: Autonomiczne akcje", "12 dni", "Approval pipeline, reports"),
        ("Quality agents", "4 dni", "Code QC + architecture review"),
        ("Security hardening", "5 dni", "Roz\u0142o\u017cone na fazy"),
        ("Structured logging", "~1 dzie\u0144", "Roz\u0142o\u017cony na Fazy 0-1"),
        ("RAZEM", "~83 dni robocze", "~4 miesi\u0105ce full-time"),
    ])

    doc.add_page_break()

    # === NEW: META WARSTWY ===
    h(doc, "12. Pami\u0119\u0107 i ci\u0105g\u0142o\u015b\u0107 sesji", 1)
    bd(doc, "Problem: nowa sesja Claude Code nie zna kontekstu poprzedniej. Rozwi\u0105zanie: 4-poziomowy system pami\u0119ci.")
    b(doc, "CLAUDE.md \u2014 instrukcje projektowe \u0142adowane automatycznie w ka\u017cdej sesji")
    b(doc, "SESSION_CONTEXT.md \u2014 auto-generowany snapshot (DB stats, coverage, synce, git)")
    b(doc, "Session handoff \u2014 summary na koniec ka\u017cdej sesji (memory/session_*.md)")
    b(doc, "Lessons learned DB \u2014 tabela bug_patterns + prevention_rules, sprawdzana przez QC agenta")
    b(doc, "Prompt versioning \u2014 tabela prompt_versions, ka\u017cda zmiana prompta = nowy rekord")

    h(doc, "13. MCP (Model Context Protocol)", 1)
    bd(doc, "Problem: Claude Code nie ma dost\u0119pu do Gilbertus API. Musi u\u017cywa\u0107 docker exec.")
    b(doc, "Przepisanie MCP server na oficjalny mcp SDK \u2014 11 tool\u00f3w")
    b(doc, "Toole: gilbertus_ask, timeline, summary, brief, alerts, status, db_stats, decide, people, evaluate, lessons")
    b(doc, "Pod\u0142\u0105czenie do Claude Code (settings.json mcpServers)")
    b(doc, "Faza 3: Omnius MCP per tenant (omnius_reh_ask, omnius_reh_docs)")
    b(doc, "Faza 5: gilbertus_action tool (approval pipeline)")

    h(doc, "14. Eliminacja powtarzanych b\u0142\u0119d\u00f3w", 1)
    b(doc, "Lessons learned DB \u2014 tabela z regu\u0142ami prewencji (seed z bug\u00f3w ekstrakcji)")
    b(doc, "Pre-commit hooks \u2014 ruff + mypy + custom checks (pool, parameterized SQL)")
    b(doc, "Code review checklist w CLAUDE.md \u2014 8 punkt\u00f3w do sprawdzenia przed commitem")
    b(doc, "Response feedback \u2014 thumbs up/down, tygodniowy raport jako\u015bci")
    b(doc, "Extraction rollback \u2014 extraction_run_id, mo\u017cliwo\u015b\u0107 cofni\u0119cia z\u0142ego runu")

    h(doc, "15. Dodatkowe ulepszenia in\u017cynieryjne", 1)
    b(doc, "Data lineage \u2014 tabela lineage (input\u2192output tracing dla ocen i insight\u00f3w)")
    b(doc, "Structured logging \u2014 structlog zamiast print() (JSON output)")
    b(doc, "Graceful worker shutdown \u2014 SIGTERM handler, doko\u0144cz bie\u017c\u0105cy chunk")
    b(doc, "API cost tracking \u2014 tabela api_costs (provider, model, tokens, cost)")
    b(doc, "Timezone awareness \u2014 TZ=Europe/Warsaw w crontab, absolutne daty w memory")

    bd(doc, "Zaktualizowana Faza 0: 12 dni (by\u0142o 6). Inwestycja w fundament eliminuj\u0105cy utrat\u0119 kontekstu i powtarzanie b\u0142\u0119d\u00f3w.")

    doc.add_page_break()

    # === RYZYKA ===
    h(doc, "16. Ryzyka", 1)

    tbl(doc, ["Ryzyko", "P", "Impact", "Mitygacja"], [
        ("Compliance blokuje people analytics", "Wys.", "Wys.", "Faza 3 BEZ danych pracown."),
        ("Roch/Krystian nie adoptuj\u0105", "\u015ared.", "Wys.", "Pain point first, demo z danymi"),
        ("Koszty API > bud\u017cet", "\u015ared.", "\u015ared.", "Haiku, cap miesi\u0119czny"),
        ("Migracja = utrata danych", "Nisk.", "Kryt.", "Backup + test restore + fallback"),
        ("Brak IT admina w REH/REF", "\u015ared.", "Wys.", "Zidentyfikowa\u0107 kto zarz\u0105dza Azure AD"),
        ("Prawnik op\u00f3\u017ania", "\u015ared.", "\u015ared.", "Zaanga\u017cowa\u0107 wcze\u015bnie, osobne zlecenia"),
        ("UODO wymaga konsultacji (DPIA)", "Nisk.", "Wys.", "Bud\u017cet +2 mc na proces"),
        ("Security breach po migracji", "Nisk.", "Kryt.", "Hardening, pentest, incident plan"),
        ("Burnout SJ (owner=dev=tester)", "\u015ared.", "Wys.", "QC agents, docelowo dev Q4 2026"),
        ("Context loss mi\u0119dzy sesjami", "Wys.", "Wys.", "CLAUDE.md + SESSION_CONTEXT.md"),
        ("Powtarzanie b\u0142\u0119d\u00f3w ekstrakcji", "\u015ared.", "\u015ared.", "Lessons learned + pre-commit"),
        ("MCP nie pod\u0142\u0105czony", "Wys.", "\u015ared.", "Przepisanie + settings.json"),
    ])

    # === WERYFIKACJA ===
    h(doc, "17. Weryfikacja", 1)
    b(doc, "Testy end-to-end: WhatsApp \u2192 odpowied\u017a z nowymi capability")
    b(doc, "Metryki DB: coverage ekstrakcji, alert count, decision count, evaluation count")
    b(doc, "Sebastian test: Czy brief jest actionable? Czy oceny trafne? Czy Roch u\u017cywa?")
    b(doc, "Code quality score: tygodniowy raport agenta QC \u2014 zero critical")
    b(doc, "Architecture adherence: tygodniowy raport \u2014 odchylenia od planu")
    b(doc, "Compliance checklist: per-faza (DPA? Polityka? DPIA?)")
    b(doc, "Budget tracking: miesi\u0119czne koszty API+infra vs bud\u017cet")
    b(doc, "Feedback: thumbs up/down, target >80% positive")

    return doc


if __name__ == "__main__":
    base = "/mnt/c/Users/jablo/Desktop"
    path = f"{base}/Gilbertus_Omnius_Plan_Q2Q3_2026.docx"
    doc = main()
    doc.save(path)
    print(f"Extended plan: {path}")
    # Also regenerate REH and REF plans (unchanged)
    from generate_plans_docx import generate_reh_plan, generate_ref_plan
    generate_reh_plan(f"{base}/Omnius_REH_Plan_dla_Rocha.docx")
    generate_ref_plan(f"{base}/Omnius_REF_Plan_dla_Krystiana.docx")
