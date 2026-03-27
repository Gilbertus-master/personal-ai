#!/usr/bin/env python3
"""Generate Masterplan V3 DOCX with professional formatting."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def shading(cell, color):
    cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>'))

def h(doc, text, level=1):
    x = doc.add_heading(text, level=level)
    for r in x.runs: r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

def lbl(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True; r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x2E, 0x5C, 0x8A); p.space_after = Pt(4)

def b(doc, text, done=False):
    prefix = "\u2705 " if done else "\u2B1C "
    p = doc.add_paragraph(prefix + text, style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    for r in p.runs: r.font.size = Pt(10)
    if done:
        for r in p.runs: r.font.color.rgb = RGBColor(0x44, 0x88, 0x44)

def bd(doc, text):
    p = doc.add_paragraph(text)
    for r in p.runs: r.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(4)

def quote(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.italic = True; r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x44, 0x44, 0x88)
    p.paragraph_format.left_indent = Cm(1)

def tbl(doc, headers, rows, hdr_color="2E5C8A"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"; table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, hd in enumerate(headers):
        c = table.rows[0].cells[i]; c.text = hd; shading(c, hdr_color)
        for p in c.paragraphs:
            for r in p.runs: r.bold = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for rd in rows:
        row = table.add_row()
        for i, v in enumerate(rd):
            row.cells[i].text = v
            for p in row.cells[i].paragraphs:
                for r in p.runs: r.font.size = Pt(9)
            if "DONE" in v: shading(row.cells[i], "E8F5E9")
            elif "TODO" in v or "TBD" in v: shading(row.cells[i], "FFF3E0")
    doc.add_paragraph()

def title_page(doc, t1, t2, t3):
    doc.add_paragraph(); doc.add_paragraph()
    x = doc.add_heading(t1, level=0); x.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in x.runs: r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(t2); r.font.size = Pt(14); r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(t3); r2.font.size = Pt(10); r2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_page_break()

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2); s.bottom_margin = Cm(2); s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)

title_page(doc,
    "GILBERTUS ALBANS + OMNIUS",
    "Masterplan V3",
    "2026-03-26 23:15 CET | Sesja 14h | 9 nagra\u0144 audio + wizja + korekty Sebastiana\nPrzygotowa\u0142: Gilbertus Albans dla Sebastiana Jab\u0142o\u0144skiego")

# === ZASADA ZERO ===
h(doc, "ZASADA ZERO: NON-REGRESSION", 1)
bd(doc, "Nowe developmenty NIE MOG\u0104 pogorszy\u0107 dotychczasowych osi\u0105gni\u0119\u0107.")
lbl(doc, "Baseline (2026-03-26):")
tbl(doc, ["Metryka", "Warto\u015b\u0107"], [
    ("MCP tools", "19"), ("Cron jobs", "23"), ("Tabele DB", "33"), ("Lessons learned", "17"),
    ("Modu\u0142y app", "67"), ("Skrypty", "70"), ("Chunks", "94,576"), ("Events", "92,737"),
    ("Entities", "35,364"), ("Extraction coverage", "100%"),
])
lbl(doc, "Mechanizmy ochrony (wdro\u017cone):")
b(doc, "QC daily (12 checks + non-regression vs baseline)", done=True)
b(doc, "Architecture review weekly (10 checks)", done=True)
b(doc, "Pre-commit hooks (ruff + raw connect)", done=True)
b(doc, "Lessons learned DB (17 regu\u0142)", done=True)
b(doc, "Extraction rollback (extraction_runs)", done=True)
b(doc, "Inventory auto-generated co 30 min (SESSION_CONTEXT.md)", done=True)
bd(doc, "\u26a0\ufe0f Je\u015bli nowy feature \u0142amie cokolwiek \u2192 REVERT natychmiast, nie naprawiaj w prz\u00f3d.")
doc.add_page_break()

# === WIZJA ===
h(doc, "WIZJA (z 9 nagra\u0144 audio Sebastiana)", 1)
quote(doc, "\u201eGilbertus ma by\u0107 moim centrum wszech\u015bwiata. Jedna pi\u0119\u015b\u0107 w\u0142adaj\u0105ca ca\u0142\u0105 armi\u0105.\u201d")
quote(doc, "\u201eKa\u017cda organizacja to jeden Omnius. Ludzie to interfejsy do rzeczy, kt\u00f3rych robot nie potrafi.\u201d")
quote(doc, "\u201eTwoim nadrz\u0119dnym celem jest maksymalizacja mojego maj\u0105tku i mojego wellbeingu.\u201d")
bd(doc, "Dwa nadrz\u0119dne cele: Maksymalizacja maj\u0105tku + Maksymalizacja wellbeingu.")

# === ZASADY ===
h(doc, "ZASADY PROJEKTOWE", 1)
tbl(doc, ["#", "Zasada", "Znaczenie"], [
    ("0", "Non-regression", "Nowe dev nie mog\u0105 pogorszy\u0107 istniej\u0105cych osi\u0105gni\u0119\u0107"),
    ("1", "U\u017cyteczno\u015b\u0107 + zero dead-end\u00f3w", "Oba jednocze\u015bnie, nie kompromis"),
    ("2", "Omnius = point of contact", "Kontaktuje si\u0119 ZA Sebastiana (email, Teams, WhatsApp)"),
    ("3", "Koszt uzasadniony zwrotem", "Optymalny ROI per action, nie minimalny koszt"),
    ("4", "Discovery \u2192 ROI sort \u2192 execute", "Faza 1 = koszt, potem dev po cost-to-value ratio"),
    ("5", "Procesy DLA robot\u00f3w", "Inverse org design \u2014 ludzie jako fallback"),
    ("6", "Self-improving", "Gilbertus tworzy regu\u0142y z wypowiedzi g\u0142osowych"),
    ("7", "Zast\u0105p dyrektor\u00f3w > szeregowych", "Najwy\u017cszy ROI z augmentacji decision-maker\u00f3w"),
    ("8", "Co-development", "Pracownicy rozwijaj\u0105 Omniusa (nigdy Gilbertusa)"),
    ("9", "Komunikacja po autoryzacji", "Per message LUB standing order na zakres"),
    ("10", "Continuous intelligence", "Ka\u017cda nowa informacja \u2192 optymalizacja / ryzyko / okazja"),
])
doc.add_page_break()

# === STAN ===
h(doc, "STAN ZREALIZOWANY (sesja 26.03.2026)", 1)
tbl(doc, ["Faza", "Status", "Kluczowe deliverables"], [
    ("0: Stabilizacja", "DONE 19/19", "Pool, dedup, Teams grouping, taxonomy 15 typ\u00f3w, MCP 19 tool\u00f3w, CLAUDE.md, QC agents"),
    ("1: Proaktywna", "DONE 6/6", "Kalendarz, brief 5 sekcji, person-aware retrieval, decision journal, correlation"),
    ("2: Oceny", "3/5 DONE", "Eval pipeline, quarterly cron, answer cache. Pending: migracja+HTTPS"),
    ("3-5: Omnius+Akcje", "KOD DONE", "Omnius REH+REF, command protocol, action pipeline, reports, feedback, inefficiency"),
    ("QC Agents", "RUNNING", "Code quality daily + Architecture weekly + Inventory auto"),
    ("Plaud", "FIXED", "ori_ready bug, 56/58 imported, pipeline automatic"),
])
doc.add_page_break()

# === FAZA A ===
h(doc, "Faza A: DEPLOY + DISCOVERY (tydzien 1-3)", 1)
bd(doc, "Cel: Uruchomi\u0107 produkcj\u0119 + Omniusa. Zebra\u0107 dane. Discovery co optymalizowa\u0107.")
tbl(doc, ["#", "Task", "Effort", "Zale\u017cno\u015bci"], [
    ("A1", "Migracja Hetzner \u2014 .env, nginx, certbot, cron", "1 dzie\u0144", "Sebastian (secrets)"),
    ("A2", "Omnius REH deploy \u2014 Docker, Graph API, SharePoint", "3 dni", "Roch + IT Admin"),
    ("A3", "Omnius REF deploy \u2014 klon REH", "1 dzie\u0144", "Krystian + IT Admin"),
    ("A4", "Onboarding Roch + Krystian", "2 dni", "A2, A3"),
    ("A5", "NDA + compliance \u2014 Miłosz Awedyk", "1 dzie\u0144", "R\u00f3wnolegle"),
    ("A6", "Discovery: infra audit REH/REF", "ongoing", "A2, A3"),
    ("A7", "Discovery: cash flow analysis", "ongoing", "A2, A3"),
    ("A8", "Cost-to-value ranking top 10", "po A6+A7", ""),
])
bd(doc, "Metryki: Serwer live z HTTPS. Roch i Krystian maj\u0105 Omniusa. Discovery report gotowy.")
doc.add_page_break()

# === FAZA B ===
h(doc, "Faza B: KOMUNIKACJA + INTELLIGENCE + COMPLIANCE (tydzien 4-10)", 1)
bd(doc, "Cel: Gilbertus komunikuje w imieniu Sebastiana. Continuous intelligence. Compliance. Optymalizacje z discovery.")

h(doc, "B1: Komunikacja w imieniu Sebastiana", 2)
bd(doc, "Tryb 1 \u2014 Autoryzacja per message: Gilbertus draftuje \u2192 Sebastian zatwierdza na WhatsApp \u2192 Gilbertus wysy\u0142a.")
bd(doc, "Tryb 2 \u2014 Standing orders: Sebastian autoryzuje zakres (np. 'emaile do Rocha o statusach, max 3/dzie\u0144, nigdy o wynagrodzeniach') \u2192 Gilbertus wysy\u0142a w scope \u2192 daily digest raportuje.")
tbl(doc, ["#", "Task", "Effort"], [
    ("B1.1", "standing_orders tabela + WhatsApp: authorize/revoke/list", "1 dzie\u0144"),
    ("B1.2", "Auto-draft engine \u2014 wykrywanie potrzeby komunikacji", "2 dni"),
    ("B1.3", "Email send (Graph API Mail.Send)", "DONE"),
    ("B1.4", "Teams send (Graph API Chat.ReadWrite)", "1 dzie\u0144"),
    ("B1.5", "WhatsApp send", "DONE"),
    ("B1.6", "Daily digest \u2014 co Gilbertus wys\u0142a\u0142 (cron 20:00)", "0.5 dnia"),
    ("B1.7", "sent_communications tabela + audit trail", "0.5 dnia"),
    ("B1.8", "Scope checker \u2014 walidacja vs standing order", "1 dzie\u0144"),
])

h(doc, "B2: Continuous Intelligence (Opportunity Detector)", 2)
bd(doc, "Gilbertus analizuje KA\u017bD\u0104 now\u0105 informacj\u0119 i szuka warto\u015bci.")
b(doc, "Co 2h: scan nowych event\u00f3w/chunk\u00f3w")
b(doc, "Klasyfikuj: OPTIMIZATION / OPPORTUNITY / RISK / NEW_BUSINESS")
b(doc, "Estimate value (PLN) + effort (h) + confidence")
b(doc, "Rank po ROI = value / effort")
b(doc, "Top items: auto-draft action + WhatsApp notification")
tbl(doc, ["#", "Task", "Effort"], [
    ("B2.1", "opportunities tabela + Opportunity Detector cron co 2h", "2 dni"),
    ("B2.2", "Value estimator (PLN per opportunity)", "1 dzie\u0144"),
    ("B2.3", "Auto-draft actions per opportunity", "1 dzie\u0144"),
    ("B2.4", "WhatsApp notifications z top discoveries", "0.5 dnia"),
    ("B2.5", "gilbertus_opportunities MCP tool + API", "0.5 dnia"),
])

lbl(doc, "Typy discoveries:")
b(doc, "Cost reduction \u2014 proces kosztuje X, automatyzacja za Y")
b(doc, "Revenue opportunity \u2014 potencjalny klient/deal, draft follow-up")
b(doc, "Risk \u2014 eskalacja/konflikt/blocker, propozycja interwencji")
b(doc, "Process inefficiency \u2014 manualne powtarzalne zadania \u2192 automatyzacja")
b(doc, "New business \u2014 nowy rynek/produkt/partner, draft exploration plan")

h(doc, "B3: Compliance Manager AI", 2)
bd(doc, "Z nagra\u0144: 'Analizuje styl komunikacji, kultur\u0119 osobist\u0105, kompas moralny. Generuje raporty flaguj\u0105ce ludzi. Trackuje czy si\u0119 poprawiaj\u0105.'")
tbl(doc, ["#", "Task", "Effort"], [
    ("B3.1", "Communication style analyzer (per person per week)", "2 dni"),
    ("B3.2", "Culture fit scoring (values, professionalism, responsiveness)", "1 dzie\u0144"),
    ("B3.3", "Improvement tracking \u2014 baseline \u2192 delta po feedbacku", "1 dzie\u0144"),
    ("B3.4", "Red flag detector (culture misfit, toxicity)", "0.5 dnia"),
    ("B3.5", "Monthly compliance report per sp\u00f3\u0142ka", "0.5 dnia"),
])

h(doc, "B4: Pierwsze optymalizacje z discovery", 2)
tbl(doc, ["#", "Task", "Effort"], [
    ("B4.1", "Top 3 optymalizacje per ROI ranking z Fazy A", "TBD po discovery"),
    ("B4.2", "Cash flow optimizer", "TBD"),
    ("B4.3", "Controlling automation", "TBD"),
])
doc.add_page_break()

# === FAZA C ===
h(doc, "Faza C: VOICE + SELF-IMPROVING (tydzien 11-16)", 1)
bd(doc, "Cel: Dialog g\u0142osowy z Gilbertusem + system samodoskonalaj\u0105cy si\u0119.")
quote(doc, "\u201eMusimy dzisiaj zrobi\u0107 interfejs g\u0142osowy.\u201d \u201eZ ca\u0142ych moich wypowiedzi musisz tworzy\u0107 zasady dla siebie.\u201d")
tbl(doc, ["#", "Task", "Effort"], [
    ("C1", "Voice interface \u2014 streaming STT + TTS + WebSocket", "3 dni"),
    ("C2", "Self-improving rules engine \u2014 g\u0142os \u2192 regu\u0142y \u2192 DB \u2192 stosuj", "2 dni"),
    ("C3", "Speaker identification \u2014 voiceprint matching w Plaud", "2 dni"),
    ("C4", "Meeting boundary detection \u2014 15-min chunking", "1 dzie\u0144"),
    ("C5", "Conceptual work assistant \u2014 voice-driven", "2 dni"),
    ("C6", "Multi-interface \u2014 Claude Code + voice + WhatsApp", "1 dzie\u0144"),
])
doc.add_page_break()

# === FAZA D ===
h(doc, "Faza D: ORGANIZACJA = OMNIUS (tydzien 17-24)", 1)
bd(doc, "Cel: Ka\u017cda organizacja to jeden Omnius. Zast\u0105p zewn\u0119trzne systemy.")
quote(doc, "\u201eKa\u017cda organizacja to by\u0142 jeden Omnius i \u017ceby nie by\u0142o \u017cadnego dodatkowego systemu, kt\u00f3ry od kogo\u015b kupujemy.\u201d")
tbl(doc, ["#", "Task", "Effort"], [
    ("D1", "Infrastructure audit \u2014 WSZYSTKIE systemy/appy", "1 dzie\u0144"),
    ("D2", "Replacement plan \u2014 sorted po savings/effort", "1 dzie\u0144"),
    ("D3", "Controlling + analytics \u2014 pierwszy replacement", "TBD"),
    ("D4", "Co-development framework \u2014 scope'y per rola", "2 dni"),
    ("D5", "Director-level augmentation \u2014 rekomendacje decyzyjne", "3 dni"),
    ("D6", "Omnius jako pe\u0142ny PoC \u2014 prowadzi konwersacje", "3 dni"),
    ("D7", "Client-facing energy apps (USP)", "TBD"),
    ("D8", "Trading strategy validation + invention", "2 dni"),
    ("D9", "Adoption monitoring", "1 dzie\u0144"),
])
doc.add_page_break()

# === FAZA E ===
h(doc, "Faza E: SCALE + PROTECT + GROW (tydzien 25-32)", 1)
bd(doc, "Cel: Skalowanie, ochrona IP, komercjalizacja, nowy biznes.")
quote(doc, "\u201eNajwa\u017cniejszy projekt do tej pory w moim \u017cyciu.\u201d \u201eTwoje sp\u00f3\u0142ki zarz\u0105dzaj\u0105 si\u0119 same... digitalowego ciebie. Dla ultra bogatych ludzi.\u201d")
tbl(doc, ["#", "Task", "Effort"], [
    ("E1", "W\u0142asny LLM (hybrydowy) \u2014 REH cloud", "TBD"),
    ("E2", "IP Box 5% \u2014 optymalizacja podatkowa", "1 dzie\u0144 (prawnik)"),
    ("E3", "NDA + non-compete \u2014 ochrona IP", "done w planie legal"),
    ("E4", "Komercyjny Omnius \u2014 ultra-bogaci target", "TBD"),
    ("E5", "Inverse org design \u2014 procesy DLA robot\u00f3w", "ongoing"),
    ("E6", "Acquisition leverage \u2014 przejmowanie sp\u00f3\u0142ek", "TBD"),
    ("E7", "Wellbeing monitor \u2014 tracking koszt\u00f3w emocjonalnych", "1 dzie\u0144"),
    ("E8", "Parallel development \u2014 ludzie w REH/REF vibe-coduj\u0105", "ongoing"),
])
doc.add_page_break()

# === KOMUNIKACJA ===
h(doc, "WARSTWA: Komunikacja w imieniu Sebastiana", 1)
h(doc, "Tryb 1: Autoryzacja per message", 2)
bd(doc, "\U0001f514 Propozycja #47: email do Rocha\nTemat: Follow-up NOFAR\n[draft tre\u015bci]\n\ntak #47 = wy\u015blij | nie #47 = odrzu\u0107 | edit #47: [zmiany]")
h(doc, "Tryb 2: Standing orders", 2)
bd(doc, "authorize: email roch@* statusy projekt\u00f3w i follow-upy, max 3/dzie\u0144, nigdy wynagrodzenia/zwolnienia/strategia")
h(doc, "Daily digest (cron 20:00)", 2)
bd(doc, "\U0001f4ca Dzi\u015b w Twoim imieniu:\n- 2 emaile do Rocha (follow-up NOFAR, status BESS)\n- 1 Teams do Maji (korekta cennikowa)\n\nStanding orders: 3 aktywne | Pending approval: 1")
bd(doc, "Audit: Ka\u017cda wiadomo\u015b\u0107 \u2192 sent_communications + audit_log. Traceability do action_items i opportunities.")
doc.add_page_break()

# === INTELLIGENCE ===
h(doc, "WARSTWA: Continuous Intelligence", 1)
bd(doc, "Opportunity Detector Pipeline (co 2h):")
b(doc, "Scan nowych event\u00f3w/chunk\u00f3w z ostatnich 2h")
b(doc, "Haiku klasyfikuje: OPTIMIZATION / OPPORTUNITY / RISK / NEW_BUSINESS")
b(doc, "Value estimate (PLN) + effort (h) + confidence")
b(doc, "ROI = value / effort, rank desc")
b(doc, "Top 5 \u2192 auto-draft action \u2192 WhatsApp notification")
bd(doc, "Tabela: opportunities | MCP: gilbertus_opportunities | Cron: co 2h")

# === COMPLIANCE ===
h(doc, "WARSTWA: Compliance Manager AI", 1)
bd(doc, "Weekly per person pipeline:")
b(doc, "Styl komunikacji (profesjonalizm, jasno\u015b\u0107, reakcja na feedback)")
b(doc, "Culture fit (warto\u015bci, zaanga\u017cowanie, wsp\u00f3\u0142praca)")
b(doc, "Responsiveness (czas odpowiedzi, follow-up na commitments)")
b(doc, "Red flags (toxicity, brak poprawy, pattern eskalacji)")
b(doc, "Score per wymiar (1-5) \u2192 por\u00f3wnanie z baseline \u2192 monthly report")
doc.add_page_break()

# === ZASOBY ===
h(doc, "ZASOBY LUDZKIE", 1)
tbl(doc, ["Osoba", "Rola", "Czas", "Kiedy"], [
    ("Roch Baranowski", "Sponsor + user Omnius REH", "16h", "Faza A-B"),
    ("Krystian Juchacz", "Sponsor + user Omnius REF", "12h", "Faza A-B"),
    ("IT Admin REH/REF", "Azure AD, Graph API", "14h", "Faza A"),
    ("Mi\u0142osz Awedyk", "DPA, NDA, DPIA, regulaminy", "~89h", "Ca\u0142y okres"),
    ("Diana Skotnicka", "Bud\u017cet, risk", "4h", "Faza A"),
    ("Monika Grudzie\u0144", "Compliance, polityka inform.", "12h", "Faza B-D"),
    ("Sebastian", "Decyzje, onboarding, testing", "~47h", "Ca\u0142y okres"),
])

# === FINANCE ===
h(doc, "FINANCE", 1)
tbl(doc, ["Pozycja", "Koszt"], [
    ("Bie\u017c\u0105ce (miesi\u0119czne)", "~3-5K PLN/mc"),
    ("Jednorazowe (prawnik + opcje)", "~35-55K PLN"),
    ("ROI (pierwszy rok)", "~3.6x (394K oszcz\u0119dno\u015bci vs 110K koszt\u00f3w)"),
])

# === WERYFIKACJA ===
h(doc, "WERYFIKACJA", 1)
lbl(doc, "Per deploy:")
b(doc, "Zasada Zero checklist (10 punkt\u00f3w)")
b(doc, "QC agent 0 errors")
b(doc, "Non-regression baseline nie spad\u0142a")
lbl(doc, "Per faza:")
b(doc, "Metryki fazy osi\u0105gni\u0119te")
b(doc, "Sebastian test \u2014 czy dzia\u0142a w praktyce?")
b(doc, "Feedback rating >80% positive")
lbl(doc, "Ongoing:")
b(doc, "Code quality daily 6:00 (12 checks)")
b(doc, "Architecture review weekly Sun 22:00 (10 checks)")
b(doc, "Inventory auto-updated co 30 min")
b(doc, "Lessons learned growing")
b(doc, "API costs tracked per module")
b(doc, "Opportunities detected and acted upon")

doc.save("/mnt/c/Users/jablo/Desktop/Gilbertus_Masterplan_V3.docx")
print("Saved!")
