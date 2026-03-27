#!/usr/bin/env python3
"""Generate 3 DOCX plans: full Gilbertus+Omnius plan, REH plan for Roch, REF plan for Krystian."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def styled_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h


def label(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x2E, 0x5C, 0x8A)
    p.space_after = Pt(4)


def bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        run.font.size = Pt(10)


def body(doc, text):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(4)


def add_table(doc, headers, rows, header_color="2E5C8A"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, header_color)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = val
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph()


def title_page(doc, title, subtitle, date="26.03.2026"):
    doc.add_paragraph()
    doc.add_paragraph()
    t = doc.add_heading(title, level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in t.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(f"Data: {date}  |  Przygotowa\u0142: Sebastian Jab\u0142o\u0144ski")
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_page_break()


def set_margins(doc):
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)


# =========================================================================
# FULL PLAN
# =========================================================================
def generate_full_plan(path):
    doc = Document()
    set_margins(doc)
    title_page(doc,
        "GILBERTUS ALBANS + OMNIUS",
        "Plan Rozwoju Q2\u2013Q3 2026  |  6 Faz  |  6 Miesi\u0119cy")

    styled_heading(doc, "1. Kontekst", 1)
    body(doc,
        "Gilbertus Albans to prywatny mentat AI Sebastiana Jab\u0142o\u0144skiego \u2014 system kognitywny, "
        "kt\u00f3ry indeksuje dane z 10+ \u017ar\u00f3de\u0142 (Teams, email, WhatsApp, audio Plaud, dokumenty), "
        "wyci\u0105ga encje i eventy, i dostarcza proaktywn\u0105 inteligencj\u0119 przez WhatsApp i Teams.")
    body(doc,
        "Omnius to agent firmowy \u2014 \u017cyje w infrastrukturze sp\u00f3\u0142ki, przetwarza dane korporacyjne "
        "(SharePoint, shared Teams, mailboxy, raporty). Ka\u017cda sp\u00f3\u0142ka ma swojego Omniusa. "
        "Omniusy raportuj\u0105 do Gilbertusa przez encrypted API.")

    label(doc, "Stan obecny (marzec 2026):")
    bullet(doc, "123 476 chunk\u00f3w, 35 543 event\u00f3w, 16 181 encji, 121 insight\u00f3w")
    bullet(doc, "15 API endpoint\u00f3w, 17+ cron job\u00f3w automatyzacji")
    bullet(doc, "Delivery: WhatsApp, Teams Bot, HTTP API")
    bullet(doc, "Infrastruktura: Docker (Postgres, Qdrant, Whisper) na WSL2")

    label(doc, "Cel docelowy:")
    body(doc,
        "Gilbertus pod nadzorem Sebastiana zarz\u0105dza sp\u00f3\u0142kami (REH, REF), optymalizuje procesy, "
        "ocenia ludzi, wy\u0142apuje nieefektywno\u015bci, wspiera decyzje tradingowe i chroni wellbeing.")

    # Diagnoza
    styled_heading(doc, "2. Diagnoza \u2014 8 problem\u00f3w do naprawy", 1)
    add_table(doc, ["#", "Problem", "Impact", "Fix"],
    [
        ("1", "Brak connection poolingu", "Skalowalno\u015b\u0107", "psycopg_pool.ConnectionPool"),
        ("2", "Teams: 17k osobnych dokument\u00f3w", "Fragmentacja kontekstu", "Grupowanie per chat per 2h"),
        ("3", "Entities odci\u0119te od People", "Morning brief bez relacji", "Bridge entity_id w people"),
        ("4", "Entity dedup surface-level", "Zaszumiony knowledge graph", "pg_trgm + fuzzy matching"),
        ("5", "Event taxonomy za w\u0105ska (8 typ\u00f3w)", "Brak deadline/escalation/blocker", "+7 nowych typ\u00f3w"),
        ("6", "Alerty = 0 (martwy system)", "Brak proaktywno\u015bci", "Aktywacja po ekstrakcji >80%"),
        ("7", "Brak pipeline\u2019u ocen", "R\u0119czne oceny pracownik\u00f3w", "Modu\u0142 app/evaluation/"),
        ("8", "Produkcja na laptopie WSL2", "Brak GPU, HTTPS, stabilno\u015bci", "Migracja Hetzner AX102"),
    ])

    # Fazy
    styled_heading(doc, "3. Plan \u2014 6 faz w 6 miesi\u0119cy", 1)

    add_table(doc, ["Faza", "Okres", "Cel", "Effort"],
    [
        ("0: Stabilizacja", "Tyg 1\u20132 (kwiecie\u0144)", "Naprawienie fundament\u00f3w", "6 dni"),
        ("1: Proaktywna inteligencja", "Tyg 3\u20136 (kwiecie\u0144\u2013maj)", "Gilbertus m\u00f3wi co wa\u017cne, zanim zapytasz", "12 dni"),
        ("2: Serwer + Oceny", "Tyg 7\u201310 (maj\u2013czerwiec)", "Hetzner + pipeline ocen pracownik\u00f3w", "10 dni"),
        ("3: Omnius Lite REH", "Tyg 11\u201316 (czerwiec\u2013lipiec)", "Roch ma AI asystenta", "12 dni"),
        ("4: Pe\u0142ne oceny + REF", "Tyg 17\u201320 (sierpie\u0144)", "Multi-perspektywiczne oceny + Krystian", "11 dni"),
        ("5: Autonomiczne akcje", "Tyg 21\u201326 (wrzesie\u0144)", "Od Q&A do proactive management", "12 dni"),
    ])

    # Faza 0
    styled_heading(doc, "Faza 0: STABILIZACJA (kwiecień, 2 tygodnie)", 2)
    bullet(doc, "Connection pooling \u2014 psycopg_pool.ConnectionPool singleton")
    bullet(doc, "Teams message grouping \u2014 per chat per 2h okno czasowe")
    bullet(doc, "Entity deduplication \u2014 jednorazowy + pg_trgm do ongoing")
    bullet(doc, "Bridge people\u2194entities \u2014 entity_id w tabeli people")
    bullet(doc, "Event taxonomy +7 typ\u00f3w: deadline, commitment, escalation, blocker, task_assignment, approval, rejection")
    bullet(doc, "Aktywacja alert\u00f3w jako cron (po ekstrakcji >80%)")

    # Faza 1
    styled_heading(doc, "Faza 1: PROAKTYWNA INTELIGENCJA (kwiecień\u2013maj, 4 tygodnie)", 2)
    bullet(doc, "Sync kalendarza z Microsoft Graph API")
    bullet(doc, "Morning brief z kontekstem kalendarza + relacji + open loops")
    bullet(doc, "Person-aware retrieval \u2014 auto-expand alias\u00f3w przy /ask")
    bullet(doc, "Decision journal via WhatsApp (keyword \u201egtd:\u201d)")
    bullet(doc, "Cross-domain correlation MVP (konflikty \u2192 trading performance)")

    # Faza 2
    styled_heading(doc, "Faza 2: SERWER + PIPELINE OCEN (maj\u2013czerwiec, 4 tygodnie)", 2)
    bullet(doc, "Migracja na Hetzner AX102 (128GB RAM, GPU-ready)")
    bullet(doc, "HTTPS + monitoring (Uptime Kuma lub Prometheus)")
    bullet(doc, "Pipeline ocen: POST /evaluate \u2192 data_collector \u2192 evaluator (Claude) \u2192 formatter (DOCX)")
    bullet(doc, "Ka\u017cda ocena z explicit confidence score i flagami perspektywy")
    bullet(doc, "Kwartalny cron auto-generowania ocen")

    # Faza 3
    styled_heading(doc, "Faza 3: OMNIUS LITE \u2014 REH (czerwiec\u2013lipiec, 6 tygodni)", 2)
    body(doc,
        "Kluczowy insight: Omnius V1 przetwarza WY\u0141\u0104CZNIE dane firmowe (SharePoint, shared Teams channels, "
        "shared mailboxy, raporty) \u2014 NIE dane pracownicze. Nie wymaga compliance na monitoring.")
    bullet(doc, "Osobny Docker Compose, osobny Postgres, osobna Qdrant collection")
    bullet(doc, "Graph API z admin consent Rocha \u2014 SharePoint + shared Teams")
    bullet(doc, "Interfejs Rocha: Teams Bot lub web chat")
    bullet(doc, "Gilbertus\u2194Omnius read-only sync (nightly summaries)")
    bullet(doc, "R\u00f3wnolegle: compliance review z prawnikiem (Mi\u0142osz Awedyk)")

    # Faza 4
    styled_heading(doc, "Faza 4: PE\u0141NE OCENY + OMNIUS REF (sierpie\u0144, 4 tygodnie)", 2)
    bullet(doc, "Enhanced evaluations z danymi Omnius (email response time, task completion)")
    bullet(doc, "Omnius REF \u2014 klon REH dla Krystiana Juchacza")
    bullet(doc, "Anomaly detection per osoba (odchylenie od baseline\u2019u komunikacji)")
    bullet(doc, "Employee scorecard dashboard")
    bullet(doc, "Decision journal outcomes tracking (auto-remind 7/30/90 dni)")

    # Faza 5
    styled_heading(doc, "Faza 5: AUTONOMICZNE AKCJE + RAPORTOWANIE (wrzesie\u0144, 6 tygodni)", 2)
    bullet(doc, "Action item pipeline: Gilbertus proponuje \u2192 Sebastian zatwierdza via WhatsApp \u2192 wykonanie")
    bullet(doc, "Automated weekly reports: per-company + cross-company + personal")
    bullet(doc, "Kwartalny evaluation cycle (auto-trigger, auto-generate, auto-archive)")
    bullet(doc, "Process inefficiency detector (\u201eten proces kosztuje 40h/mc, automatyzacja 30h\u201d)")
    bullet(doc, "Feedback loop: thumbs up/down na odpowiedzi")

    # Architektura
    styled_heading(doc, "4. Architektura docelowa", 1)
    body(doc,
        "Gilbertus (prywatna infra) = god-view. Omnius REH + Omnius REF (infra sp\u00f3\u0142ek) = agenci firmowi. "
        "Izolacja: Omniusy nie komunikuj\u0105 si\u0119 mi\u0119dzy sob\u0105. Gilbertus odbiera summaries, nie raw data.")

    # Ryzyka
    styled_heading(doc, "5. Ryzyka", 1)
    add_table(doc, ["Ryzyko", "Prawdop.", "Impact", "Mitygacja"],
    [
        ("Compliance blokuje people analytics", "Wysokie", "Wysoki", "Faza 3 dzia\u0142a BEZ danych pracowniczych"),
        ("Roch/Krystian nie adoptuj\u0105 Omnius", "\u015arednie", "Wysoki", "Start od pain point, demo z ich danymi"),
        ("Koszty API rosn\u0105", "\u015arednie", "\u015aredni", "Haiku do ekstrakcji, Sonnet do syntezy"),
        ("Migracja serwera = utrata danych", "Niskie", "Krytyczny", "Full backup + test restore + fallback"),
    ])

    doc.save(path)
    print(f"Full plan: {path}")


# =========================================================================
# REH PLAN (for Roch)
# =========================================================================
def generate_reh_plan(path):
    doc = Document()
    set_margins(doc)
    title_page(doc,
        "OMNIUS REH",
        "AI Asystent dla Respect Energy Holding\nPlan wdro\u017cenia Q2\u2013Q3 2026")

    styled_heading(doc, "1. Czym jest Omnius?", 1)
    body(doc,
        "Omnius to firmowy asystent AI \u2014 system, kt\u00f3ry indeksuje dokumenty firmowe, "
        "rozmowy z shared Teams channels, raporty i dokumentacj\u0119 projektow\u0105, "
        "a nast\u0119pnie umo\u017cliwia szybkie wyszukiwanie i analiz\u0119 informacji "
        "za pomoc\u0105 naturalnego j\u0119zyka.")
    body(doc,
        "Omnius NIE monitoruje prywatnych rozm\u00f3w pracownik\u00f3w. Przetwarza wy\u0142\u0105cznie "
        "wsp\u00f3\u0142dzielone dane firmowe: SharePoint, publiczne kana\u0142y Teams, shared mailboxy, "
        "raporty finansowe i dokumenty projektowe.")

    styled_heading(doc, "2. Co Omnius potrafi?", 1)
    label(doc, "Faza 1 (czerwiec\u2013lipiec 2026):")
    bullet(doc, "Wyszukiwanie dokument\u00f3w firmowych: \u201eznajd\u017a ostatni\u0105 wersj\u0119 umowy z NOFAR\u201d")
    bullet(doc, "Status projekt\u00f3w: \u201ejaki jest status projektu BESS?\u201d")
    bullet(doc, "Analiza rozm\u00f3w: \u201eco trading team omawia\u0142 w kanale energia w zesz\u0142ym tygodniu?\u201d")
    bullet(doc, "Kontrakty: \u201epoka\u017c kontrakty >1M PLN podpisane w Q1 2026\u201d")
    bullet(doc, "Morning brief: codzienny raport z najwa\u017cniejszymi informacjami")

    label(doc, "Faza 2 (sierpie\u0144 2026, po review compliance):")
    bullet(doc, "People analytics: wzorce komunikacji, response time, task completion")
    bullet(doc, "Anomaly detection: spadki aktywno\u015bci, zmiany wzorc\u00f3w")
    bullet(doc, "Employee evaluations: automatyczne oceny kwartalne")

    label(doc, "Faza 3 (wrzesie\u0144 2026):")
    bullet(doc, "Wykonywanie akcji: wysy\u0142anie emaili, tworzenie ticket\u00f3w, scheduling spotka\u0144")
    bullet(doc, "Weekly reports: automatyczne raporty tygodniowe")
    bullet(doc, "Process optimization: wykrywanie nieefektywno\u015bci")

    styled_heading(doc, "3. Interfejs", 1)
    body(doc,
        "Omnius dost\u0119pny jest przez Teams Bot lub web chat. "
        "Pytania w j\u0119zyku naturalnym \u2014 po polsku lub angielsku. "
        "Odpowiedzi zawieraj\u0105 \u017ar\u00f3d\u0142a (link do dokumentu, nazw\u0119 kana\u0142u, dat\u0119).")

    styled_heading(doc, "4. Dane i bezpiecze\u0144stwo", 1)
    bullet(doc, "Omnius REH \u017cyje na dedykowanym serwerze (Hetzner, Niemcy)")
    bullet(doc, "Osobna baza danych \u2014 dane REH nie mieszaj\u0105 si\u0119 z innymi sp\u00f3\u0142kami")
    bullet(doc, "Szyfrowanie: at-rest (disk encryption) + in-transit (HTTPS/mTLS)")
    bullet(doc, "Audit log: ka\u017cdy dost\u0119p i zapytanie logowane")
    bullet(doc, "Faza 1: TYLKO wsp\u00f3\u0142dzielone dane firmowe (SharePoint, shared Teams, raporty)")
    bullet(doc, "Faza 2: people analytics dopiero po review compliance z prawnikiem")

    styled_heading(doc, "5. Wymagania od REH", 1)
    bullet(doc, "Admin consent na Microsoft Graph API (SharePoint + shared Teams channels)")
    bullet(doc, "Wskazanie kluczowych SharePoint libraries i Teams kana\u0142\u00f3w do zaindeksowania")
    bullet(doc, "Feedback od Rocha: jakie pytania zadaje najcz\u0119\u015bciej? Czego szuka?")
    bullet(doc, "Czas: 2-3 spotkania onboardingowe (demo, pierwsze pytania, iteracja)")

    styled_heading(doc, "6. Harmonogram", 1)
    add_table(doc, ["Etap", "Termin", "Co dostarczamy"],
    [
        ("Setup infrastruktury", "Czerwiec 2026, tydz. 1\u20132", "Serwer, baza danych, indeksowanie dokument\u00f3w"),
        ("Graph API + indeksowanie", "Czerwiec 2026, tydz. 2\u20133", "SharePoint, shared Teams, raporty"),
        ("Interfejs Rocha (Teams Bot)", "Czerwiec 2026, tydz. 3\u20134", "Dzia\u0142aj\u0105cy bot do pyta\u0144"),
        ("Onboarding + iteracja", "Lipiec 2026, tydz. 1\u20132", "Demo, feedback, dostrojenie"),
        ("People analytics (warunkowe)", "Sierpie\u0144 2026", "Po review compliance"),
        ("Autonomiczne akcje", "Wrzesie\u0144 2026", "Emaile, tickety, raporty"),
    ])

    styled_heading(doc, "7. Przyk\u0142ady pyta\u0144", 1)
    body(doc, "Roch mo\u017ce pyta\u0107 Omniusa:")
    bullet(doc, "\u201eJaki jest status projektu BESS?\u201d")
    bullet(doc, "\u201ePoka\u017c wszystkie pending kontrakty powy\u017cej 1M PLN\u201d")
    bullet(doc, "\u201eCo Maja Kalinowska raportowa\u0142a w zesz\u0142ym tygodniu?\u201d")
    bullet(doc, "\u201eZnajd\u017a ostatni raport finansowy za Q4 2025\u201d")
    bullet(doc, "\u201ePodsumuj dyskusj\u0119 o GoldenPeaks z ostatnich 2 tygodni\u201d")
    bullet(doc, "\u201eKto odpowiada za kontrakt z Cognor i jaki jest deadline?\u201d")

    doc.save(path)
    print(f"REH plan: {path}")


# =========================================================================
# REF PLAN (for Krystian)
# =========================================================================
def generate_ref_plan(path):
    doc = Document()
    set_margins(doc)
    title_page(doc,
        "OMNIUS REF",
        "AI Asystent dla Respect Energy Fuels\nPlan wdro\u017cenia Q3 2026")

    styled_heading(doc, "1. Czym jest Omnius?", 1)
    body(doc,
        "Omnius to firmowy asystent AI \u2014 system, kt\u00f3ry indeksuje dokumenty firmowe, "
        "rozmowy z shared Teams channels, raporty i dokumentacj\u0119 projektow\u0105, "
        "a nast\u0119pnie umo\u017cliwia szybkie wyszukiwanie i analiz\u0119 informacji "
        "za pomoc\u0105 naturalnego j\u0119zyka.")
    body(doc,
        "Omnius NIE monitoruje prywatnych rozm\u00f3w pracownik\u00f3w. Przetwarza wy\u0142\u0105cznie "
        "wsp\u00f3\u0142dzielone dane firmowe: SharePoint, publiczne kana\u0142y Teams, shared mailboxy, "
        "raporty finansowe i dokumenty projektowe.")

    styled_heading(doc, "2. Co Omnius potrafi?", 1)
    label(doc, "Faza 1 (sierpie\u0144 2026):")
    bullet(doc, "Wyszukiwanie dokument\u00f3w firmowych: \u201eznajd\u017a umow\u0119 na dostawy LNG z Q1 2026\u201d")
    bullet(doc, "Status projekt\u00f3w: \u201ejaki jest status regasyfikacji?\u201d")
    bullet(doc, "Analiza rozm\u00f3w: \u201eco zesp\u00f3\u0142 logistyki omawia\u0142 w zesz\u0142ym tygodniu?\u201d")
    bullet(doc, "Kontrakty: \u201epoka\u017c aktywne kontrakty gazowe\u201d")
    bullet(doc, "Morning brief: codzienny raport z najwa\u017cniejszymi informacjami REF")

    label(doc, "Faza 2 (wrzesie\u0144 2026, po compliance):")
    bullet(doc, "People analytics: wzorce komunikacji, response time")
    bullet(doc, "Anomaly detection: spadki aktywno\u015bci, zmiany wzorc\u00f3w")
    bullet(doc, "Employee evaluations: automatyczne oceny kwartalne")
    bullet(doc, "Autonomiczne akcje: emaile, tickety, raporty tygodniowe")

    styled_heading(doc, "3. Interfejs", 1)
    body(doc,
        "Omnius dost\u0119pny jest przez Teams Bot lub web chat. "
        "Pytania w j\u0119zyku naturalnym \u2014 po polsku lub angielsku. "
        "Odpowiedzi zawieraj\u0105 \u017ar\u00f3d\u0142a (link do dokumentu, nazw\u0119 kana\u0142u, dat\u0119).")

    styled_heading(doc, "4. Dane i bezpiecze\u0144stwo", 1)
    bullet(doc, "Omnius REF \u017cyje na dedykowanym serwerze (Hetzner, Niemcy)")
    bullet(doc, "Osobna baza danych \u2014 dane REF nie mieszaj\u0105 si\u0119 z innymi sp\u00f3\u0142kami")
    bullet(doc, "Szyfrowanie: at-rest (disk encryption) + in-transit (HTTPS/mTLS)")
    bullet(doc, "Audit log: ka\u017cdy dost\u0119p i zapytanie logowane")
    bullet(doc, "TYLKO wsp\u00f3\u0142dzielone dane firmowe \u2014 bez prywatnych czat\u00f3w pracownik\u00f3w")

    styled_heading(doc, "5. Wymagania od REF", 1)
    bullet(doc, "Admin consent na Microsoft Graph API (SharePoint + shared Teams channels)")
    bullet(doc, "Wskazanie kluczowych SharePoint libraries i Teams kana\u0142\u00f3w do zaindeksowania")
    bullet(doc, "Feedback od Krystiana: jakie pytania zadaje najcz\u0119\u015bciej? Czego szuka?")
    bullet(doc, "Czas: 2-3 spotkania onboardingowe (demo, pierwsze pytania, iteracja)")

    styled_heading(doc, "6. Harmonogram", 1)
    add_table(doc, ["Etap", "Termin", "Co dostarczamy"],
    [
        ("Omnius REH (pilot)", "Czerwiec\u2013lipiec 2026", "Wdro\u017cenie pilotowe na REH, walidacja podej\u015bcia"),
        ("Setup infrastruktury REF", "Sierpie\u0144 2026, tydz. 1\u20132", "Serwer, baza danych, indeksowanie"),
        ("Graph API + indeksowanie", "Sierpie\u0144 2026, tydz. 2\u20133", "SharePoint, shared Teams, raporty"),
        ("Interfejs Krystiana", "Sierpie\u0144 2026, tydz. 3\u20134", "Teams Bot dzia\u0142aj\u0105cy"),
        ("Onboarding + iteracja", "Wrzesie\u0144 2026, tydz. 1\u20132", "Demo, feedback, dostrojenie"),
        ("Zaawansowane funkcje", "Wrzesie\u0144 2026+", "People analytics, akcje, raporty"),
    ])

    styled_heading(doc, "7. Przyk\u0142ady pyta\u0144", 1)
    body(doc, "Krystian mo\u017ce pyta\u0107 Omniusa:")
    bullet(doc, "\u201eJaki jest status regasyfikacji i sprzeda\u017cy?\u201d")
    bullet(doc, "\u201ePoka\u017c aktywne kontrakty gazowe z Q1 2026\u201d")
    bullet(doc, "\u201eCo zesp\u00f3\u0142 logistyki raportowa\u0142 o dostawach LNG?\u201d")
    bullet(doc, "\u201eZnajd\u017a raport P&L za stycze\u0144 2026\u201d")
    bullet(doc, "\u201ePodsumuj ustalenia z ostatniego spotkania zarz\u0105du\u201d")
    bullet(doc, "\u201eKto odpowiada za kontrakt z PGNiG i kiedy jest deadline?\u201d")

    doc.save(path)
    print(f"REF plan: {path}")


if __name__ == "__main__":
    base = "/mnt/c/Users/jablo/Desktop"
    generate_full_plan(f"{base}/Gilbertus_Omnius_Plan_Q2Q3_2026.docx")
    generate_reh_plan(f"{base}/Omnius_REH_Plan_dla_Rocha.docx")
    generate_ref_plan(f"{base}/Omnius_REF_Plan_dla_Krystiana.docx")
