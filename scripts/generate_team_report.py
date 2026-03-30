#!/usr/bin/env python3
"""Generate Gilbertus AI Team Building Report as DOCX."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import sys

doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Helper functions
def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    return h

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Medium Shading 1 Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10)
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)
            for p in row_cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()
    return table

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(f" {text}")
    else:
        p.add_run(text)
    return p

# ============================================================
# TITLE PAGE
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading('GILBERTUS ALBANS', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_heading('Strategia Budowy Zespolu AI Development', level=1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('Przygotowane dla: Sebastian Jablonski\n').bold = True
meta.add_run('Data: 30 marca 2026\n')
meta.add_run('Wersja: 1.0\n')
meta.add_run('Klasyfikacja: POUFNE')
doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS
# ============================================================
add_heading('Spis tresci', 1)
toc_items = [
    '1. Streszczenie wykonawcze',
    '2. Stan projektu Gilbertus — kontekst',
    '3. Analiza najlepszych praktyk teamow AI',
    '4. Rekomendowany sklad zespolu',
    '5. Profile stanowisk i wymagania',
    '6. Model wspolpracy i workflow',
    '7. Budzety i koszty',
    '8. Proces onboardingu',
    '9. KPI i metryki sukcesu',
    '10. Ryzyka i mitygacja',
    '11. Harmonogram budowy zespolu',
]
for item in toc_items:
    doc.add_paragraph(item, style='List Number')
doc.add_page_break()

# ============================================================
# 1. EXECUTIVE SUMMARY
# ============================================================
add_heading('1. Streszczenie wykonawcze', 1)
doc.add_paragraph(
    'Gilbertus Albans to zaawansowany system AI mentat, ktory osiagnal etap rozwoju wymagajacy '
    'profesjonalnego zespolu. Projekt obejmuje ~135,000 linii kodu (Python + TypeScript), '
    '128 tabel w bazie danych, 189 endpointow API, 44 narzedzia MCP, 56 cronow automatyzacji '
    'i integracje z 10+ zrodlami danych.'
)
doc.add_paragraph(
    'Na podstawie analizy najlepszych praktyk budowy teamow AI w 2025-2026 oraz specyfiki '
    'projektu Gilbertus, rekomenduje budowe zespolu 4-5 osob o nastepujacej strukturze:'
)

add_table(
    ['Rola', 'Priorytet', 'Forma'],
    [
        ['Senior AI/Backend Engineer (Tech Lead)', 'KRYTYCZNY — hire #1', 'B2B full-time'],
        ['Senior Data & Integration Engineer', 'KRYTYCZNY — hire #2', 'B2B full-time'],
        ['Mid/Senior Frontend + DevOps Engineer', 'WYSOKI — hire #3', 'B2B full-time'],
        ['AI/ML Specialist (RAG & Embeddings)', 'WYSOKI — hire #4', 'B2B part-time lub full-time'],
        ['QA & Security Engineer (opcjonalnie)', 'SREDNI — hire #5', 'B2B part-time'],
    ]
)

doc.add_paragraph(
    'Szacowany miesieczny budzet: 85,000-130,000 PLN netto (B2B). '
    'ROI: przyspieszenie rozwoju 3-5x, eliminacja single point of failure (Sebastian), '
    'produkcyjne wdrozenie na Hetzner w ciagu 4-6 tygodni od sformowania zespolu.'
)
doc.add_page_break()

# ============================================================
# 2. STAN PROJEKTU
# ============================================================
add_heading('2. Stan projektu Gilbertus — kontekst', 1)

add_heading('2.1 Metryki codebase', 2)
add_table(
    ['Metryka', 'Wartosc', 'Komentarz'],
    [
        ['Python LoC', '~82,500', 'Backend, AI, extraction, analysis'],
        ['JavaScript/TypeScript LoC', '~52,400', 'Frontend Omnius, integracje'],
        ['Shell scripts', '109', 'Automatyzacja, deployment, crony'],
        ['Pliki testowe', '11', 'KRYTYCZNIE NISKI — wymaga natychmiastowej poprawy'],
        ['Tabele PostgreSQL', '128', 'Partycjonowane, z migracjami'],
        ['Kolekcje Qdrant', '1+', '~105k chunków wektorowych'],
        ['API endpointy', '189', 'FastAPI, rate-limited'],
        ['MCP tools', '44', '6 grup: core, system, people, business, finance, omnius'],
        ['Cron joby', '56', 'Od co 2 min do monthly'],
    ]
)

add_heading('2.2 Kluczowe obszary technologiczne', 2)
areas = [
    ('Agentic AI Architecture:', 'LangGraph workflows, Claude API, multi-model routing, tool orchestration'),
    ('Data Ingestion Pipeline:', 'Microsoft Graph (Teams, Email, Calendar), WhatsApp, Plaud audio, dokumenty, ChatGPT'),
    ('Vector Search & RAG:', 'Qdrant, OpenAI embeddings, semantic chunking, retrieval pipeline'),
    ('Database Layer:', 'PostgreSQL 16, 128 tabel, migracje, connection pooling, partycjonowanie'),
    ('Delivery & Interaction:', 'WhatsApp Bot, Teams Bot, REST API, Voice (TTS/STT)'),
    ('Multi-tenant System (Omnius):', 'React frontend, RBAC, tenant isolation, user management'),
    ('Security & Compliance:', 'Audit trail, rate limiting, gitleaks, pip-audit, input validation'),
    ('Monitoring & Self-healing:', 'Data Guardian, extraction watchdog, non-regression, quality checks'),
]
for prefix, text in areas:
    add_bullet(text, prefix)

add_heading('2.3 Obecne wyzwania wymagajace zespolu', 2)
challenges = [
    'Single point of failure — Sebastian jest jedynym developerem',
    'Pokrycie testami ~2% — krytycznie niskie dla systemu produkcyjnego',
    'Deployment wciaz na WSL2 — potrzebna migracja na Hetzner VPS',
    'Graph API sync niestabilny — 400 errors na Teams pagination',
    'Koszty API bez kontroli — $724/24h na extraction (naprawione dzis)',
    'Brak code review — caly kod przechodzi bez peer review',
    'Dokumentacja architekturalna fragmentaryczna',
    'Frontend Omnius wymaga dopracowania UX',
]
for c in challenges:
    add_bullet(c)
doc.add_page_break()

# ============================================================
# 3. BEST PRACTICES
# ============================================================
add_heading('3. Analiza najlepszych praktyk teamow AI', 1)

add_heading('3.1 Model "Lean AI Pod" (2025-2026)', 2)
doc.add_paragraph(
    'Wedlug badan OpenAI, Anthropic i wiodacych firm AI w 2025-2026, najefektywniejsze '
    'zespoly AI to male, wysoce wyspecjalizowane "pody" skladajace sie z 3-5 seniorow, '
    'gdzie kazdy czlonek pokrywa 2-3 domeny. Kluczowe zasady:'
)
principles = [
    ('Industrialization over Novelty:', 'Zespol 2026 musi priorytetyzowac niezawodnosc i architekture systemu nad tuning parametrow modeli.'),
    ('AI-Augmented Productivity:', 'Jeden senior AI engineer z narzediami AI (Claude Code, Copilot) osiaga produktywnosc 5+ tradycyjnych developerow.'),
    ('Cross-functional Ownership:', 'Kazdy czlonek zespolu jest wlascicielem end-to-end swojego obszaru — od kodu po monitoring.'),
    ('60:40 Senior/Junior Ratio:', 'Dla malych teamow: minimum 60% seniorow. W naszym przypadku: 100% senior/mid — uzasadnione zlozonoscia systemu.'),
    ('Flat Structure:', 'Brak hierarchii managerskiej. Tech Lead + peers. Sebastian jako Product Owner.'),
]
for prefix, text in principles:
    add_bullet(text, prefix)

add_heading('3.2 Dlaczego NIE 10-osobowy team', 2)
doc.add_paragraph(
    'Dla projektu o tej zlozonosci i na tym etapie rozwoju, duzy zespol bylby kontrproduktywny:'
)
reasons = [
    'Overhead komunikacyjny rosnie kwadratowo z rozmiarem zespolu (prawo Brooksa)',
    'Gilbertus ma specyficzny context biznesowy (REH/REF), ktorego nauczenie wymaga czasu',
    'Codebase jest dobrze ustrukturyzowany ale slabo udokumentowany — onboarding jest waskim gardlem',
    'AI-augmented workflow (Claude Code) mnozy efektywnosc seniorow 3-5x',
    'Sebastian musi miec pelna kontrole nad kierunkiem — mniej osob = latwiejszy oversight',
]
for r in reasons:
    add_bullet(r)

add_heading('3.3 Referencyjne struktury teamow', 2)
add_table(
    ['Model', 'Rozmiar', 'Zastosowanie', 'Fit dla Gilbertusa'],
    [
        ['Startup Pod', '3-4', 'MVP, rapid iteration', 'Za maly — brak pokrycia specjalizacji'],
        ['Growth Squad', '5-7', 'Scaling, production hardening', 'OPTYMALNY — pelne pokrycie'],
        ['CoE (Center of Excellence)', '8-15', 'Enterprise, multiple products', 'Za duzy — overhead'],
        ['Full Stack AI Team', '10+', 'Duze organizacje, R&D', 'Nieadekwatny dla 1-product team'],
    ]
)
doc.add_page_break()

# ============================================================
# 4. REKOMENDOWANY SKLAD
# ============================================================
add_heading('4. Rekomendowany sklad zespolu', 1)

doc.add_paragraph(
    'Rekomenduje model "Growth Squad" — 4 core + 1 opcjonalny, z Sebastianem jako Product Owner. '
    'Kazda rola pokrywa 2-3 domeny, co daje pelne pokrycie przy minimalnym headcount.'
)

add_heading('Struktura zespolu', 2)

# Role diagram as table
add_table(
    ['#', 'Rola', 'Domeny', 'FTE', 'Hire Priority'],
    [
        ['1', 'Senior AI/Backend Engineer\n(Tech Lead)', 'Python/FastAPI, LangGraph,\nClaude API, architektura', '1.0', 'Miesiac 1\nKRYTYCZNY'],
        ['2', 'Senior Data & Integration\nEngineer', 'PostgreSQL, Qdrant, Graph API,\npipeline ingestion', '1.0', 'Miesiac 1-2\nKRYTYCZNY'],
        ['3', 'Mid/Senior Frontend +\nDevOps Engineer', 'React/TS (Omnius), Docker,\nHetzner, CI/CD, monitoring', '1.0', 'Miesiac 2-3\nWYSOKI'],
        ['4', 'AI/ML Specialist\n(RAG & Agents)', 'Embeddings, RAG optimization,\nprompt engineering, evals', '0.5-1.0', 'Miesiac 2-3\nWYSOKI'],
        ['5', 'QA & Security Engineer\n(opcjonalnie)', 'Testy, security audits,\npen-testing, compliance', '0.5', 'Miesiac 4+\nSREDNI'],
    ]
)

add_heading('Role Sebastiana', 2)
doc.add_paragraph(
    'Sebastian pelni role Product Owner & Domain Expert — jedyna osoba z pelnym kontekstem '
    'biznesowym REH/REF. Odpowiedzialny za:'
)
sebastian_roles = [
    'Priorytetyzacja backlogu i roadmapy (Masterplan)',
    'Akceptacja / odrzucanie PR-ow (final approval)',
    'Domain knowledge transfer — kontekst biznesowy, relacje, procesy',
    'Testowanie z perspektywy uzytkownika koncowego',
    'Decyzje strategiczne: co budowac, co kupic, co odlozyc',
]
for r in sebastian_roles:
    add_bullet(r)
doc.add_page_break()

# ============================================================
# 5. PROFILE STANOWISK
# ============================================================
add_heading('5. Profile stanowisk i wymagania', 1)

# ROLE 1
add_heading('5.1 Senior AI/Backend Engineer (Tech Lead)', 2)
doc.add_paragraph('Kluczowa rola — pierwsza osoba do zatrudnienia. Bedzie architektem technicznym i oparciem dla zespolu.')

add_heading('Wymagania twarde', 3)
reqs = [
    '5+ lat doswiadczenia w Python (FastAPI/Django/Flask)',
    'Doswiadczenie z LLM APIs (Claude, OpenAI) w produkcji',
    'Znajomosc LangChain/LangGraph lub podobnych frameworkow agentowych',
    'PostgreSQL — zaawansowane (partycjonowanie, optymalizacja, migracje)',
    'Doswiadczenie z vector databases (Qdrant, Pinecone, Weaviate)',
    'Docker, Linux, CI/CD',
    'Git workflow, code review jako reviewer',
]
for r in reqs:
    add_bullet(r)

add_heading('Wymagania miekkie', 3)
soft = [
    'Zdolnosc do samodzielnej pracy i podejmowania decyzji architekturalnych',
    'Komunikatywnosc — bedzie interfejsem miedzy Sebastianem a reszta zespolu',
    'Proaktywnosc w identyfikacji problemow i proponowaniu rozwiazan',
    'Doswiadczenie w mentoringu juniorow/midow',
]
for s in soft:
    add_bullet(s)

add_heading('Kluczowe zadania (pierwsze 3 miesiace)', 3)
tasks = [
    'Audit codebase i stworzenie dokumentacji architekturalnej',
    'Setup CI/CD pipeline (GitHub Actions + testy)',
    'Pokrycie testami krytycznych sciezek (extraction, ingestion, API)',
    'Migracja na Hetzner VPS (production deployment)',
    'Refactoring connection pooling i error handling',
    'Code review wszystkich PR-ow',
]
for t in tasks:
    add_bullet(t)

doc.add_paragraph()

# ROLE 2
add_heading('5.2 Senior Data & Integration Engineer', 2)
doc.add_paragraph('Odpowiedzialny za przeplywy danych — od ingestion po delivery.')

add_heading('Wymagania twarde', 3)
reqs2 = [
    '4+ lat doswiadczenia z pipeline\'ami danych (ETL/ELT)',
    'PostgreSQL zaawansowany (query optimization, indeksy, EXPLAIN ANALYZE)',
    'Doswiadczenie z REST API integracjami (Microsoft Graph API to duzy plus)',
    'Python — ETL, async, concurrent processing',
    'Doswiadczenie z message queues (opcjonalnie: Redis, RabbitMQ)',
    'Monitoring i alerting (Grafana, Prometheus lub custom)',
    'Znajomosc WhatsApp Business API lub podobnych messaging APIs',
]
for r in reqs2:
    add_bullet(r)

add_heading('Kluczowe zadania', 3)
tasks2 = [
    'Naprawa i stabilizacja Graph API sync (Teams, Email, Calendar)',
    'Optymalizacja extraction pipeline (cost reduction, throughput)',
    'Monitoring ingestion health — dashboardy, alerty',
    'Qdrant optimization — re-chunking, embedding quality',
    'Backup strategy hardening (point-in-time recovery)',
    'Data quality pipeline — walidacja, anomaly detection',
]
for t in tasks2:
    add_bullet(t)

doc.add_paragraph()

# ROLE 3
add_heading('5.3 Mid/Senior Frontend + DevOps Engineer', 2)
doc.add_paragraph('Hybryda: Omnius UI + infrastruktura produkcyjna.')

add_heading('Wymagania twarde', 3)
reqs3 = [
    '3+ lat React/TypeScript',
    'Docker Compose, Docker networking',
    'Linux server administration (Debian/Ubuntu)',
    'Nginx/Caddy jako reverse proxy z SSL',
    'GitHub Actions lub GitLab CI',
    'Monitoring (uptimerobot, healthchecks.io, custom)',
    'Podstawowa znajomosc PostgreSQL i REST API',
]
for r in reqs3:
    add_bullet(r)

add_heading('Kluczowe zadania', 3)
tasks3 = [
    'Deployment Omnius na Hetzner (production-ready)',
    'CI/CD pipeline: lint, test, build, deploy',
    'Omnius UX improvements — dashboardy, real-time updates',
    'SSL, domain, firewall, backup automation',
    'Monitoring stack: uptime, latency, error rates',
    'Multi-tenant UI (REH Roch, REF Krystian)',
]
for t in tasks3:
    add_bullet(t)

doc.add_paragraph()

# ROLE 4
add_heading('5.4 AI/ML Specialist (RAG & Agents)', 2)
doc.add_paragraph('Specjalista od jakosci AI — odpowiedzi, embedding, agentow.')

add_heading('Wymagania twarde', 3)
reqs4 = [
    '3+ lat doswiadczenia z NLP/ML w produkcji',
    'Gleboka znajomosc RAG (chunking strategies, re-ranking, hybrid search)',
    'Prompt engineering — systematic evaluation, A/B testing',
    'Doswiadczenie z evaluation frameworks (RAGAS, LangSmith, custom)',
    'Vector databases — indexing strategies, similarity metrics',
    'Fine-tuning lub distillation (nice to have)',
    'Znajomosc Anthropic Claude API (tool use, system prompts)',
]
for r in reqs4:
    add_bullet(r)

add_heading('Kluczowe zadania', 3)
tasks4 = [
    'Audit jakosci odpowiedzi Gilbertusa (evaluation pipeline)',
    'Optymalizacja chunking strategy (rozmiar, overlap, metadata)',
    'Implementacja re-ranking i hybrid search',
    'Prompt optimization dla extraction (entities, events, commitments)',
    'RAG evaluation dashboard — metryki jakosci w czasie',
    'Agentic workflow optimization (LangGraph)',
]
for t in tasks4:
    add_bullet(t)

doc.add_paragraph()

# ROLE 5
add_heading('5.5 QA & Security Engineer (opcjonalnie)', 2)
doc.add_paragraph('Part-time, od miesiaca 4+. Odpowiedzialny za quality gates i security.')

add_heading('Wymagania', 3)
reqs5 = [
    'Doswiadczenie z pytest, integration testing',
    'Security auditing (OWASP Top 10, pen-testing)',
    'API security (rate limiting, auth, injection prevention)',
    'Compliance awareness (RODO/GDPR, data retention)',
    'Load testing (locust, k6)',
]
for r in reqs5:
    add_bullet(r)
doc.add_page_break()

# ============================================================
# 6. MODEL WSPOLPRACY
# ============================================================
add_heading('6. Model wspolpracy i workflow', 1)

add_heading('6.1 Narzedzia', 2)
add_table(
    ['Kategoria', 'Narzedzie', 'Cel'],
    [
        ['Kod', 'GitHub (private repo)', 'Wersjonowanie, PR, code review'],
        ['CI/CD', 'GitHub Actions', 'Testy, lint, deploy'],
        ['Komunikacja', 'Discord lub Slack (prywatny)', 'Daily comms, alerty'],
        ['Task management', 'GitHub Issues + Projects', 'Backlog, sprint planning'],
        ['Dokumentacja', 'Markdown w repo + Notion', 'Architecture Decision Records (ADR)'],
        ['AI Coding', 'Claude Code', 'Kazdy developer uzywa CC do codziennej pracy'],
        ['Monitoring', 'Grafana + custom dashboardy', 'System health, metryki'],
    ]
)

add_heading('6.2 Rytm pracy', 2)
add_table(
    ['Kiedy', 'Co', 'Kto', 'Format'],
    [
        ['Codziennie 9:00', 'Async standup', 'Wszyscy', 'Wiadomosc na Discord/Slack\n(co zrobilem, co robie, blokery)'],
        ['Poniedzialek 10:00', 'Sprint planning (1h)', 'Wszyscy + Sebastian', 'Video call — priorytetyzacja tygodnia'],
        ['Sroda 15:00', 'Tech sync (30 min)', 'Developerzy', 'Sync techniczny, code review backlog'],
        ['Piatek 16:00', 'Demo & retro (1h)', 'Wszyscy + Sebastian', 'Demo zrobionego, feedback, retro'],
        ['Ad hoc', 'Pair programming', '2 osoby', 'Przy zlozonych problemach'],
    ]
)

add_heading('6.3 Git workflow', 2)
doc.add_paragraph('Trunk-based development z feature branches:')
git_rules = [
    'main branch = production (protected, wymaga PR + approval)',
    'Feature branches: feature/TICKET-description',
    'Kazdy PR wymaga: 1 approval + CI green + testy',
    'Squash merge do main',
    'Sebastian ma final approval na krytyczne zmiany (architektura, security, prompty)',
    'Tech Lead ma approval na standardowe zmiany',
    'Kazdy commit musi przejsc pre-commit hooks (gitleaks, ruff, mypy)',
]
for r in git_rules:
    add_bullet(r)

add_heading('6.4 Zasady pracy z AI', 2)
ai_rules = [
    'Kazdy developer MUSI uzywac Claude Code do codziennej pracy (wymagane)',
    'AI-generated code podlega identycznym standardom code review jak reczny',
    'Prompty systemowe i extraction prompty — zmiany tylko po aprovacie Tech Leada',
    'Koszty API monitorowane dziennie — hard cap na kazdym module',
    'CLAUDE.md w repo to "zrodlo prawdy" — kazdy developer musi go przeczytac',
]
for r in ai_rules:
    add_bullet(r)
doc.add_page_break()

# ============================================================
# 7. BUDZETY
# ============================================================
add_heading('7. Budzety i koszty', 1)

add_heading('7.1 Wynagrodzenia (B2B netto/miesiac, Polska)', 2)
add_table(
    ['Rola', 'Stawka min', 'Stawka max', 'Rekomendacja', 'Roczny koszt'],
    [
        ['Senior AI/Backend (Tech Lead)', '25,000 PLN', '35,000 PLN', '30,000 PLN', '360,000 PLN'],
        ['Senior Data & Integration', '22,000 PLN', '30,000 PLN', '26,000 PLN', '312,000 PLN'],
        ['Mid/Senior Frontend + DevOps', '18,000 PLN', '25,000 PLN', '22,000 PLN', '264,000 PLN'],
        ['AI/ML Specialist (0.5-1.0 FTE)', '12,000 PLN', '28,000 PLN', '20,000 PLN', '240,000 PLN'],
        ['QA & Security (0.5 FTE)', '9,000 PLN', '14,000 PLN', '11,000 PLN', '132,000 PLN'],
    ]
)

add_heading('7.2 Scenariusze budzetowe', 2)
add_table(
    ['Scenariusz', 'Sklad', 'Miesiecznie', 'Rocznie', 'Uwagi'],
    [
        ['Minimum (3 osoby)', 'Tech Lead + Data Eng\n+ Frontend/DevOps', '78,000 PLN\n(~18,400 EUR)', '936,000 PLN', 'Pokrywa krytyczne potrzeby.\nBrak specjalisty RAG i QA.'],
        ['Optymalny (4 osoby)', 'Minimum + AI/ML\nSpecialist (full-time)', '98,000 PLN\n(~23,100 EUR)', '1,176,000 PLN', 'REKOMENDOWANY.\nPelne pokrycie AI quality.'],
        ['Pelny (5 osob)', 'Optymalny + QA/Security\n(part-time)', '109,000 PLN\n(~25,700 EUR)', '1,308,000 PLN', 'Produkcyjny gold standard.\nPelne quality gates.'],
    ]
)

add_heading('7.3 Dodatkowe koszty', 2)
add_table(
    ['Pozycja', 'Miesiecznie', 'Rocznie', 'Uwagi'],
    [
        ['Hetzner VPS (AX102)', '~800 PLN', '~9,600 PLN', 'Production server'],
        ['GitHub Team', '~200 PLN', '~2,400 PLN', 'Private repos, Actions'],
        ['Claude API (Anthropic)', '~3,000-5,000 PLN', '~48,000 PLN', 'Po optymalizacji kosztow'],
        ['OpenAI API (embeddings)', '~500 PLN', '~6,000 PLN', 'Embedding generation'],
        ['Monitoring (Grafana Cloud)', '~0-500 PLN', '~6,000 PLN', 'Free tier moze wystarczyc'],
        ['Narzedzia (Notion, Slack)', '~500 PLN', '~6,000 PLN', 'Opcjonalnie'],
        ['RAZEM infra', '~5,500 PLN', '~78,000 PLN', ''],
    ]
)

add_heading('7.4 ROI analizy', 2)
doc.add_paragraph('Zwrot z inwestycji w zespol:')
roi_items = [
    ('Przyspieszenie rozwoju 3-5x:', 'Masterplan V10 (Kroki A-E) w 3 miesiace zamiast 12'),
    ('Eliminacja SPOF:', 'Sebastian nie jest jedynym developerem — ciaglosc operacji'),
    ('Jakosc produkcyjna:', 'Testy, code review, monitoring = mniej awarii, mniej recznej interwencji'),
    ('Deployment 24/7:', 'Hetzner VPS = Gilbertus dziala non-stop, nie tylko na laptopie'),
    ('Skalowanie Omnius:', 'Wielu uzytkownikow (Roch, Krystian, przyszli) = wiekszy impact'),
    ('Oszczednosc czasu Sebastiana:', 'Delegacja developmentu = wiecej czasu na strategiczne decyzje'),
]
for prefix, text in roi_items:
    add_bullet(text, prefix)
doc.add_page_break()

# ============================================================
# 8. ONBOARDING
# ============================================================
add_heading('8. Proces onboardingu', 1)

doc.add_paragraph(
    'Onboarding jest KRYTYCZNY dla sukcesu zespolu. Gilbertus to zlozony system z unikalna '
    'domena biznesowa. Zle przeprowadzony onboarding moze opoznic produktywnosc o miesiace. '
    'Ponizszy proces jest zaprojektowany aby nowy developer byl produktywny w 2 tygodnie '
    'i w pelni autonomiczny w 4-6 tygodni.'
)

add_heading('8.1 Faza 0: Pre-onboarding (przed dniem 1)', 2)
pre_onboarding = [
    'Podpisanie NDA i umowy B2B',
    'Setup kont: GitHub (dodanie do repo), Discord/Slack, Anthropic API key',
    'Wyslanie "Onboarding Pack" (patrz sekcja 8.5)',
    'Przygotowanie CLAUDE.md jako "biblia projektu"',
    'Konfiguracja dostepu do srodowisk (dev, staging)',
]
for p in pre_onboarding:
    add_bullet(p)

add_heading('8.2 Tydzien 1: Orientacja i pierwszy PR', 2)
add_table(
    ['Dzien', 'Aktywnosc', 'Cel', 'Deliverable'],
    [
        ['Pn', 'Spotkanie z Sebastianem (2h)\n+ setup srodowiska', 'Kontekst biznesowy REH/REF,\nwizja Gilbertusa, docker-compose up', 'Dzialajace srodowisko\nlokalne'],
        ['Wt', 'Code walkthrough z Tech Leadem\n(lub Sebastianem jesli to hire #1)', 'Architektura: app/, scripts/,\nmcp_gilbertus/, migracje', 'Notatki architekturalne'],
        ['Sr', 'Deep dive w swoj obszar\n+ Claude Code setup', 'Zrozumienie kodu w swoim\nobszarze odpowiedzialnosci', 'Lista pytan / problemow'],
        ['Cz', 'Pierwszy "good first issue"', 'Praktyczne wejscie w codebase:\nbug fix lub maly feature', 'Pierwszy PR (draft)'],
        ['Pt', 'Code review + retro\nz zespolem', 'Feedback na PR, dostosowanie\ndo konwencji projektu', 'Merged PR #1'],
    ]
)

add_heading('8.3 Tydzien 2: Glebokie zanurzenie', 2)
week2 = [
    'Samodzielne rozwiazanie 2-3 issues ze sprint backlogu',
    'Participation in all team ceremonies (standup, planning, demo)',
    'Nauka MCP tools — uzycie gilbertus_ask, gilbertus_status, gilbertus_brief',
    'Napisanie pierwszych testow dla swojego obszaru',
    'Peer code review (review PR innego developera)',
    'Dokumentacja "co zrozumialem" — wklej do repo jako ADR',
]
for w in week2:
    add_bullet(w)

add_heading('8.4 Tygodnie 3-4: Autonomia', 2)
week34 = [
    'Samodzielne wlascicielstwo nad 1 streafem Masterplanu',
    'Prowadzenie code review dla swoich PR-ow (jako reviewer innego PR)',
    'On-call rotation (monitoring alertow Gilbertusa)',
    'Proponowanie usprawnien — pierwszy "tech proposal"',
    '30-day check-in z Sebastianem — feedback obustrony',
]
for w in week34:
    add_bullet(w)

add_heading('8.5 Onboarding Pack (dokumenty)', 2)
add_table(
    ['Dokument', 'Zawartosc', 'Format'],
    [
        ['CLAUDE.md', 'Konwencje, zasady, architektura,\nkomendy statusowe', 'Markdown w repo'],
        ['SESSION_CONTEXT.md', 'Aktualny stan systemu\n(auto-generowany)', 'Markdown w repo'],
        ['Masterplan V10', 'Roadmapa, fazy, priorytety', 'Plan file'],
        ['Architecture Overview', 'Diagram komponentow,\nflow danych, integracje', 'DO STWORZENIA\nprzez Tech Leada'],
        ['Domena biznesowa', 'REH/REF, trading energetyczny,\nstruktura spolki, ludzie', 'Spotkanie z Sebastianem\n+ dokument'],
        ['Security Guidelines', 'NDA, RODO, access control,\nsecrets management', 'Markdown w repo'],
        ['Development Setup', 'docker-compose, .env,\nvenv, pre-commit hooks', 'Markdown w repo'],
    ]
)

add_heading('8.6 "Good First Issues" — gotowe zadania onboardingowe', 2)
doc.add_paragraph('Przygotowane zadania dla nowych developerow (wg roli):')

add_table(
    ['Rola', 'Good First Issue', 'Difficulty', 'Czas'],
    [
        ['Tech Lead', 'Napisz testy dla app/api/ask.py', 'Medium', '1 dzien'],
        ['Tech Lead', 'Refactor connection pool timeout handling', 'Medium', '1 dzien'],
        ['Data Eng', 'Fix Teams sync 400 errors (skiptoken handling)', 'Hard', '2-3 dni'],
        ['Data Eng', 'Dodaj monitoring dashboard dla ingestion health', 'Medium', '1-2 dni'],
        ['Frontend/DevOps', 'Setup GitHub Actions CI (lint + test)', 'Medium', '1 dzien'],
        ['Frontend/DevOps', 'Docker-compose prod profile na Hetzner', 'Medium', '2 dni'],
        ['AI/ML', 'Evaluation pipeline: measure answer quality', 'Medium', '2 dni'],
        ['AI/ML', 'Optimize chunking strategy (compare sizes)', 'Medium', '2 dni'],
        ['QA', 'Setup pytest fixtures for DB + API tests', 'Medium', '1-2 dni'],
    ]
)
doc.add_page_break()

# ============================================================
# 9. KPI
# ============================================================
add_heading('9. KPI i metryki sukcesu', 1)

add_heading('9.1 Metryki zespolu', 2)
add_table(
    ['KPI', 'Cel (miesiac 1)', 'Cel (miesiac 3)', 'Cel (miesiac 6)'],
    [
        ['Test coverage', '5%', '30%', '60%'],
        ['PR review time', '< 24h', '< 8h', '< 4h'],
        ['Deploy frequency', '1x/tydzien', '3x/tydzien', 'Daily'],
        ['Uptime (po Hetzner)', 'N/A', '99%', '99.5%'],
        ['MTTR (mean time to repair)', '> 4h', '< 2h', '< 30 min'],
        ['API error rate', '< 5%', '< 2%', '< 1%'],
        ['Extraction coverage', '97.8%', '99%', '99.5%'],
        ['API cost / chunk', '$0.030', '$0.020', '$0.015'],
    ]
)

add_heading('9.2 Metryki produktowe (Masterplan)', 2)
add_table(
    ['Krok Masterplan', 'Status dzis', 'Cel (miesiac 3)', 'Cel (miesiac 6)'],
    [
        ['Krok E: Non-regression', '80%', '100%', '100%'],
        ['Krok A: Strategic Radar', '85%', '90%', '95%'],
        ['Krok B: Auto-actions', '60%', '75%', '85%'],
        ['Krok D: Decision journal', '40%', '70%', '85%'],
        ['Krok C: Feedback loop', '80%', '85%', '90%'],
        ['Hetzner deployment', '0%', '100%', '100%'],
        ['Omnius multi-tenant', '70%', '90%', '100%'],
    ]
)
doc.add_page_break()

# ============================================================
# 10. RYZYKA
# ============================================================
add_heading('10. Ryzyka i mitygacja', 1)

add_table(
    ['Ryzyko', 'Prawdopodobienstwo', 'Impact', 'Mitygacja'],
    [
        ['Trudnosc znalezienia\nsenior AI eng w PL', 'Wysokie', 'Krytyczny', 'Szukaj remote (EU).\nUzyj sieci Sebastiana.\nRozwaz headhuntera.'],
        ['Dlogi onboarding\n(zlozony codebase)', 'Srednie', 'Wysoki', 'Onboarding Pack gotowy\nprzed hire. Good First\nIssues przygotowane.'],
        ['Knowledge drain\n(odejscie czlonka)', 'Niskie', 'Wysoki', 'Dokumentacja ADR.\nPair programming.\nMin 2 osoby znaja\nkazdy modul.'],
        ['Scope creep\n(za duzo na raz)', 'Srednie', 'Sredni', 'Sprint planning.\nSebastian priorytetyzuje.\nMax 1-2 inicjatywy\nrownolegle.'],
        ['Konflikt wizji\n(Tech Lead vs Sebastian)', 'Niskie', 'Wysoki', 'Sebastian ma final say.\nJasne zasady w CLAUDE.md.\n30-day check-in.'],
        ['Koszty API\nwymkna sie', 'Srednie', 'Sredni', 'Cost caps (juz wdrozone).\nDaily monitoring.\nBudget alerts.'],
        ['Security breach\n(nowy developer)', 'Niskie', 'Krytyczny', 'NDA. Secrets w .env\n(nie w repo). Gitleaks\npre-commit. Audit trail.'],
    ]
)
doc.add_page_break()

# ============================================================
# 11. HARMONOGRAM
# ============================================================
add_heading('11. Harmonogram budowy zespolu', 1)

add_table(
    ['Okres', 'Aktywnosc', 'Milestone'],
    [
        ['Tydzien 1-2\n(kwiecien)', 'Przygotowanie:\n- Onboarding Pack\n- Job descriptions\n- Good First Issues\n- Architecture doc', 'Gotowy hiring pipeline'],
        ['Tydzien 3-4\n(kwiecien)', 'Rekrutacja #1:\nSenior AI/Backend (Tech Lead)\n- LinkedIn, NoFluffJobs, JustJoinIT\n- Siec kontaktow\n- Headhunter (opcja)', 'Kandydat #1 wybrany'],
        ['Tydzien 5-6\n(maj)', 'Onboarding Tech Lead\n+ rekrutacja #2:\nData & Integration Eng', 'Tech Lead produktywny\nKandydat #2 wybrany'],
        ['Tydzien 7-8\n(maj)', 'Tech Lead: CI/CD + testy\nOnboarding Data Eng\n+ rekrutacja #3', 'CI/CD live\nData Eng onboarded\nKandydat #3 wybrany'],
        ['Tydzien 9-12\n(czerwiec)', '3-osobowy team:\n- Hetzner deployment\n- Teams sync fix\n- Test coverage 20%+\nRekrutacja #4 (AI/ML)', 'PRODUCTION DEPLOY\nna Hetzner'],
        ['Tydzien 13-16\n(lipiec)', '4-osobowy team:\n- Masterplan Kroki A-E\n- RAG optimization\n- Omnius UX\n- Opcja: hire #5 QA', 'Krok E: 100%\nKrok A: 90%\nTest coverage 30%+'],
        ['Tydzien 17-24\n(sie-wrz)', 'Pelny team:\n- Zakonczenie Masterplanu\n- Self-healing 95%\n- Multi-tenant Omnius', 'PIESC ZAMKNIETA\n92% capability'],
    ]
)

doc.add_paragraph()
doc.add_paragraph(
    'Przy prawidlowej egzekucji, Gilbertus osiagnie pelna operacyjnosc produkcyjna '
    '(cel "Piesc Zamknieta" z Masterplan V10) w Q3 2026 — 6 miesiecy od sformowania zespolu.'
)

doc.add_page_break()

# ============================================================
# APPENDIX
# ============================================================
add_heading('Zalacznik A: Gdzie szukac ludzi', 1)

add_table(
    ['Kanal', 'Typ', 'Dla jakich rol', 'Uwagi'],
    [
        ['justjoin.it', 'Job board (PL)', 'Wszystkie', 'Najwiekszy IT job board w Polsce'],
        ['nofluffjobs.com', 'Job board (PL)', 'Wszystkie', 'Transparentne widleki, B2B friendly'],
        ['LinkedIn', 'Direct search', 'Senior roles', 'Boolean search + InMail'],
        ['Bulldogjob', 'Job board (PL)', 'Mid/Senior', 'Dobry zasieg w PL'],
        ['Anthropic Discord', 'Community', 'AI/ML Specialist', 'Claude-focused developers'],
        ['Headhunter (IT)', 'Rekrutacja', 'Tech Lead', '15-20% fee, ale szybciej'],
        ['Siec Sebastiana', 'Referral', 'Wszystkie', 'Najlepsza jakosc kandydatow'],
        ['GitHub/OSS', 'Community', 'AI/ML, Backend', 'Sprawdz kontrybutorów LangChain/Qdrant'],
    ]
)

add_heading('Zalacznik B: Template ogloszenia (Tech Lead)', 1)
doc.add_paragraph(
    'SENIOR AI/BACKEND ENGINEER (TECH LEAD) — Prywatny projekt AI\n\n'
    'Szukamy doswiadczonego inzyniera Python/AI do roli Tech Leada w innowacyjnym '
    'projekcie AI mentat — prywatnym systemie inteligencji biznesowej zbudowanym na '
    'Claude (Anthropic), FastAPI, PostgreSQL i Qdrant.\n\n'
    'Projekt obejmuje: 135k LoC, 44 narzedzia MCP, 189 endpointow API, integracje z '
    'Microsoft Graph, WhatsApp, audio transcription. Stack: Python 3.12, FastAPI, '
    'LangGraph, Qdrant, Docker, PostgreSQL 16.\n\n'
    'Wymagamy: 5+ lat Python, doswiadczenie z LLM APIs w produkcji, PostgreSQL advanced, '
    'Docker/Linux, code review leadership. Mile widziane: LangChain/LangGraph, Qdrant, '
    'Microsoft Graph API.\n\n'
    'Forma: B2B, remote-first (spotkania w Warszawie 1-2x/miesiac). '
    'Stawka: 25,000-35,000 PLN netto/miesiac.'
)

# Save
output_path = sys.argv[1] if len(sys.argv) > 1 else '/mnt/c/Users/jablo/Desktop/Gilbertus_Team_Building_Strategy.docx'
doc.save(output_path)
print(f"Saved to: {output_path}")
