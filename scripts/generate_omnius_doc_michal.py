"""Generate Omnius onboarding document for Michał Schulte."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# Styles
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

# ── Title ──────────────────────────────────────────────────────────────────

title = doc.add_heading('Omnius REF — Dokumentacja techniczna', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Przewodnik wdrożeniowy dla Michała Schulty\nIT Support Specialist, Respect Energy Fuels')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.add_run('Data: ').bold = True
meta.add_run('27 marca 2026')
meta.add_run('    |    ')
meta.add_run('Wersja: ').bold = True
meta.add_run('1.0')
meta.add_run('    |    ')
meta.add_run('Poufność: ').bold = True
meta.add_run('Wewnętrzna')

doc.add_page_break()

# ── 1. Co to jest Omnius ───────────────────────────────────────────────────

doc.add_heading('1. Co to jest Omnius', level=1)

doc.add_paragraph(
    'Omnius to korporacyjny asystent AI dedykowany dla Respect Energy Fuels (REF). '
    'Jest to inteligentny system, który indeksuje dane firmowe z Microsoft 365 '
    '(Teams, SharePoint, email, kalendarz), przetwarza je i udostępnia zarządowi '
    'w formie odpowiedzi na pytania, analiz i raportów.'
)

doc.add_paragraph(
    'Omnius jest częścią ekosystemu Gilbertus Albans — prywatnego systemu AI '
    'Sebastiana Jabłońskiego. Gilbertus jest "mózgiem" całego systemu i kontroluje '
    'Omniusy w poszczególnych spółkach. Dane firmowe REF nigdy nie opuszczają '
    'infrastruktury REF — Gilbertus łączy się z Omniusem przez zabezpieczone API.'
)

doc.add_heading('Kluczowe cechy:', level=2)

items = [
    ('Sztuczna inteligencja', 'Omnius używa modeli Claude (Anthropic) i OpenAI do analizy danych i odpowiadania na pytania.'),
    ('Microsoft 365 integration', 'Automatycznie pobiera dane z Teams, SharePoint, email i kalendarza REF.'),
    ('RBAC (Role-Based Access Control)', '7 poziomów dostępu — od CEO po specjalistę. Każdy widzi tylko to, do czego ma uprawnienia.'),
    ('Governance', 'CEO i zarząd mogą rozwijać Omniusa, ale nie mogą usuwać funkcjonalności ani zmniejszać zakresu danych.'),
    ('Teams Bot', 'Zarząd korzysta z Omniusa bezpośrednio w Microsoft Teams — pisze pytania, dostaje odpowiedzi.'),
    ('Web interface', 'Alternatywnie: przeglądarkowy chat na dedykowanej stronie z logowaniem Azure AD.'),
    ('Plaud Pin S', 'CEO i zarząd mogą mieć urządzenia do nagrywania spotkań — transkrypcje automatycznie trafiają do Omniusa.'),
]

for title_text, desc in items:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(title_text + ': ').bold = True
    p.add_run(desc)

# ── 2. Architektura ────────────────────────────────────────────────────────

doc.add_page_break()
doc.add_heading('2. Architektura systemu', level=1)

doc.add_paragraph(
    'System działa w architekturze hybrydowej — Gilbertus na serwerze Sebastiana (Hetzner), '
    'Omnius REF na infrastrukturze Azure REF. Dane firmowe zostają w firmie.'
)

doc.add_heading('Schemat:', level=2)

arch = doc.add_paragraph()
arch.paragraph_format.space_before = Pt(6)
arch_text = """
┌─────────────────────────┐         ┌─────────────────────────┐
│   HETZNER (Sebastian)   │         │      AZURE REF          │
│                         │         │                         │
│  Gilbertus API          │  HTTPS  │  Omnius REF API         │
│  ├── Dane prywatne      │────────>│  ├── /ask (dane REF)    │
│  ├── Cross-company      │  API    │  ├── /commands/*        │
│  └── WhatsApp           │  Key    │  ├── Teams Bot          │
│                         │         │  ├── Web chat           │
│  Gilbertus DB           │         │  └── /metrics           │
│  (PostgreSQL)           │         │                         │
│                         │         │  Omnius DB (PostgreSQL) │
│                         │         │  Qdrant (vector search) │
│                         │         │                         │
│                         │         │  Azure AD SSO           │
│                         │         │  ├── Krystian (CEO)     │
│                         │         │  ├── Edgar (zarząd)     │
│                         │         │  ├── Witold (zarząd)    │
│                         │         │  └── Michał (operator)  │
└─────────────────────────┘         └─────────────────────────┘
"""
run = arch.add_run(arch_text)
run.font.name = 'Consolas'
run.font.size = Pt(8)

doc.add_heading('Kluczowe zasady:', level=2)

rules = [
    'Dane REF NIGDY nie opuszczają Azure REF.',
    'Gilbertus → Omnius: jednokierunkowy flow (push config, query).',
    'Omnius NIE woła Gilbertusa — nie ma dostępu do danych prywatnych Sebastiana.',
    'Deploy kodu: Gilbertus/Sebastian pushuje przez SSH + rsync.',
]
for r in rules:
    doc.add_paragraph(r, style='List Bullet')

# ── 3. RBAC ────────────────────────────────────────────────────────────────

doc.add_page_break()
doc.add_heading('3. Role i uprawnienia (RBAC)', level=1)

doc.add_paragraph(
    'System ma 7 poziomów dostępu. Każda rola ma zdefiniowane uprawnienia zarówno '
    'do danych jak i do rozwoju systemu.'
)

table = doc.add_table(rows=8, cols=4)
table.style = 'Light Shading Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['Rola', 'Level', 'Dostęp do danych', 'Zakres developmentu']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

roles = [
    ('gilbertus_admin', '99', 'Wszystko', 'Pełna kontrola — push config, reguły, deploy'),
    ('operator (Ty)', '70', 'ZERO danych biznesowych', 'Import, credentials, sync, infra, dev tasks'),
    ('ceo', '60', 'Wszystkie dane firmowe', 'Zarządzanie użytkownikami, config, queries'),
    ('board', '50', 'Dane firmowe (bez CEO-private)', 'Zarządzanie users poniżej director'),
    ('director', '40', 'Dane departamentu', 'Config widoków w departamencie'),
    ('manager', '30', 'Dane zespołu', 'Config własnych widoków'),
    ('specialist', '20', 'Własne taski', 'Preferencje osobiste'),
]

for i, (role, level, data, dev) in enumerate(roles):
    row = table.rows[i + 1]
    row.cells[0].text = role
    row.cells[1].text = level
    row.cells[2].text = data
    row.cells[3].text = dev

doc.add_paragraph()

doc.add_heading('Twoja rola: operator', level=2)

doc.add_paragraph(
    'Jako operator jesteś "human interface" — wykonujesz zadania, których AI nie może '
    'zrobić sam (np. podanie credentials, consent w Azure AD, konfiguracja infra). '
    'NIE masz dostępu do danych biznesowych firmy — nie możesz czytać emaili, '
    'dokumentów ani transkrypcji.'
)

op_items = [
    'Dostajesz zadania od Gilbertusa w Teams Bot (np. "podaj credentials SharePoint")',
    'Odpowiadasz: "done #N [wynik]" gdy zadanie jest wykonane',
    'Wspierasz zarząd w korzystaniu z Omniusa (support techniczny)',
    'Importujesz dane, konfigurujesz synchronizacje',
    'NIE widzisz treści pytań i odpowiedzi zarządu',
]
for item in op_items:
    doc.add_paragraph(item, style='List Bullet')

# ── 4. Wartość dla Krystiana i zarządu ─────────────────────────────────────

doc.add_page_break()
doc.add_heading('4. Wartość dla Krystiana i zarządu', level=1)

doc.add_paragraph(
    'Omnius daje zarządowi REF bezprecedensowy dostęp do wiedzy firmowej. '
    'Zamiast szukać informacji w dziesiątkach miejsc, zarząd pyta Omniusa.'
)

doc.add_heading('Co może Krystian (CEO):', level=2)

ceo_items = [
    ('Pytania o dane firmowe', 'Np. "Jakie kontrakty gazowe mamy na Q2?" — Omnius przeszukuje Teams, email, SharePoint i odpowiada z kontekstem.'),
    ('Analiza dokumentów', 'Upload raportu → Omnius analizuje, wyciąga wnioski, porównuje z historycznymi danymi.'),
    ('Nagrywanie spotkań (Plaud)', 'Plaud Pin S automatycznie transkrybuje i importuje spotkania. Krystian może potem zapytać: "Co ustaliliśmy na spotkaniu z NOFAR?"'),
    ('Proponowanie nowych funkcji', 'Pisze w Teams: "nowa funkcjonalność: dashboard z KPI handlowymi" — Omnius ocenia wartość i tworzy task.'),
    ('Zarządzanie użytkownikami', 'Może dodawać nowych użytkowników (dyrektorów, managerów) z odpowiednimi uprawnieniami.'),
    ('Pełna prywatność', 'Jego pytania i odpowiedzi nie są widoczne dla nikogo poza nim i Sebastianem.'),
]

for title_text, desc in ceo_items:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(title_text + ': ').bold = True
    p.add_run(desc)

doc.add_heading('Co mogą Edgar i Witold (zarząd):', level=2)

board_items = [
    'Wszystko co CEO, ale BEZ dostępu do dokumentów oznaczonych "ceo_only"',
    'Mogą proponować nowe funkcjonalności (z oceną wartości przez Omniusa)',
    'Mogą zarządzać użytkownikami poniżej poziomu director',
    'Mają własne nagrania Plaud (jeśli dostaną urządzenie)',
]
for item in board_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Governance — czego zarząd NIE MOŻE:', level=2)

forbidden = [
    'Usuwać istniejących funkcjonalności',
    'Zmniejszać zakresu danych dostępnych dla Omniusa',
    'Wyłączać synchronizacji lub cronów',
    'Modyfikować uprawnień (RBAC) lub reguł governance',
    'Usuwać źródeł danych',
]
for item in forbidden:
    p = doc.add_paragraph(item, style='List Bullet')
    p.runs[0].font.color.rgb = RGBColor(180, 0, 0)

# ── 5. Co musisz dostarczyć ────────────────────────────────────────────────

doc.add_page_break()
doc.add_heading('5. Co musisz dostarczyć / zrobić', level=1)

doc.add_paragraph(
    'Poniżej lista rzeczy potrzebnych do uruchomienia Omniusa REF. '
    'Każdy punkt to osobny task — realizuj je po kolei.'
)

doc.add_heading('5.1. Azure VM', level=2)

doc.add_paragraph('Potrzebna maszyna wirtualna na Azure (subskrypcja REF):')

vm_items = [
    'Ubuntu 22.04 LTS lub nowszy',
    'Min. 4 vCPU, 8 GB RAM, 100 GB SSD',
    'Publiczny IP z DNS (np. omnius-ref.re-fuels.com)',
    'Porty otwarte: 443 (HTTPS), 22 (SSH — TYLKO z IP Hetznera: podaj mi IP)',
    'Docker + Docker Compose zainstalowane',
]
for item in vm_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('5.2. Azure AD App Registration', level=2)

doc.add_paragraph('Zarejestruj aplikację w Azure AD (tenant REF):')

ad_steps = [
    ('Nazwa aplikacji', 'Omnius REF'),
    ('Redirect URI', 'https://omnius-ref.re-fuels.com/auth/callback (type: SPA)'),
    ('API Permissions (Application)', 'Mail.Read, Chat.Read.All, ChannelMessage.Read.All, Sites.Read.All, Calendars.Read, User.Read.All'),
    ('API Permissions (Delegated)', 'openid, profile, email'),
    ('Client Secret', 'Wygeneruj i przekaż mi bezpiecznie (NIE emailem — użyj szyfrowanego kanału)'),
    ('Admin Consent', 'Potrzebna zgoda admina na application permissions'),
]

for step, desc in ad_steps:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(step + ': ').bold = True
    p.add_run(desc)

doc.add_heading('5.3. Teams Bot Registration', level=2)

doc.add_paragraph('W Azure Portal → Bot Services:')

bot_steps = [
    'Utwórz Bot Channel Registration',
    'Messaging endpoint: https://omnius-ref.re-fuels.com/teams/webhook',
    'Włącz kanał Microsoft Teams',
    'Zainstaluj bota w tenancie REF (Teams Admin Center)',
    'Przekaż mi: App ID + App Secret',
]
for item in bot_steps:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('5.4. SSH Key', level=2)

doc.add_paragraph(
    'Potrzebuję dostępu SSH do Azure VM żeby deployować kod. '
    'Stwórz użytkownika "omnius-deploy" i przekaż mi:'
)

ssh_items = [
    'Adres serwera (IP lub DNS)',
    'Port SSH (jeśli nie 22)',
    'Ja wygeneruję klucz SSH — Ty dodasz public key do authorized_keys',
]
for item in ssh_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('5.5. Credentials do przekazania', level=2)

doc.add_paragraph(
    'Wszystkie credentials przekaż BEZPIECZNIE — przez szyfrowany kanał, '
    'nigdy emailem ani Teamsem w plain text.'
)

creds_table = doc.add_table(rows=8, cols=3)
creds_table.style = 'Light Shading Accent 1'

creds_headers = ['Credential', 'Źródło', 'Jak przekazać']
for i, h in enumerate(creds_headers):
    creds_table.rows[0].cells[i].text = h
    creds_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

creds = [
    ('Azure Tenant ID', 'Azure Portal → Azure AD → Overview', 'Może być plain text'),
    ('Azure Client ID', 'App Registration → Overview', 'Może być plain text'),
    ('Azure Client Secret', 'App Registration → Certificates & Secrets', 'SZYFROWANY kanał'),
    ('Teams Bot App ID', 'Bot Channel Registration', 'Może być plain text'),
    ('Teams Bot Secret', 'Bot Channel Registration → Settings', 'SZYFROWANY kanał'),
    ('SSH host + port', 'Azure VM', 'Może być plain text'),
    ('Anthropic API Key', 'Dostarczę ja (Sebastian)', 'N/A'),
]

for i, (cred, src, how) in enumerate(creds):
    row = creds_table.rows[i + 1]
    row.cells[0].text = cred
    row.cells[1].text = src
    row.cells[2].text = how

# ── 6. Jak działa deploy ───────────────────────────────────────────────────

doc.add_page_break()
doc.add_heading('6. Jak działa deploy', level=1)

doc.add_paragraph(
    'Po dostarczeniu infrastruktury i credentials, deploy wygląda tak:'
)

deploy_steps = [
    'Sebastian/Gilbertus edytuje kod w omnius/',
    'Gilbertus wykonuje: bash scripts/deploy_omnius.sh ref',
    'Skrypt synchronizuje kod przez rsync + SSH',
    'Na serwerze Azure: docker compose up --build',
    'Automatyczny health check — jeśli OK, deploy zakończony',
    'Gilbertus potwierdza sukces lub raportuje problem',
]

for i, step in enumerate(deploy_steps):
    p = doc.add_paragraph(style='List Number')
    p.text = step

doc.add_paragraph(
    'Nie musisz nic robić przy deploy — to się dzieje automatycznie. '
    'Twoja rola to utrzymanie infrastruktury (VM, sieć, certyfikaty).'
)

# ── 7. Technologie ────────────────────────────────────────────────────────

doc.add_heading('7. Stack technologiczny', level=1)

tech_table = doc.add_table(rows=11, cols=2)
tech_table.style = 'Light Shading Accent 1'
tech_table.rows[0].cells[0].text = 'Komponent'
tech_table.rows[0].cells[1].text = 'Technologia'
for cell in tech_table.rows[0].cells:
    cell.paragraphs[0].runs[0].bold = True

techs = [
    ('Backend', 'Python 3.12 + FastAPI'),
    ('Baza danych', 'PostgreSQL 16'),
    ('Vector search', 'Qdrant'),
    ('AI Models', 'Claude Haiku (Anthropic) + OpenAI embeddings'),
    ('Auth', 'Azure AD SSO (JWT) + API keys'),
    ('Bot', 'Microsoft Bot Framework → Teams'),
    ('Deploy', 'Docker Compose + rsync + SSH'),
    ('Monitoring', 'Prometheus metrics (/metrics endpoint)'),
    ('Secrets', 'Azure Key Vault (produkcja) / .env (dev)'),
    ('Migrations', 'Alembic + raw SQL'),
]

for i, (comp, tech) in enumerate(techs):
    tech_table.rows[i + 1].cells[0].text = comp
    tech_table.rows[i + 1].cells[1].text = tech

# ── 8. Kontakt ─────────────────────────────────────────────────────────────

doc.add_heading('8. Kontakt i eskalacja', level=1)

doc.add_paragraph(
    'W przypadku problemów technicznych lub pytań:'
)

contact_items = [
    ('Sebastian Jabłoński', 'Właściciel projektu, decyzje architektoniczne — WhatsApp/Teams'),
    ('Gilbertus (AI)', 'Przydziela Ci zadania automatycznie przez Teams Bot — odpowiadaj "done #N"'),
    ('Krystian Juchacz', 'CEO REF, pierwszy użytkownik — feedback o UX i wartości'),
]

for name, desc in contact_items:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(name + ': ').bold = True
    p.add_run(desc)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('WAŻNE: ').bold = True
p.add_run(
    'Nie masz dostępu do danych biznesowych w Omniusie. Nie próbuj odpytywać '
    'endpointu /ask ani czytać treści dokumentów. Twoja rola to infrastruktura '
    'i development, nie analiza danych.'
)

# ── Save ───────────────────────────────────────────────────────────────────

output_path = '/mnt/c/Users/jablo/Desktop/Omnius_REF_Onboarding_Michal_Schulte.docx'
doc.save(output_path)
print(f'Saved to: {output_path}')
