"""Convert ARCHITECTURE.md to DOCX and save to Windows Desktop."""
import re
from docx import Document
from docx.shared import Pt

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

with open('docs/technical/ARCHITECTURE.md', 'r') as f:
    lines = f.readlines()

in_code_block = False

for line in lines:
    line = line.rstrip()

    # Code block toggle
    if line.startswith('```'):
        in_code_block = not in_code_block
        continue

    if in_code_block:
        p = doc.add_paragraph(line)
        p.style = doc.styles['Normal']
        for run in p.runs:
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
        continue

    # Headers
    if line.startswith('# '):
        doc.add_heading(line[2:], level=1)
    elif line.startswith('## '):
        doc.add_heading(line[3:], level=2)
    elif line.startswith('### '):
        doc.add_heading(line[4:], level=3)
    elif line.startswith('#### '):
        doc.add_heading(line[5:], level=4)
    # Bullet points
    elif line.startswith('- '):
        text = line[2:]
        text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
        text = re.sub(r'`([^`]+)`', r'\1', text)
        doc.add_paragraph(text, style='List Bullet')
    # Numbered lists
    elif re.match(r'^\d+\.', line):
        text = re.sub(r'^\d+\.\s*', '', line)
        text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
        text = re.sub(r'`([^`]+)`', r'\1', text)
        doc.add_paragraph(text, style='List Number')
    # Table separator
    elif line.startswith('|--') or line.startswith('| --'):
        continue
    # Table rows
    elif line.startswith('|'):
        cells = [c.strip() for c in line.split('|')[1:-1]]
        if cells:
            doc.add_paragraph('  |  '.join(cells))
    # Empty lines
    elif not line.strip():
        doc.add_paragraph('')
    # Regular text
    else:
        text = re.sub(r'\*\*([^*]+)\*\*', r'\1', line)
        text = re.sub(r'`([^`]+)`', r'\1', text)
        doc.add_paragraph(text)

output = '/mnt/c/Users/jablo/Desktop/Gilbertus_Technical_Documentation.docx'
doc.save(output)
print(f'Saved: {output}')
